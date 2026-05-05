import os, sys
import hashlib

# Keep Socket.IO on threading mode to avoid deprecated/fragile eventlet runtime on Gunicorn.
ASYNC_MODE = 'threading'

import uuid
import re
import concurrent.futures

import re
import io
import json
import time
import random
import datetime
import math
import smtplib
import traceback
import resend
from types import SimpleNamespace

# Compatibility shim for libraries that incorrectly call datetime.utcnow() on the module.
# This must run before importing third-party libraries to catch import-time calls.
if not hasattr(datetime, 'utcnow'):
    datetime.utcnow = datetime.datetime.utcnow

from sqlalchemy.exc import IntegrityError
from sqlalchemy import text as sql_text
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import timedelta
from dotenv import load_dotenv, dotenv_values
from flask import Flask, request, jsonify, Response, stream_with_context, redirect
from flask_cors import CORS
from flask_socketio import SocketIO, join_room, leave_room, emit as socket_emit
from werkzeug.exceptions import HTTPException
from flask_jwt_extended import (
    JWTManager, create_access_token, get_jwt_identity, jwt_required, verify_jwt_in_request
)
from passlib.hash import pbkdf2_sha256
from google import genai as google_genai
from google.genai import types as google_genai_types
from models import db, History, Answer, User, Conversation, ConversationSummary, MemoryItem, MemoryNode, MemoryEdge, Snippet, PasswordResetToken, UserFollow, Notification, Favorite, Project, ProjectFile, UserBadge, SharedSession, XPEvent, CollaborationReview, CollaborationComment, TokenBalance, TokenTransaction, TokenPackage, TokenPurchase, ApiKey, VSCodeLoginState, VSCodeOTP, PostLike, AnswerLike, NotificationRead, NotificationHidden, Feedback, FeedbackDetail, UserTheme
from anthropic import Anthropic, APIError
from openai import OpenAI
import stripe
import iyzipay
from utils.language_detector import LanguageDetector
from utils.model_router import ModelRouter
from utils.standardizer import CodeStandardizer
from utils.github_parser import GitHubParser
from utils.timeout_utils import to_gemini_timeout

# Global registry for cancelled requests
CANCELLED_REQUESTS = {}

def is_request_cancelled(request_id):
    return CANCELLED_REQUESTS.get(request_id, False)
from utils.memory_utils import (
    build_minimum_continuation_capsule,
    build_memory_retrieval_plan,
    build_structured_memory_capsule,
    detect_memory_conflicts,
    extract_memory_candidates,
)
from services.lifecycle_orchestrator import start_worker, LifecycleOrchestrator
from services.agent_runtime import AgentToolRuntime, run_agent_turn, stream_text_chunks, AgentAbortException


def _classify_openai_error(error: Exception) -> str:
    message = str(error or '').strip()
    lowered = message.lower()

    if 'authentication' in lowered or 'api key' in lowered or 'unauthorized' in lowered:
        return 'authentication_or_key_issue'
    if 'insufficient_quota' in lowered or 'quota' in lowered or 'rate limit' in lowered or '429' in lowered:
        return 'quota_or_rate_limit'
    if 'model' in lowered and ('not found' in lowered or 'does not exist' in lowered or 'unsupported' in lowered):
        return 'model_not_available'
    if 'timeout' in lowered or 'timed out' in lowered:
        return 'timeout'
    if 'connection' in lowered or 'network' in lowered or 'temporary failure' in lowered:
        return 'network_or_transport'

    status_code = getattr(error, 'status_code', None)
    if status_code == 401 or status_code == 403:
        return 'authentication_or_permission_issue'
    if status_code == 404:
        return 'model_not_available'
    if status_code == 429:
        return 'quota_or_rate_limit'
    if status_code and int(status_code) >= 500:
        return 'provider_server_error'

    return 'unknown_openai_error'


def _normalize_gemini_model_name(model_name):
    if not model_name:
        return model_name
    # Ensure the model string uses the full resource path expected by the
    # Google GenAI client (e.g. "models/gemini-2.5-flash"). Some places in
    # config may omit the "models/" prefix while others include it; always
    # normalize to include the prefix to avoid mismatches that cause some
    # Gemini variants to fail when called.
    return model_name if model_name.startswith('models/') else f"models/{model_name}"


def _utcnow():
    return datetime.datetime.now(datetime.timezone.utc).replace(tzinfo=None)


def _get_iyzico_frontend_base_url(origin: str | None = None) -> str:
    # Use origin if provided, otherwise fallback to env, otherwise use the current request host
    default_url = request.host_url.rstrip('/') if request else 'http://localhost'
    return (origin or os.getenv('FRONTEND_URL') or os.getenv('APP_FRONTEND_URL') or default_url).rstrip('/')


def _clean_gemma_output(text: str, question: str = None) -> str:
    import re
    if not text:
        return ""

    # 1. Block-based extraction (The most robust for Gemma)
    # If the model used <answer> or [RESPONSE] tags, we ONLY take what's inside.
    answer_match = re.search(r'<(answer|response)>(.*?)</\1>', text, re.DOTALL | re.IGNORECASE)
    if answer_match:
        return answer_match.group(2).strip()
    
    # Heuristic: Find the LAST quoted line that looks like a response
    # Often Gemma outputs: * "Selam! How can I help?" at the end of its analysis.
    quotes = re.findall(r'^\s*(\*|\-)?\s*["\'](.+?)["\']\s*$', text, re.MULTILINE)
    if quotes:
        # Take the last one, it's usually the final draft
        return quotes[-1][1].strip()
    
    # 2. Block tags removal (legacy)
    text = re.sub(r'<(think|thought|thinking|reasoning|analysis)>.*?</\1>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<(think|thought|thinking|reasoning|analysis)>.*', '', text, flags=re.DOTALL | re.IGNORECASE)
    
    # Remove the tags themselves if they are dangling
    text = re.sub(r'</?(answer|think|thought|thinking|reasoning|analysis)>', '', text, flags=re.IGNORECASE)

    # 3. Individual label lines (Pure metadata)
    # This surgical regex catches: "* **Label:** content" and replaces with "content"
    # If content is empty, the line effectively becomes empty.
    # Added more characters to handle "1-2 short sentences?", "Natural?" etc.
    label_pattern = r'(?i)^\s*(\*|\-)?\s*(\*\*|\*)?([a-z0-9\/\?\(\)\-\+ ]{2,50}):(\*\*|\*)?\s*(.*)'
    
    # Common words found in Gemma's echoed instructions/checklists
    forbidden_keywords = [
        'metadata', 'persona', 'thinking process', 'reasoning steps', 
        'short sentences', 'direct answer', 'echoing instructions',
        'helpful ai assistant', 'user profile', 'expertise level',
        'natural conversation', 'internal analysis', 'thinking block'
    ]
    
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        lower_line = line.lower()
        # Keyword filtering: if the line looks like an echoed instruction
        if any(kw in lower_line for kw in forbidden_keywords):
            continue
            
        # Echo filtering: if the line looks like an echo of the user's question
        if question and question.lower().strip() in lower_line:
            # But allow it if the actual answer IS the question (rare)
            if len(lower_line) > len(question) + 10: 
                continue

        # Try the generic label pattern
        match = re.match(label_pattern, line)
        if match:
            # We found a label. Keep only the content after it.
            # If the content is just "Yes", "No", "True", "False" (common in Gemma junk), discard the whole line.
            content = match.group(5).strip()
            if content.lower() in ['yes', 'no', 'true', 'false', 'yes.', 'no.', 'turkish', 'english', 'natural']:
                continue
            else:
                cleaned_lines.append(content)
        else:
            # Not a label line, keep as is
            cleaned_lines.append(line)
    
    text = '\n'.join(cleaned_lines)

    # Legacy specific patterns for multi-line or complex cases
    line_patterns = [
        r'(?i)^\s*#+\s*thinking.*',
        r'(?i)^\s*#+\s*thought.*',
        r'(?i)^\s*#+\s*reasoning.*',
        r'(?i)^\s*#+\s*plan.*',
        r'(?i)^\s*#+\s*analysis.*',
        r'(?i)^\s*#+\s*strategy.*',
        r'(?i)^\s*(\*|\-)?\s*(\*\*|\*)?no internal (analysis|reasoning|thinking).*:?(\*\*|\*)?.*',
        r'(?i)^\s*(\*|\-)?\s*(\*\*|\*)?no (reasoning|thinking) blocks?.*',
        r'(?i)^\s*(\*|\-)?\s*(\*\*|\*)?no markdown labels?.*',
        r'(?i)^\s*(\*|\-)?\s*(\*\*|\*)?no image generation.*',
        r'(?i)^\s*(\*|\-)?\s*(\*\*|\*)?confidence score:?(\*\*|\*)?.*',
    ]
    for pattern in line_patterns:
        text = re.sub(pattern, '', text, flags=re.MULTILINE)

    # 4. Final Polish: deduplication and quote stripping
    text = re.sub(r'\n{3,}', '\n\n', text).strip()
    # Deduplicate back-to-back sentences (common in some model outputs)
    text = re.sub(r'^"(.+?)"\s+\1\s*$', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'^"(.+?)"\s+"\1"\s*$', r'\1', text, flags=re.DOTALL)
    # Strip surrounding quotes if the whole reply is quoted
    text = re.sub(r'^"(.+)"$', r'\1', text, flags=re.DOTALL)
    
    return text.strip()

def _extract_gemini_text(response_obj, question: str = None):
    text = getattr(response_obj, 'text', None)
    if isinstance(text, str) and text:
        return _clean_gemma_output(text, question)

    # Best-effort fallback for non-text response shapes.
    candidates = getattr(response_obj, 'candidates', None)
    if candidates:
        parts = getattr(getattr(candidates[0], 'content', None), 'parts', None)
        if parts:
            chunks = []
            for part in parts:
                ptext = getattr(part, 'text', None)
                if ptext:
                    chunks.append(ptext)
            if chunks:
                return _clean_gemma_output(''.join(chunks), question)

    return ''


class _GeminiCompatModel:
    def __init__(self, client, model_name):
        self._client = client
        self._model_name = _normalize_gemini_model_name(model_name)

    def _normalize_content_part(self, part):
        if isinstance(part, dict) and part.get("mime_type") and part.get("data") is not None:
            return google_genai_types.Part.from_bytes(
                data=part.get("data"),
                mime_type=str(part.get("mime_type") or "application/octet-stream"),
            )
        if isinstance(part, list):
            return [self._normalize_content_part(item) for item in part]
        return part

    def _normalize_contents(self, contents):
        if isinstance(contents, list):
            return [self._normalize_content_part(item) for item in contents]
        return self._normalize_content_part(contents)

    def generate_content(self, contents, stream=False, request_options=None, question=None, system_instruction=None):
        # request_options is accepted for backward compatibility with old SDK call sites.
        # It is assumed to be in seconds and is converted to milliseconds with a 10s floor.
        timeout_sec = (request_options or {}).get('timeout') if isinstance(request_options, dict) else None
        gemini_timeout = to_gemini_timeout(timeout_sec)
        normalized_contents = self._normalize_contents(contents)
        
        config = {'http_options': {'timeout': gemini_timeout}}
        if system_instruction:
            config['system_instruction'] = system_instruction


        if stream:
            def _iter_stream():
                stream_iter = self._client.models.generate_content_stream(
                    model=self._model_name,
                    contents=normalized_contents,
                    config=config,
                )
                
                buffer = ""
                in_answer_mode = False
                yielded_any = False
                full_raw_for_fallback = ""
                
                for item in stream_iter:
                    # Get raw text
                    raw_chunk = ""
                    candidates = getattr(item, 'candidates', None)
                    if candidates:
                        parts = getattr(getattr(candidates[0], 'content', None), 'parts', None)
                        if parts:
                            raw_chunk = "".join([getattr(p, 'text', '') for p in parts])
                    else:
                        raw_chunk = getattr(item, 'text', '')

                    if not raw_chunk:
                        continue
                    
                    full_raw_for_fallback += raw_chunk

                    # Aggressive <answer> tag filtering
                    chunk_lower = raw_chunk.lower()
                    
                    if not in_answer_mode:
                        if '<answer>' in chunk_lower:
                            # Start yielding from after the tag
                            in_answer_mode = True
                            parts = re.split(r'<answer>', raw_chunk, flags=re.IGNORECASE)
                            raw_chunk = parts[1] if len(parts) > 1 else ""
                        else:
                            # Still haven't found the answer tag, skip this chunk
                            continue
                    
                    if in_answer_mode:
                        if '</answer>' in chunk_lower:
                            # Final part of the answer
                            parts = re.split(r'</answer>', raw_chunk, flags=re.IGNORECASE)
                            final_part = parts[0]
                            if final_part:
                                yield SimpleNamespace(text=_clean_gemma_output(final_part, question))
                                yielded_any = True
                            break # We are done
                        else:
                            # Mid-answer content
                            if raw_chunk:
                                yield SimpleNamespace(text=_clean_gemma_output(raw_chunk, question))
                                yielded_any = True
                
                # Fallback if no output was yielded (model ignored <answer> tags)
                if not yielded_any and full_raw_for_fallback:
                    yield SimpleNamespace(text=_clean_gemma_output(full_raw_for_fallback, question))

                # Some Gemini variants expose no usable text through the stream API.
                # Retry once with the non-stream response shape before giving up.
                if not yielded_any:
                    response = self._client.models.generate_content(
                        model=self._model_name,
                        contents=normalized_contents,
                        config=config,
                    )
                    final_text = _extract_gemini_text(response, question)
                    if final_text:
                        print(f"Gemini stream fallback succeeded for {self._model_name} (len={len(final_text)})")
                        yield SimpleNamespace(text=final_text)
                    else:
                        print(f"Gemini stream and fallback were empty for {self._model_name}")


            return _iter_stream()

        response = self._client.models.generate_content(
            model=self._model_name,
            contents=normalized_contents,
            config=config,
        )
        text = _extract_gemini_text(response, question)
        if not text:
            print(f"Gemini non-stream response was empty for {self._model_name}")
        return SimpleNamespace(
            text=text,
            content=[SimpleNamespace(text=text)] if text else []
        )



class _GeminiCompat:
    def __init__(self):
        self._client = None

    @property
    def models(self):
        if not self._client:
            raise RuntimeError('Gemini client is not configured')
        return self._client.models

    def configure(self, api_key=None):
        if not api_key:
            self._client = None
            return
        # to_gemini_timeout ensures 10s floor and millisecond conversion
        self._client = google_genai.Client(api_key=api_key, http_options={'timeout': to_gemini_timeout(300)})

    def GenerativeModel(self, model_name):
        if not self._client:
            raise RuntimeError('Gemini client is not configured')
        return _GeminiCompatModel(self._client, model_name)

    def embed_content(self, model, content, task_type=None):
        if not self._client:
            raise RuntimeError('Gemini client is not configured')

        config = {'task_type': task_type} if task_type else None
        response = self._client.models.embed_content(
            model=_normalize_gemini_model_name(model),
            contents=content,
            config=config,
        )

        # Keep old extraction logic compatible by returning dict with `embedding`.
        embedding_values = None
        if getattr(response, 'embeddings', None):
            first = response.embeddings[0]
            embedding_values = getattr(first, 'values', None) or getattr(first, 'embedding', None)
        if embedding_values is None:
            embedding_values = getattr(response, 'embedding', None)

        return {'embedding': embedding_values or []}


genai = _GeminiCompat()

# Cache for AI responses to reduce quota usage
health_narrative_cache = {}
CACHE_TTL = 3600  # 1 saat cache

# Project embedding cache (RAG-lite for project chat)
project_embedding_cache = {}
PROJECT_EMBED_CACHE_TTL = 1800  # 30 minutes

PLAN_LIMITS = {
    'free': {
        'daily_requests': 30,
        'monthly_tokens': 120000,
    },
    'premium': {
        'daily_requests': 250,
        'monthly_tokens': 1200000,
    }
}

# -------------------------------------------------------
# 💰 TOKEN EKONOMİSİ — Model başına token maliyeti
# Kullanıcı kararı: Gemini=10, GPT=2, Claude=12
# Model Blend ve Compare daha pahalı (çoklu çağrı içerir)
# -------------------------------------------------------
TOKEN_COSTS = {
    # Gemini ailesi
    'gemini': 5,
    'gemini-3-flash-preview': 5,
    'gemini-2.5-flash': 4,
    'gemini-3.1-flash-lite-preview': 2,
    'gemini-2.5-flash-lite': 1,
    'gemini-1.5-flash': 1,
    'gemma': 3,
    'gemma-4-26b-a4b-it': 3,
    'gemma-4-31b-it': 4,
    'gemma-2-27b-it': 3,
    'gemma-2-9b-it': 2,
    # OpenAI ailesi
    'gpt': 10,
    'gpt-4o': 10,
    'gpt-4o-mini': 2,
    'gpt-3.5-turbo': 2,
    # Anthropic ailesi
    'claude': 10,
    'claude-sonnet-4-5-20250929': 10,
    'claude-opus-4-5-20251101': 18,
    'claude-sonnet': 10,
    'claude-opus': 18,
    'claude-haiku': 5,
    # Özel modlar
    'model-blend': 20,
    'model-compare': 25,
    # Varsayılan (bilinmeyen model)
    'default': 5,
}

SIGNUP_GRANT_TOKENS = 100  # Yeni kullanıcıya verilen ücretsiz başlangıç token'ı
MONTHLY_GRANT_TOKENS = 100  # Her ay başında korunacak minimum token hakkı

DEFAULT_TOKEN_PACKAGES = [
    {
        'id': 'starter',
        'name': 'Starter',
        'description': 'Solo kullanım ve hafif deneme akışları için.',
        'tokens': 500,
        'price_usd': 5.0,
        'price_try': 175.0,
        'bonus_pct': 0,
    },
    {
        'id': 'pro-pack',
        'name': 'Pro Pack',
        'description': 'Sürekli üretim akışı olan bireyler ve küçük ekipler için.',
        'tokens': 2000,
        'price_usd': 15.0,
        'price_try': 495.0,
        'bonus_pct': 5,
    },
    {
        'id': 'heavy-user-bundle',
        'name': 'Heavy User Bundle',
        'description': 'Yoğun kullanım ve ekip içi denemeler için.',
        'tokens': 8000,
        'price_usd': 49.0,
        'price_try': 1590.0,
        'bonus_pct': 10,
    },
    {
        'id': 'studio-upgrade',
        'name': 'Studio Upgrade',
        'description': 'Kurumsal ekipler ve yüksek hacimli kullanım için.',
        'tokens': 20000,
        'price_usd': 99.0,
        'price_try': 2990.0,
        'bonus_pct': 15,
    },
]


def _cosine_similarity(v1, v2):
    """Compute cosine similarity without extra dependencies."""
    if not v1 or not v2 or len(v1) != len(v2):
        return -1.0
    dot = 0.0
    n1 = 0.0
    n2 = 0.0
    for a, b in zip(v1, v2):
        dot += a * b
        n1 += a * a
        n2 += b * b
    if n1 <= 0.0 or n2 <= 0.0:
        return -1.0
    return dot / (math.sqrt(n1) * math.sqrt(n2))


def _project_signature(project_files):
    """Return a stable signature for cache invalidation when files change."""
    sig_parts = []
    for pf in project_files:
        updated = pf.updated_at.isoformat() if pf.updated_at else ''
        sig_parts.append(f"{pf.id}:{pf.name}:{updated}:{len(pf.content or '')}")
    return '|'.join(sig_parts)


def _chunk_text(text, chunk_size=1200, overlap=200):
    """Split text into overlapping chunks."""
    text = (text or '').strip()
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0
    text_len = len(text)
    while start < text_len:
        end = min(text_len, start + chunk_size)
        chunks.append(text[start:end])
        if end >= text_len:
            break
        start = max(end - overlap, start + 1)
    return chunks


def _extract_embedding_values(resp):
    """Extract embedding vector from different SDK response shapes."""
    try:
        # dict style
        if isinstance(resp, dict):
            if isinstance(resp.get('embedding'), list):
                return resp.get('embedding')
            emb_list = resp.get('embeddings')
            if isinstance(emb_list, list) and emb_list:
                first = emb_list[0]
                if isinstance(first, dict):
                    if isinstance(first.get('values'), list):
                        return first.get('values')
                    if isinstance(first.get('embedding'), list):
                        return first.get('embedding')

        # object style
        if hasattr(resp, 'embedding') and isinstance(resp.embedding, list):
            return resp.embedding

        if hasattr(resp, 'embeddings') and resp.embeddings:
            first = resp.embeddings[0]
            if hasattr(first, 'values') and isinstance(first.values, list):
                return first.values
            if hasattr(first, 'embedding') and isinstance(first.embedding, list):
                return first.embedding
    except Exception:
        return None

    return None


def _embed_text_with_fallback(text, task_type='RETRIEVAL_DOCUMENT'):
    """Embed text via Gemini with model fallback for compatibility."""
    if not text or not GEMINI_API_KEY:
        return None, None

    candidates = []
    env_model = os.getenv('EMBEDDING_MODEL_NAME', 'models/gemini-embedding-2-preview')
    candidates.append(env_model)
    candidates.extend([
        'models/gemini-embedding-2-preview',
        'gemini-embedding-2-preview',
        'models/gemini-embedding-001',
        'gemini-embedding-001',
    ])

    # Dedupe while preserving order
    seen = set()
    unique_candidates = []
    for c in candidates:
        if c and c not in seen:
            seen.add(c)
            unique_candidates.append(c)

    for model_name in unique_candidates:
        try:
            resp = genai.embed_content(
                model=model_name,
                content=text,
                task_type=task_type,
            )
            values = _extract_embedding_values(resp)
            if values:
                return values, model_name
        except Exception as e:
            print(f"Embedding call failed for {model_name}: {e}")
            continue

    return None, None


def _build_project_embedding_index(project_or_id):
    """Build or reuse cached embedding index for project files."""
    with app.app_context():
        # Force re-fetch to ensure we are in a session
        pid = project_or_id if isinstance(project_or_id, (int, str)) else getattr(project_or_id, 'id', None)
        if not pid:
            return None
            
        project = db.session.get(Project, int(pid))
        if not project:
            return None
            
        files = project.files.order_by(ProjectFile.name).all()
        if not files:
            return None

        signature = _project_signature(files)
        now = time.time()

        cached = project_embedding_cache.get(pid)
        if cached and cached.get('signature') == signature and (now - cached.get('timestamp', 0) < PROJECT_EMBED_CACHE_TTL):
            return cached

        raw_items = []
        for pf in files:
            content = (pf.content or '')[:12000]  # hard limit per file for cost control
            for idx, chunk in enumerate(_chunk_text(content, chunk_size=1200, overlap=200)):
                raw_items.append({
                    'pf_name': pf.name,
                    'pf_lang': pf.language or 'plaintext',
                    'text': chunk,
                    'chunk_index': idx
                })

        if not raw_items:
            return None

        def _embed_task(item):
            emb, model_used = _embed_text_with_fallback(item['text'], task_type='RETRIEVAL_DOCUMENT')
            if not emb:
                return None
            return {
                'file': item['pf_name'],
                'language': item['pf_lang'],
                'text': item['text'],
                'embedding': emb,
                'chunk_index': item['chunk_index'],
                'model_used': model_used,
            }

        chunks = []
        # Use max_workers=10 for fast parallel embedding. Most API quotas allow this.
        with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
            results = list(executor.map(_embed_task, raw_items))
        
        chunks = [r for r in results if r]

        if not chunks:
            return None

        index_data = {
            'signature': signature,
            'timestamp': now,
            'chunks': chunks,
        }
        project_embedding_cache[pid] = index_data
        return index_data


def build_project_context_for_question(project_or_id, question, top_k=6):
    """Return concise, relevance-ranked project context for a user question."""
    with app.app_context():
        if isinstance(project_or_id, (int, str)):
            try:
                pid = int(project_or_id)
                project = db.session.get(Project, pid)
            except Exception:
                return ""
        else:
            project = project_or_id

        if not project:
            return ""

        semantic_result = get_project_semantic_hits(project, question, top_k=top_k)
        if not semantic_result:
            return f"[System: Bu sohbet '{project.name}' projesine aittir.]"

        query_model = semantic_result.get('query_model', {})
        selected = semantic_result.get('hits', [])

        ctx = [f"[System: Bu sohbet '{project.name}' projesine aittir."]
        if project.description:
            ctx.append(f"Proje açıklaması: {project.description}")
        ctx.append(f"Soruya göre embedding tabanlı en alakalı dosya parçaları (query_model={query_model}):")

        total_chars = 0
        for item in selected:
            score = item.get('score', 0)
            snippet = item.get('snippet', '')
            total_chars += len(snippet)
            if total_chars > 7000:
                break
            ctx.append(
                f"\n## Dosya: {item['file']} ({item['language']}) | Benzerlik: {score:.3f}\n"
                f"```{item['language']}\n{snippet}\n```"
            )

        ctx.append(']')
        return '\n'.join(ctx)


def get_project_semantic_hits(project_or_id, question, top_k=6):
    """Return top semantic hits for a question within project files."""
    index_data = _build_project_embedding_index(project_or_id)
    if not index_data:
        return None

    query_emb, query_model = _embed_text_with_fallback(question or '', task_type='RETRIEVAL_QUERY')
    if not query_emb:
        return None

    scored = []
    for item in index_data['chunks']:
        score = _cosine_similarity(query_emb, item['embedding'])
        if score > 0:
            scored.append((score, item))

    if not scored:
        return None

    scored.sort(key=lambda x: x[0], reverse=True)
    selected = scored[:top_k]

    hits = []
    for score, item in selected:
        hits.append({
            'score': round(float(score), 4),
            'file': item['file'],
            'language': item['language'],
            'chunk_index': item['chunk_index'],
            'snippet': item['text'][:1200],
            'chunk_model': item.get('model_used'),
        })

    return {
        'query_model': query_model,
        'hits': hits,
        'total_chunks': len(index_data['chunks']),
    }


def invalidate_project_embedding_cache(project_id):
    try:
        pid = int(project_id)
    except Exception:
        return
    project_embedding_cache.pop(pid, None)


def _agent_project_search(project_id, query, top_k=6):
    with app.app_context():
        try:
            pid = int(project_id)
        except Exception:
            return None
        project = db.session.get(Project, pid)
        if not project:
            return None
        result = get_project_semantic_hits(project, query, top_k=top_k)
        if not result:
            return None
        return {
            'hits': [
                {
                    'file': hit.get('file'),
                    'score': hit.get('score'),
                    'text': hit.get('snippet'),
                    'chunk_index': hit.get('chunk_index'),
                }
                for hit in (result.get('hits') or [])
            ]
        }


def _agent_db_read(query, limit=100):
    """Run read-only SQL for agent db_read tool with a server-side row cap."""
    with app.app_context():
        sql = str(query or '').strip()
        if not sql:
            return {'ok': False, 'error': 'query is required'}

        lowered = sql.lower().strip()
        if not (lowered.startswith('select') or lowered.startswith('with')):
            return {'ok': False, 'error': 'Only SELECT/WITH queries are allowed.'}

        safe_limit = max(1, min(500, int(limit or 100)))
        wrapped = f"SELECT * FROM ({sql}) AS _agent_read LIMIT :agent_limit"
        rows = db.session.execute(sql_text(wrapped), {'agent_limit': safe_limit}).mappings().all()
        return {
            'ok': True,
            'limit': safe_limit,
            'row_count': len(rows),
            'rows': [dict(r) for r in rows],
        }


def _parse_workspace_files_payload(raw_value, max_files=80, max_chars_per_file=18000):
    parsed = raw_value
    if isinstance(raw_value, str):
        raw_text = raw_value.strip()
        if not raw_text:
            return []
        try:
            parsed = json.loads(raw_text)
        except Exception:
            return []

    if not isinstance(parsed, list):
        return []

    files = []
    for item in parsed:
        if not isinstance(item, dict):
            continue
        path = str(item.get('path') or item.get('name') or '').strip()
        if not path:
            continue
        content = str(item.get('content') or item.get('text') or '')
        language = str(item.get('language') or item.get('lang') or 'plaintext')
        files.append({
            'path': path,
            'content': content[:max_chars_per_file],
            'language': language,
        })
        if len(files) >= max_files:
            break
    return files


def _resolve_agent_provider_model(selected_model):
    raw_model = str(selected_model or '').strip()
    normalized = raw_model.replace('models/', '', 1) if raw_model.startswith('models/') else raw_model
    model_lc = normalized.lower()

    if 'claude' in model_lc:
        return 'anthropic', normalized
    if 'gpt' in model_lc or model_lc.startswith('o1') or model_lc.startswith('o3'):
        return 'openai', normalized
    if 'gemini' in model_lc or 'gemma' in model_lc:
        return 'gemini', normalized
    return None, normalized


SUPPORTED_GEMINI_AGENT_MODELS = (
    'gemini-3-flash-preview',
    'gemini-3.1-flash-lite-preview',
    'gemini-2.5-flash',
    'gemini-2.5-flash-lite',
    'gemini-2.0-flash',
    'gemini-1.5-flash',
    'gemini-1.5-pro',
)


def _is_agent_model_supported(provider: str, model: str) -> bool:
    provider_lc = str(provider or '').lower().strip()
    model_lc = str(model or '').lower().strip()

    if provider_lc in {'openai', 'anthropic'}:
        return True

    if provider_lc == 'gemini':
        if model_lc.startswith('gemma-'):
            return True
        # Agent Mode allows Gemini text/tool-calling families.
        # Also allow any gemini model that contains 'flash' (more permissive)
        if any(model_lc.startswith(prefix) for prefix in SUPPORTED_GEMINI_AGENT_MODELS):
            return True
        # Fallback: if it's a gemini model with 'flash' in the name, allow it
        if 'gemini' in model_lc and 'flash' in model_lc:
            return True

    return False


def _resolve_agent_project(user, conversation, payload_project_id):
    if conversation and conversation.project_id:
        project = db.session.get(Project, conversation.project_id)
        if project and user and project.user_id == user.id:
            return project, 'conversation'
        if project and user is None:
            return project, 'conversation'

    if payload_project_id and user:
        try:
            project_id_int = int(payload_project_id)
        except Exception:
            project_id_int = None
        if project_id_int:
            project = Project.query.filter_by(id=project_id_int, user_id=user.id).first()
            if project:
                return project, 'payload'

    # --- FALLBACK: Auto-Discovery Removed ---
    # We no longer automatically link conversations to the last project. 
    # Users must explicitly select a project or be in a project workspace.
    
    return None, None



def _clip_agent_metadata(trace, changed_files, max_trace=10, max_changed=20):
    safe_trace = trace if isinstance(trace, list) else []
    safe_changed = changed_files if isinstance(changed_files, list) else []
    return {
        'trace': safe_trace[:max_trace],
        'changed_files': safe_changed[:max_changed],
        'trace_total': len(safe_trace),
        'changed_total': len(safe_changed),
        'trace_truncated': len(safe_trace) > max_trace,
        'changed_truncated': len(safe_changed) > max_changed,
    }


# Lightweight in-memory state for external tool conversations (VS Code extension, API clients).
_external_conversation_state = {}
# VS Code login state is now database-backed via VSCodeLoginState model
VSCODE_LOGIN_STATE_TTL_SECONDS = 600


def _cleanup_vscode_login_state():
    now = time.time()
    try:
        expired = VSCodeLoginState.query.filter(VSCodeLoginState.expires_at < now).all()
        for s in expired:
            db.session.delete(s)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        print(f"[VSCODE-AUTH] Cleanup error: {e}")


def _is_valid_vscode_state(value: str) -> bool:
    return bool(re.fullmatch(r'[A-Za-z0-9_-]{16,128}', str(value or '').strip()))


def _normalise_path_for_match(value):
    return str(value or '').replace('\\', '/').strip().lower()


def _conversation_state_key(user_id, conversation_id):
    conv_part = str(conversation_id).strip() if conversation_id is not None else 'default'
    return f"{int(user_id)}:{conv_part}"


def _is_short_confirmation_intent(text):
    q = str(text or '').strip().lower()
    if not q:
        return False

    compact = re.sub(r'\s+', ' ', q)
    compact = re.sub(r'[!?.,;:]+$', '', compact).strip()

    explicit = {
        'değiştir', 'degistir', 'o zaman değiştir', 'o zaman degistir',
        'uygula', 'apply', 'apply changes', 'ok', 'ok yap', 'tamam',
        'evet', 'do it', 'change it', 'yap', 'yap gitsin', 'onayla',
    }
    if compact in explicit:
        return True

    # Keep this narrow to avoid false positives on regular questions.
    words = compact.split()
    if len(words) <= 4 and any(token in compact for token in ('değiştir', 'degistir', 'uygula', 'apply')):
        return True

    return False


def _resolve_workspace_file_content(workspace_files, active_file):
    target_norm = _normalise_path_for_match(active_file)
    if not target_norm:
        return '', ''

    for item in (workspace_files or []):
        path = str(item.get('path') or '').strip()
        content = str(item.get('content') or '')
        path_norm = _normalise_path_for_match(path)
        if not path_norm:
            continue
        if path_norm == target_norm or path_norm.endswith('/' + target_norm) or target_norm.endswith('/' + path_norm):
            return path, content

    return '', ''

# Load environment variables
basedir = os.path.abspath(os.path.dirname(__file__))
env_path = os.path.join(basedir, '.env')

# Load .env file
load_dotenv(env_path, override=True, encoding='utf-8')

# Fallback: ensure critical keys are available even if dotenv parser misses entries.
required_env_keys = (
    'GEMINI_API_KEY',
    'OPENAI_API_KEY',
    'ANTHROPIC_API_KEY',
    'STRIPE_SECRET_KEY',
    'STRIPE_PUBLISHABLE_KEY',
    'STRIPE_WEBHOOK_SECRET',
)
if any(not os.getenv(key) for key in required_env_keys):
    try:
        parsed = dotenv_values(env_path, encoding='utf-8')
        for key, value in parsed.items():
            if key and value is not None:
                os.environ[key] = str(value).strip()
    except Exception:
        pass  # Silent fail - individual feature checks handle missing keys.


# static_folder has to be absolute or relative to this file. 
# In render-build.sh, we copy client/dist/* to server/static.
static_folder_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'static')
app = Flask(__name__, static_folder=static_folder_path, static_url_path='')

# SaaS-Grade Lifecycle Orchestration (Start Async Worker)
start_worker(app)
basedir = os.path.abspath(os.path.dirname(__file__))
# Veritabanı dosyasını instance klasöründe tutuyoruz (Flask standardı)
instance_path = os.path.join(basedir, 'instance')
if not os.path.exists(instance_path):
    os.makedirs(instance_path, exist_ok=True)

db_path = os.path.join(instance_path, 'codebrain.db')
# SQLite default fallback
default_db = 'sqlite:///' + db_path
database_url = os.getenv('DATABASE_URL', default_db)

# Heroku compatibility: postgres:// to postgresql://
if database_url.startswith("postgres://"):
    database_url = database_url.replace("postgres://", "postgresql://", 1)

app.config['SQLALCHEMY_DATABASE_URI'] = database_url
# Pool Optimization for Concurrency
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    'pool_size': 50,
    'max_overflow': 150,
    'pool_recycle': 1800,
    'pool_pre_ping': True
}
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
# JWT Security: Require strong secret key in production
_jwt_secret = os.getenv('JWT_SECRET_KEY')
if not _jwt_secret or _jwt_secret == 'dev-secret-key':
    import secrets
    _jwt_secret = secrets.token_hex(32)
    print("WARNING: JWT_SECRET_KEY not set or using default. Generated temporary key. Set a strong JWT_SECRET_KEY in .env for production!")
app.config['JWT_SECRET_KEY'] = _jwt_secret
app.config['JWT_ACCESS_TOKEN_EXPIRES'] = timedelta(hours=6)
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB limit

if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

from werkzeug.utils import secure_filename
from flask import send_from_directory
import base64
import mimetypes

# Register MIME types for static assets
mimetypes.add_type('application/javascript', '.js')
mimetypes.add_type('text/css', '.css')

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({'status': 'ok'}), 200

CORS(app, resources={r"/*": {"origins": "*"}})
db.init_app(app)
jwt = JWTManager(app)


@app.before_request
def sanitize_auth_headers():
    """
    Pre-sanitize Authorization headers before JWT library processing.
    
    Purpose:
    - Catch malformed/binary Authorization headers early
    - Replace bad headers with empty string so JWT sees "missing token" instead of "malformed"
    - Prevents UTF-8 codec errors from propagating through JWT library
    
    Approach:
    - Try to get and decode Authorization header safely
    - If it fails or contains non-printable chars, remove it
    - API keys (ca-...) and Bearer tokens will still work as they're valid ASCII
    """
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header:
            return  # No header, nothing to do
        
        # Try to convert to string safely
        auth_str = str(auth_header)
        
        # Check if it's a valid JWT or API key format
        # Valid formats: "Bearer <token>" or "ca-<key>"
        if auth_str.lower().startswith('bearer '):
            # Try to extract bearer token
            parts = auth_str.split(None, 1)
            if len(parts) != 2:
                # Malformed bearer header, remove it
                print(f"DEBUG: Removing malformed Bearer header")
                request.environ['HTTP_AUTHORIZATION'] = ''
                return
            token = parts[1].strip()
        elif auth_str.startswith('ca-'):
            # API key format, keep it as-is
            token = auth_str
        else:
            # Unknown format, could be garbage - remove it
            print(f"DEBUG: Removing unknown auth format: {repr(auth_str[:60])}")
            request.environ['HTTP_AUTHORIZATION'] = ''
            return
        
        # Validate JWT structure if it's a Bearer token
        if auth_str.lower().startswith('bearer '):
            # Bearer tokens must be 3 dot-separated segments (header.payload.signature)
            if token.count('.') != 2:
                # Malformed JWT, remove it
                print(f"DEBUG: Removing JWT with invalid segment count: {token.count('.') + 1} segments instead of 3")
                request.environ['HTTP_AUTHORIZATION'] = ''
                return
        
        # Try to validate token is safe ASCII/printable
        try:
            # If token decodes to valid UTF-8 and contains only printable chars, it's safe
            token_bytes = token.encode('utf-8', errors='strict')
            token.encode('ascii', errors='strict')  # JWT tokens should be ASCII
            # If we got here, token looks valid
        except (UnicodeEncodeError, UnicodeDecodeError) as e:
            # Token contains non-ASCII or invalid UTF-8, remove header
            print(f"DEBUG: Removing non-ASCII auth token: {e}")
            request.environ['HTTP_AUTHORIZATION'] = ''
            return
    except Exception as e:
        # Safety: if anything goes wrong in header sanitization, remove the header
        print(f"DEBUG: Exception in sanitize_auth_headers: {e}, removing Authorization header")
        try:
            request.environ['HTTP_AUTHORIZATION'] = ''
        except:
            pass


# SocketIO async mode is pinned to threading for cross-platform stability.
socketio = SocketIO(
    app,
    cors_allowed_origins="*",
    async_mode=ASYNC_MODE,
    logger=False,
    engineio_logger=False,
    ping_timeout=60,
    ping_interval=25,
)

# --- SOCKETIO EVENT HANDLERS ---
@socketio.on('join_room')
def handle_join_room(data):
    token = data.get('token')
    user_name = data.get('user_name', 'Guest')
    if token:
        join_room(token)
        print(f"DEBUG: User {user_name} joined room {token}")
        # Odaya katılan yeni kullanıcıyı diğerlerine duyur
        socketio.emit('user_joined', {'user_name': user_name, 'token': token}, room=token)

@socketio.on('leave_room')
def handle_leave_room(data):
    token = data.get('token')
    user_name = data.get('user_name', 'Guest')
    if token:
        leave_room(token)
        print(f"DEBUG: User {user_name} left room {token}")
        socketio.emit('user_left', {'user_name': user_name, 'token': token}, room=token)

@socketio.on('connect')
def handle_connect():
    print(f"DEBUG: Socket connected: {request.sid}")

@socketio.on('disconnect')
def handle_disconnect():
    print(f"DEBUG: Socket disconnected: {request.sid}")

# Initialize database tables (Critical for Gunicorn/Render)
with app.app_context():
    try:
        db.create_all()
        print("Database tables initialized successfully.")

        # --- Startup Seeding: Ensure token packages match defaults ---
        try:
            current_count = TokenPackage.query.count()
            if current_count == 0:
                for pkg_data in DEFAULT_TOKEN_PACKAGES:
                    pkg = TokenPackage(
                        name=pkg_data['name'],
                        description=pkg_data.get('description', ''),
                        tokens=pkg_data['tokens'],
                        price_usd=pkg_data['price_usd'],
                        bonus_pct=pkg_data.get('bonus_pct', 0),
                        is_active=True,
                    )
                    db.session.add(pkg)
                db.session.commit()
                print(f"Seeded {len(DEFAULT_TOKEN_PACKAGES)} default token packages.")
            else:
                # Update existing if they match tokens (alignment strategy)
                for pkg_data in DEFAULT_TOKEN_PACKAGES:
                    existing = TokenPackage.query.filter_by(tokens=pkg_data['tokens']).first()
                    if existing:
                        existing.name = pkg_data['name']
                        existing.description = pkg_data['description']
                        existing.price_usd = pkg_data['price_usd']
                        existing.bonus_pct = pkg_data['bonus_pct']
                db.session.commit()
                print("Token packages aligned with defaults.")
        except Exception as seed_err:
            print(f"Warning: Could not seed/align token packages: {seed_err}")
            db.session.rollback()

    except Exception as e:
        print(f"Error initializing database: {e}")

# Database initialization handled above.

@jwt.invalid_token_loader
def invalid_token_callback(error):
    error_str = str(error)
    print(f"DEBUG: JWT invalid_token_loader called: {error_str}")
    
    # Catch UTF-8 encoding errors from malformed headers
    if 'utf-8' in error_str.lower() or 'codec' in error_str.lower() or 'decode' in error_str.lower():
        print(f"DEBUG: ENCODING ERROR in JWT decode: {error_str}")
        try:
            auth = request.headers.get('Authorization', '')
            if auth:
                print(f"DEBUG: Problem auth header (first 80 chars): {repr(str(auth)[:80])}")
        except:
            pass
        return jsonify({
            'error': 'Authorization header encoding error',
            'code': 'AUTH_ENCODING_ERROR'
        }), 401
    
    return jsonify({'error': f'Invalid token: {error}'}), 422

@jwt.unauthorized_loader
def missing_token_callback(error):
    print(f"DEBUG: JWT missing_token_loader called: {error}")
    return jsonify({'error': f'Missing token: {error}'}), 401

@jwt.expired_token_loader
def expired_token_callback(jwt_header, jwt_payload):
    print(f"DEBUG: JWT expired_token_loader called: {jwt_payload}")
    return jsonify({'error': 'Token expired', 'code': 'TOKEN_EXPIRED'}), 401

# --- MIDDLEWARE & ERROR HANDLERS ---

@app.errorhandler(Exception)
def handle_exception(e):
    error_str = str(e)
    
    # Catch encoding errors early and return 401 instead of 500
    if 'utf-8' in error_str.lower() or 'codec' in error_str.lower() or 'decode' in error_str.lower():
        print(f"DEBUG: GLOBAL UTF-8 ENCODING ERROR: {error_str}")
        try:
            auth = request.headers.get('Authorization', '')
            if auth:
                print(f"DEBUG: Problem auth header: {repr(str(auth)[:80])}")
        except:
            pass
        return jsonify({
            'error': 'Request header encoding error',
            'code': 'HEADER_ENCODING_ERROR'
        }), 400
    
    # Preserve HTTP errors like 404/405 instead of converting all of them to 500.
    if isinstance(e, HTTPException):
        return jsonify({
            'error': e.name,
            'details': e.description
        }), e.code

    # Log unexpected errors with traceback.
    print(f"🔥 GLOBAL ERROR: {str(e)}")
    import traceback
    traceback.print_exc()
    return jsonify({
        'error': 'An internal server error occurred.',
        'details': str(e)
    }), 500

def call_gemini_with_retry(prompt, model_name='gemini-2.5-flash', max_retries=2):
    """Calls Gemini API with exponential backoff for 429 errors."""
    for i in range(max_retries):
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(prompt)
            if response and response.text:
                return response.text
            return None
        except Exception as e:
            err_msg = str(e)
            if "429" in err_msg or "quota" in err_msg.lower():
                if i < max_retries - 1:
                    wait_time = (i + 1) * 3
                    print(f"Gemini 429 received. Waiting {wait_time}s before retry {i+1}...")
                    time.sleep(wait_time)
                    continue
                else:
                    print("Gemini quota exceeded after retries. Triggering fallback.")
                    return None # Explicitly return None to trigger fallback in route
            raise e
    return None

# --- GITHUB FEATURES ROTALARI ---

# Legacy endpoint kept for reference. Disabled to avoid route collision.
# @app.route('/api/github/blueprint', methods=['GET'])
def get_repo_blueprint():
    print("DEBUG: Blueprint route hit!")
    """Generates a project blueprint markdown with a Mermaid diagram."""
    repo_url = request.args.get('repo')
    if not repo_url:
        return jsonify({'error': 'Repo URL is required'}), 400

    parser = GitHubParser()
    tree = parser.get_repo_tree(repo_url)
    
    if not tree:
        return jsonify({'error': 'Failed to fetch repository tree'}), 404

    tree_str = parser.format_tree_for_prompt(tree)
    
    prompt = f"""
    Analyze the following repository structure and generate a comprehensive Project Blueprint in Markdown.
    
    Repository: {repo_url}
    
    Structure:
    {tree_str}
    
    The blueprint should include:
    1. **Overview**: High-level purpose of the project.
    2. **Architecture**: A Mermaid.js classDiagram or graphTD representing the project structure.
    3. **Key Components**: Description of main directories and files.
    4. **Tech Stack**: Inferred technologies based on file extensions.
    
    Respond ONLY with the Markdown content.
    """
    
    try:
        content = call_gemini_with_retry(prompt)
        if content is None:
            raise Exception("Gemini quota exceeded or empty response")
        return jsonify({'markdown': content})
    except Exception as e:
        # Dynamic Heuristic Fallback for Blueprint
        repo_name = repo_url.split('/')[-1]
        paths = [item['path'] for item in tree]
        top_dirs = sorted(list(set([p.split('/')[0] for p in paths if '/' in p])))[:7]
        dir_list = "\n".join([f"- **{d}/**: Identified project component." for d in top_dirs])
        
        fallback_markdown = f"""# Project Blueprint: {repo_name}

## Overview
Automated structure analysis for {repo_url}.

## Architecture
```mermaid
graph TD
    Root[Project Root]
"""
        # Add some dynamic nodes to mermaid
        for i, d in enumerate(top_dirs[:4]):
            fallback_markdown += f"    Root --> Dir{i}[{d}]\n"
        
        fallback_markdown += f"""```

## Key Components
{dir_list if dir_list else "- *No major directories identified.*"}

## Tech Stack
- Multi-language support (Standard heuristics applied).

> [!NOTE]
> This is a dynamically generated heuristic blueprint as the AI service is currently at capacity.
"""
        return jsonify({'markdown': fallback_markdown})

# DUPLICATE ENDPOINT - DEVRE DIŞI (Gerçek endpoint satır ~3550'de)
# @app.route('/api/github/health', methods=['GET'])
# def get_repo_health():
#     """Analyzes repository health and returns metrics + narrative."""
#     ...eski kod kaldırıldı...

# --- 1. GEMINI KONFIGURASYONU ---
GEMINI_MODEL = os.getenv('GEMINI_MODEL_NAME', 'models/gemini-2.5-flash')
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')

if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
else:
    print("Warning: GEMINI_API_KEY not defined. Gemini calls disabled.")

# --- 2. CLAUDE KONFIGURASYONU ---
# Varsayılan model olarak hızlı ve zeki olan Sonnet 4.5'i seçtik
ANTHROPIC_MODEL = os.getenv('CLAUDE_MODEL_NAME', 'claude-sonnet-4-5-20250929')
ANTHROPIC_API_KEY = os.getenv('ANTHROPIC_API_KEY')

claude_client = None
if ANTHROPIC_API_KEY:
    try:
        claude_client = Anthropic(api_key=ANTHROPIC_API_KEY)
    except Exception as e:
        print(f"Warning: Failed to initialize Anthropic client: {e}")
else:
    print("Warning: ANTHROPIC_API_KEY not defined. Claude calls disabled.")

# --- 3. OPENAI (GPT) KONFIGURASYONU ---
OPENAI_MODEL = os.getenv('OPENAI_MODEL_NAME', 'gpt-4o')
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')


openai_client = None
openai_init_error = None

if OPENAI_API_KEY:
    try:
        openai_client = OpenAI(api_key=OPENAI_API_KEY)
        print("OpenAI client initialized successfully.")
    except Exception as e:
        openai_init_error = str(e)
        print(f"Warning: Failed to initialize OpenAI client: {e}")
else:
    print("Warning: OPENAI_API_KEY not defined. GPT calls disabled.")


# --- MODEL FONKSİYONLARI ---

# Initialize Utils
language_detector = LanguageDetector(os.getenv('GEMINI_API_KEY'))
model_router = ModelRouter()


# --- YARDIMCI FONKSİYONLAR ---
def transcribe_audio_with_gemini(audio_path):
    """Gemini kullanarak ses dosyasını metne çevirir."""
    try:
        if not GEMINI_API_KEY:
            return None

        import mimetypes
        
        mime_type, _ = mimetypes.guess_type(audio_path)
        # Fallback mime types
        file_ext = os.path.splitext(audio_path)[1].lower()
        if not mime_type:
            if file_ext == '.mp3': mime_type = 'audio/mpeg'
            elif file_ext == '.wav': mime_type = 'audio/wav'
            elif file_ext == '.m4a': mime_type = 'audio/mp4'
            elif file_ext == '.webm': mime_type = 'audio/webm'
            elif file_ext == '.ogg': mime_type = 'audio/ogg'
            else: mime_type = 'audio/mp3'

        # Denenecek model listesi (En yüksek ücretsiz kotadan en düşüğe)
        model_candidates = [
            GEMINI_MODEL.replace('models/', ''), # .env'deki model
            "gemini-3.1-flash-lite-preview",
            "gemini-2.5-flash-lite",
            "gemini-3-flash-preview",
            "gemini-2.5-flash"
        ]
        
        last_error = None
        for m_name in model_candidates:
            try:
                model = genai.GenerativeModel(m_name)
                
                with open(audio_path, 'rb') as audio_file:
                    audio_bytes = audio_file.read()
                    
                audio_part = {
                    "mime_type": mime_type,
                    "data": audio_bytes
                }
                
                response = model.generate_content([
                    audio_part, 
                    "Please transcribe this audio exactly as it is spoken in Turkish. Just return the text."
                ])
                
                if response and response.text:
                    return response.text.strip()
            except Exception as e:
                print(f"Transcription trial with {m_name} failed: {e}")
                last_error = e
                continue
                
        return None
    except Exception as e:
        print(f"Transcription error: {e}")
        return None

def generate_image_with_dalle(prompt):
    """OpenAI DALL-E 3 kullanarak resim oluşturur."""
    if not openai_client:
        return "Error: OpenAI API key not found."
    
    try:
        print(f"Generating image for prompt: {prompt}")
        response = openai_client.images.generate(
            model="dall-e-3",
            prompt=prompt,
            size="1024x1024",
            quality="standard",
            n=1,
        )
        
        image_url = response.data[0].url
        
        # Resmi indirip yerel olarak kaydet (URL'lerin süresi doluyor)
        import requests
        from datetime import datetime
        
        img_data = requests.get(image_url).content
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"generated_{timestamp}.png"
        
        # Klasör yoksa oluştur
        save_dir = os.path.join(app.root_path, 'static', 'generated')
        os.makedirs(save_dir, exist_ok=True)
        
        save_path = os.path.join(save_dir, filename)
        
        with open(save_path, 'wb') as handler:
            handler.write(img_data)
            
        # Frontend için erişilebilir URL döndür (Relative path is safer for proxy/cors)
        local_url = f"/generated/{filename}"
        return f"![Generated Image]({local_url})\n\n**Generated for:** *{prompt}*"
        
    except Exception as e:
        print(f"DALL-E Error: {e}")
        return f"Sorry, I couldn't generate the image. Error: {str(e)}"

def generate_gemini_answer(question: str, code: str, history_context: list = None, requested_model: str = None, image_path: str = None, prefs: dict = None, github_context: str = None, depth: int = 0):
    """Gemini API çağrısı yapar. Sadece seçilen modeli kullanır."""
    if depth > 2:
        yield "[System Message]: Error: Maximum fallback depth reached. AI services are currently unavailable."
        return
    if not GEMINI_API_KEY:
        yield "Error: GEMINI_API_KEY missing."
        return

    # User Preferences: Style Prompt
    style_prompt = ""
    if prefs:
        if prefs.get('response_style') == 'concise':
            style_prompt = "Keep your answers very concise and short. "
        elif prefs.get('response_style') == 'detailed':
            style_prompt = "Provide detailed and comprehensive explanations. "

    persona_info = ""
    if prefs:
        persona = prefs.get('persona', 'General User')
        expertise = prefs.get('expertise', 'Intermediate')
        persona_info = f"I am a {persona} with {expertise} expertise level. "

    # Model Seçimi
    if requested_model and ('gemini' in requested_model or 'gemma' in requested_model):
        model_mapping = {
            # Gemini 3
            'gemini-3-flash-preview': 'gemini-3-flash-preview',
            # Gemini 3.1
            'gemini-3.1-flash-lite': 'gemini-3.1-flash-lite-preview',
            'gemini-3.1-flash-lite-preview': 'gemini-3.1-flash-lite-preview',
            # Gemini 2.5
            'gemini-2.5-flash-lite': 'gemini-2.5-flash-lite',
            'gemini-2.5-flash': 'gemini-2.5-flash',
            # Gemma 4
            'gemma-4-26b': 'gemma-4-26b-a4b-it',
            'gemma-4-26b-it': 'gemma-4-26b-a4b-it',
            'gemma-4-26b-a4b-it': 'gemma-4-26b-a4b-it',
            'gemma-4-31b': 'gemma-4-31b-it',
            'gemma-4-31b-it': 'gemma-4-31b-it',
            # Gemma 2
            'gemma-2-27b': 'gemma-2-27b-it',
            'gemma-2-27b-it': 'gemma-2-27b-it',
            'gemma-2-9b': 'gemma-2-9b-it',
            'gemma-2-9b-it': 'gemma-2-9b-it',
        }
        
        target_id = model_mapping.get(requested_model, requested_model)
        
        # Fallback Zinciri Hazırla
        fallback_chain = []
        
        # 1. Hedef modeli ekle
        fallback_chain.append(target_id)
        
        # 2. Kota sırasına göre yedek modeller
        for alt in ['gemini-3.1-flash-lite-preview', 'gemini-2.5-flash-lite', 'gemini-2.5-flash']:
            if alt not in fallback_chain:
                fallback_chain.append(alt)

        print(f"--- Model Zinciri: {fallback_chain} ---")
        
        system_instruction = (
            "You are a helpful AI assistant. Communicate with the user in a natural conversation style. "
            "If the user asks a question about code, software, or a technical topic, "
            "provide detailed technical assistance and give code examples if necessary (in Markdown code block). "
            "IMPORTANT: Always respond in the same language as the user's question (e.g., if the question is in Turkish, respond in Turkish). "
            "Never output internal analysis, background thinking, or instruction labels like <thought>, Thinking:, Input, Role, etc. "
            "CRITICAL: Provide ONLY the final response without any reasoning steps or internal dialogue. "
            "If context, memory, or repository information is provided below, you MUST strictly adhere to it and incorporate it into your response to ensure perfect continuity."
        )
        if github_context:
            system_instruction += f"\n\nCONTEXT FROM SYSTEM MEMORY OR REPOSITORY:\n{github_context}"
    else:
        # Varsayılan (Fallback zinciri ile)
        fallback_chain = [GEMINI_MODEL, 'gemini-3.1-flash-lite-preview', 'gemini-2.5-flash', 'gemini-2.5-flash-lite']
        system_instruction = (
            "You are a helpful AI assistant. Communicate with the user in a natural conversation style. "
            "If the user asks a question about code, software, or a technical topic, "
            "provide detailed technical assistance and give code examples if necessary (in Markdown code block). "
            "IMPORTANT: Always respond in the same language as the user's question (e.g., if the question is in Turkish, respond in Turkish). "
            "Never output internal analysis, background thinking, or instruction labels like <thought>, Thinking:, Input, Role, etc. "
            "CRITICAL: Provide ONLY the final response without any reasoning steps or internal dialogue. "
            "If context, memory, or repository information is provided below, you MUST strictly adhere to it and incorporate it into your response to ensure perfect continuity."
        )
        if github_context:
            system_instruction += f"\n\nCONTEXT FROM SYSTEM MEMORY OR REPOSITORY:\n{github_context}"

    # Restore the system prompt into prompt_parts to ensure context, memory, and persona are received
    system_prompt = system_instruction
    if persona_info or style_prompt:
        system_prompt = f"{persona_info}{style_prompt}\n\n{system_prompt}"
    
    prompt_parts = [system_prompt]
    

    # Eğer history_context yoksa, prompta asla örnek diyalog veya geçmiş başlığı eklenmesin
    model_name_lower = (requested_model or "").lower()
    if history_context:
        # Sadece geçmiş varsa few-shot ve başlık ekle
        if 'gemma' in model_name_lower:
            few_shot = [
                "User: Selam",
                "Assistant: Selam! İyiyim, teşekkür ederim. Siz nasılsınız? Size nasıl yardımcı olabilirim?",
                "User: Python'da liste nasıl sıralanır?",
                "Assistant: Python'da bir listeyi `sort()` metodu veya `sorted()` fonksiyonu ile sıralayabilirsiniz. Örnek: `liste.sort()`.",
            ]
            prompt_parts.extend(few_shot)

        filtered_history = []
        for turn in history_context:
            u_text = turn.get('user', '').strip()
            a_text = turn.get('ai', '').strip()
            if u_text or a_text:
                filtered_history.append((u_text, a_text))
        if filtered_history:
            prompt_parts.append("--- Previous Conversation ---")
            for u_text, a_text in filtered_history:
                if u_text: prompt_parts.append(f"User: {u_text}")
                if a_text: prompt_parts.append(f"Assistant: {a_text}")
            prompt_parts.append("--- New Message ---")

    prompt_parts.append(f"User: {question.strip() or 'Hello'}")

    if code and code.strip():
        prompt_parts.append("Related Code:\n```\n" + code.strip() + "\n```")
    
    if image_path:
        # Dosya uzantısını kontrol et
        file_ext = os.path.splitext(image_path)[1].lower()
        text_extensions = ['.txt', '.py', '.js', '.jsx', '.ts', '.tsx', '.json', '.md', '.csv', 
                          '.html', '.css', '.xml', '.yaml', '.yml', '.log', '.sql', '.sh', 
                          '.bat', '.ps1', '.c', '.cpp', '.h', '.java', '.rb', '.go', '.rs', 
                          '.php', '.swift', '.kt', '.r', '.m']
        
        if file_ext in text_extensions:
            # Metin dosyası - içeriği oku ve prompt'a ekle
            try:
                with open(image_path, 'r', encoding='utf-8', errors='ignore') as f:
                    file_content = f.read()
                file_name = os.path.basename(image_path)
                prompt_parts.append(f"\n--- Uploaded File: {file_name} ---\n```\n{file_content}\n```\n")
                prompt_parts.append("Analyze the uploaded file content and answer the user's question about it.")
                print(f"Text file read successfully: {file_name} ({len(file_content)} chars)")
            except Exception as e:
                print(f"Text file read error: {e}")
        
        elif file_ext == '.pdf':
            # PDF dosyası - pypdf ile oku
            try:
                from pypdf import PdfReader
                reader = PdfReader(image_path)
                pdf_text = ""
                for page in reader.pages:
                    pdf_text += page.extract_text() or ""
                file_name = os.path.basename(image_path)
                prompt_parts.append(f"\n--- Uploaded PDF: {file_name} ---\n{pdf_text[:10000]}\n")  # İlk 10K karakter
                prompt_parts.append("Analyze the uploaded PDF content and answer the user's question about it.")
                print(f"PDF file read successfully: {file_name} ({len(pdf_text)} chars)")
            except Exception as e:
                print(f"PDF read error: {e}")
        
        elif file_ext in ['.doc', '.docx']:
            # Word dosyası - python-docx ile oku
            try:
                from docx import Document
                doc = Document(image_path)
                doc_text = "\n".join([para.text for para in doc.paragraphs])
                file_name = os.path.basename(image_path)
                prompt_parts.append(f"\n--- Uploaded Word Document: {file_name} ---\n{doc_text[:10000]}\n")  # İlk 10K karakter
                prompt_parts.append("Analyze the uploaded Word document content and answer the user's question about it.")
                print(f"Word file read successfully: {file_name} ({len(doc_text)} chars)")
            except Exception as e:
                print(f"Word read error: {e}")
        
        else:
            # Resim veya Ses dosyası - Gemini için Part nesnesi kullan
            try:
                import PIL.Image
                import base64
                import mimetypes
                
                # MIME type'ı tespit et
                mime_type, _ = mimetypes.guess_type(image_path)
                
                # Ses dosyası kontrolü
                audio_extensions = ['.mp3', '.wav', '.webm', '.m4a', '.ogg', '.aac', '.flac']
                if any(file_ext == ext for ext in audio_extensions):
                    if not mime_type or not mime_type.startswith('audio/'):
                        if file_ext == '.webm': mime_type = 'audio/webm'
                        elif file_ext == '.mp3': mime_type = 'audio/mpeg'
                        elif file_ext == '.wav': mime_type = 'audio/wav'
                        elif file_ext == '.m4a': mime_type = 'audio/mp4'
                        elif file_ext == '.ogg': mime_type = 'audio/ogg'
                        elif file_ext == '.aac': mime_type = 'audio/aac'
                        elif file_ext == '.flac': mime_type = 'audio/flac'
                    
                    print(f"Audio processing: {image_path} ({mime_type})")
                    
                    with open(image_path, 'rb') as audio_file:
                        audio_bytes = audio_file.read()
                    
                    print(f"Audio file size: {len(audio_bytes)} bytes")
                    
                    # SDK dictionary format
                    audio_part = {
                        "mime_type": mime_type,
                        "data": audio_bytes
                    }
                    prompt_parts.append(audio_part)
                    
                    # Eğer kullanıcı metin yazmadıysa, ses mesajı için özel talimat ekle
                    if not question.strip() or question.strip() == 'Hello':
                        prompt_parts.append(
                            "The user has sent a voice message. Please:\n"
                            "1. First, transcribe what the user said in the audio.\n"
                            "2. Then, respond to their message/question appropriately.\n"
                            "Format your response as:\n"
                            "**You said:** [transcription]\n\n"
                            "[Your response to their message]"
                        )
                    else:
                        prompt_parts.append(
                            "The user has sent a voice message along with their text. "
                            "Listen to the audio and consider both the audio content and the written text when responding."
                        )
                    print("Audio part added to prompt successfully")
                    
                else:
                    # Resim (Default fallback)
                    if not mime_type or not mime_type.startswith('image/'):
                        mime_type = 'image/jpeg'  # Varsayılan
                    
                    # Resmi base64'e çevir
                    with open(image_path, 'rb') as img_file:
                        image_bytes = img_file.read()
                    
                    # SDK dictionary format
                    image_part = {
                        "mime_type": mime_type,
                        "data": image_bytes
                    }
                    
                    prompt_parts.append(image_part)
                    prompt_parts.append("Answer the question related to this image.")
                    print(f"Image added successfully: {os.path.basename(image_path)} ({mime_type})")

            except Exception as e:
                print(f"Media upload error: {e}")
                import traceback
                traceback.print_exc()
                yield f"Error processing media file: {str(e)}"
                return

    # Sadece kod varsa veya teknik soru gibiyse maddeler halinde yanıtla
    if code and code.strip():
        prompt_parts.append("Answer in bullet points and support with example code.")
    
    # Fallback Zinciri Üzerinde Dön
    model_success = False
    
    for model_name in fallback_chain:
        current_model_id = f"models/{model_name}" if not model_name.startswith("models/") else model_name
        
        try:
            print(f"Gemini Deneniyor: {current_model_id}")
            # model = genai.GenerativeModel(current_model_id)
            model = _GeminiCompatModel(genai, current_model_id)
            
            # Remove "models/" prefix for clean comparison
            clean_first = fallback_chain[0].replace('models/', '')
            clean_current = model_name.replace('models/', '')

            if clean_current != clean_first and model_name != fallback_chain[0]:
                yield f"\n\n*> [System]: Previous model failed, trying **{clean_current}**...*\n\n"

            # Use the compat model which has the streaming filter
            # Separate the system prompt from the content parts for better instruction following
            system_instruction = prompt_parts[0] if prompt_parts else None
            user_contents = prompt_parts[1:] if len(prompt_parts) > 1 else []
            
            # If there's no user content yet (e.g. only system prompt), the model call might fail.
            # Ensure at least one user part exists.
            if not user_contents:
                user_contents = [question or "Hello"]

            response_iter = model.generate_content(
                user_contents, 
                stream=True, 
                question=question,
                system_instruction=system_instruction
            )

            
            for chunk in response_iter:
                if chunk.text:
                    yield chunk.text
            
            model_success = True
            break # Başarılı olduysa döngüden çık

        except Exception as exc:
            error_str = str(exc)
            
            # Hata Analizi
            is_quota = "429" in error_str or "TooManyRequests" in error_str or "quota" in error_str.lower()
            is_not_found = "404" in error_str or "NotFound" in error_str
            is_internal = "500" in error_str or "internal" in error_str.lower()
            
            if is_quota or is_not_found or is_internal or "503" in error_str:
                print(f"Hata ({model_name}): {error_str} -> Sıradaki modele geçiliyor.")
                continue # Sonraki modele geç
            else:
                # Kritik ve bilinmeyen bir hata ise direkt bildir ve dur
                yield f"[Critical Error ({model_name})]: {exc}"
                return

    if not model_success:
        yield "\n\n*> [System]: All Gemini models failed (Quota/Service). Falling back to **Claude Opus 4.6**...*\n\n"
        yield from generate_claude_answer(question, code, history_context, 'claude-opus-4-5-20251101', image_path, prefs, github_context, depth + 1)


def generate_claude_answer(question: str, code: str, history_context: list = None, requested_model: str = None, image_path: str = None, prefs: dict = None, github_context: str = None, depth: int = 0):
    """Claude API çağrısı yapar (Streaming)."""
    if depth > 2:
        yield "[System Message]: Error: Maximum fallback depth reached. AI services are currently unavailable."
        return
    if not claude_client:
        yield "Error: ANTHROPIC_API_KEY missing."
        return
    
    if requested_model and 'claude' in requested_model:
        target_model = requested_model
    else:
        target_model = ANTHROPIC_MODEL if ANTHROPIC_MODEL else 'claude-sonnet-4-5-20250929'

    print(f"Claude İsteği (Stream) şu modelle yapılıyor: {target_model}")

    style_prompt = ""
    if prefs:
        if prefs.get('response_style') == 'concise':
            style_prompt = "Keep your answers very concise and short. "
        elif prefs.get('response_style') == 'detailed':
            style_prompt = "Provide detailed and comprehensive explanations. "

    # User Persona info
    persona_info = ""
    if prefs:
        persona = prefs.get('persona', 'General User')
        expertise = prefs.get('expertise', 'Intermediate')
        interests = ", ".join(prefs.get('interests', []))
        persona_info = f"User Profile: {persona} (Expertise: {expertise}). "
        if interests:
            persona_info += f"User is interested in: {interests}. "

    system_prompt = (
        "You are a helpful AI assistant. Communicate with the user in a natural conversation style. "
        f"{persona_info}"
        f"{style_prompt}"
        "provide detailed technical assistance and give code examples if necessary (in Markdown code block). "
        "IMPORTANT: Always respond in the same language as the user's question (e.g., if the question is in Turkish, respond in Turkish). "
        "For greetings or small talk, respond naturally in 1-2 short sentences only. "
        "Never output internal analysis, background thinking, or instruction labels like <thought>, Thinking:, Input, Role, etc. "
        "CRITICAL: Provide ONLY the final response without any reasoning steps or internal dialogue. "
        "CRITICAL: You CANNOT generate images directly. DO NOT output markdown image links (e.g. ![](/static/...)). If the user asks for an image, explain that you are a text model."
    )

    if github_context:
        system_prompt += f"\n\nCRITICAL: If context, memory, or repository information is provided below, you MUST strictly adhere to it and incorporate it into your response to ensure perfect continuity.\n\nCONTEXT FROM SYSTEM MEMORY OR REPOSITORY:\n{github_context}"

    user_message = f"Question: {question.strip() or 'Unspecified'}"
    if code and code.strip():
        user_message += "\n\nRelated Code:\n```\n" + code.strip() + "\n```"

    messages = []
    if history_context:
        for turn in history_context:
            u_text = turn.get('user', '').strip()
            a_text = turn.get('ai', '').strip()
            if u_text:
                messages.append({"role": "user", "content": u_text})
            if a_text:
                messages.append({"role": "assistant", "content": a_text})
    
    if image_path:
        # Dosya uzantısını kontrol et
        file_ext = os.path.splitext(image_path)[1].lower()
        text_extensions = ['.txt', '.py', '.js', '.jsx', '.ts', '.tsx', '.json', '.md', '.csv', 
                          '.html', '.css', '.xml', '.yaml', '.yml', '.log', '.sql', '.sh', 
                          '.bat', '.ps1', '.c', '.cpp', '.h', '.java', '.rb', '.go', '.rs', 
                          '.php', '.swift', '.kt', '.r', '.m']
        
        if file_ext in text_extensions:
            # Metin dosyası - içeriği oku ve mesaja ekle
            try:
                with open(image_path, 'r', encoding='utf-8', errors='ignore') as f:
                    file_content = f.read()
                file_name = os.path.basename(image_path)
                user_message += f"\n\n--- Uploaded File: {file_name} ---\n```\n{file_content}\n```\n"
                user_message += "\nAnalyze the uploaded file content and answer the user's question about it."
                messages.append({"role": "user", "content": user_message})
                print(f"Claude: Text file read successfully: {file_name} ({len(file_content)} chars)")
            except Exception as e:
                print(f"Claude: Text file read error: {e}")
                messages.append({"role": "user", "content": user_message})
        
        elif file_ext == '.pdf':
            # PDF dosyası - pypdf ile oku
            try:
                from pypdf import PdfReader
                reader = PdfReader(image_path)
                pdf_text = ""
                for page in reader.pages:
                    pdf_text += page.extract_text() or ""
                file_name = os.path.basename(image_path)
                user_message += f"\n\n--- Uploaded PDF: {file_name} ---\n{pdf_text[:10000]}\n"
                user_message += "\nAnalyze the uploaded PDF content and answer the user's question about it."
                messages.append({"role": "user", "content": user_message})
                print(f"Claude: PDF file read successfully: {file_name} ({len(pdf_text)} chars)")
            except Exception as e:
                print(f"Claude: PDF read error: {e}")
                messages.append({"role": "user", "content": user_message})
        
        elif file_ext in ['.doc', '.docx']:
            # Word dosyası - python-docx ile oku
            try:
                from docx import Document
                doc = Document(image_path)
                doc_text = "\n".join([para.text for para in doc.paragraphs])
                file_name = os.path.basename(image_path)
                user_message += f"\n\n--- Uploaded Word Document: {file_name} ---\n{doc_text[:10000]}\n"
                user_message += "\nAnalyze the uploaded Word document content and answer the user's question about it."
                messages.append({"role": "user", "content": user_message})
                print(f"Claude: Word file read successfully: {file_name} ({len(doc_text)} chars)")
            except Exception as e:
                print(f"Claude: Word read error: {e}")
                messages.append({"role": "user", "content": user_message})
        
        else:
            # Resim dosyası - mevcut mantık
            try:
                with open(image_path, "rb") as image_file:
                    encoded_string = base64.b64encode(image_file.read()).decode('utf-8')
                    mime_type, _ = mimetypes.guess_type(image_path)
                    if not mime_type: mime_type = 'image/jpeg'
                    
                    messages.append({
                        "role": "user",
                        "content": [
                            {"type": "image", "source": {"type": "base64", "media_type": mime_type, "data": encoded_string}},
                            {"type": "text", "text": user_message}
                        ]
                    })
            except Exception as e:
                print(f"Claude image error: {e}")
                messages.append({"role": "user", "content": user_message})
    else:
        messages.append({"role": "user", "content": user_message})

    try:
        with claude_client.messages.stream(
            model=target_model,
            max_tokens=4096,
            system=system_prompt,
            messages=messages
        ) as stream:
            for text in stream.text_stream:
                yield text

    except Exception as exc:
        yield f"\n\n*> [System]: Claude Error ({target_model}): {exc}. Falling back to Gemini...*\n\n"
        yield from generate_gemini_answer(question, code, history_context, 'gemini-2.5-flash', image_path, prefs, github_context, depth + 1)


def generate_gpt_answer(question: str, code: str, history_context: list = None, requested_model: str = None, image_path: str = None, prefs: dict = None, github_context: str = None, depth: int = 0):
    """OpenAI GPT API'sini kullanarak cevap üretir (Streaming)."""
    if depth > 2:
        yield "[System Message]: Error: Maximum fallback depth reached. AI services are currently unavailable."
        return
    if not openai_client:
        if openai_init_error:
            yield f"Error: OpenAI client init failed: {openai_init_error}"
        else:
            yield "Error: OPENAI_API_KEY missing."
        return

    if requested_model and 'gpt' in requested_model:
        target_model = requested_model
    else:
        target_model = OPENAI_MODEL

    # Model info logged without sensitive data
    print(f"GPT Request (Stream) with model: {target_model}")

    style_prompt = ""
    if prefs:
        if prefs.get('response_style') == 'concise':
            style_prompt = "Keep your answers very concise and short. "
        elif prefs.get('response_style') == 'detailed':
            style_prompt = "Provide detailed and comprehensive explanations. "

    # User Persona info
    persona_info = ""
    if prefs:
        persona = prefs.get('persona', 'General User')
        expertise = prefs.get('expertise', 'Intermediate')
        interests = ", ".join(prefs.get('interests', []))
        persona_info = f"User Profile: {persona} (Expertise: {expertise}). "
        if interests:
            persona_info += f"User is interested in: {interests}. "

    system_prompt = (
        "You are a helpful AI assistant. Communicate with the user in a natural conversation style. "
        f"{persona_info}"
        f"{style_prompt}"
        "If the user asks a question about code, software, or a technical topic, "
        "provide detailed technical assistance and give code examples if necessary (in Markdown code block). "
        "IMPORTANT: Always respond in the same language as the user's question (e.g., if the question is in Turkish, respond in Turkish). "
        "For greetings or small talk, respond naturally in 1-2 short sentences only. "
        "Never output internal analysis, background thinking, or instruction labels like <thought>, Thinking:, Input, Role, etc. "
        "CRITICAL: Provide ONLY the final response without any reasoning steps or internal dialogue."
    )
    
    if github_context:
        system_prompt += f"\n\nCRITICAL: If context, memory, or repository information is provided below, you MUST strictly adhere to it and incorporate it into your response to ensure perfect continuity.\n\nCONTEXT FROM SYSTEM MEMORY OR REPOSITORY:\n{github_context}"

    user_message = f"Question: {question.strip() or 'Unspecified'}"
    if code and code.strip():
        user_message += "\n\nRelated Code:\n```\n" + code.strip() + "\n```"

    if github_context:
        system_prompt += f"\n\nCRITICAL: If context, memory, or repository information is provided below, you MUST strictly adhere to it and incorporate it into your response to ensure perfect continuity.\n\nCONTEXT FROM SYSTEM MEMORY OR REPOSITORY:\n{github_context}"

    messages = [{"role": "system", "content": system_prompt}]
    
    if history_context:
        for turn in history_context:
            u_text = turn.get('user', '').strip()
            a_text = turn.get('ai', '').strip()
            if u_text:
                messages.append({"role": "user", "content": u_text})
            if a_text:
                messages.append({"role": "assistant", "content": a_text})

    if image_path:
        # Dosya uzantısını kontrol et
        file_ext = os.path.splitext(image_path)[1].lower()
        text_extensions = ['.txt', '.py', '.js', '.jsx', '.ts', '.tsx', '.json', '.md', '.csv', 
                          '.html', '.css', '.xml', '.yaml', '.yml', '.log', '.sql', '.sh', 
                          '.bat', '.ps1', '.c', '.cpp', '.h', '.java', '.rb', '.go', '.rs', 
                          '.php', '.swift', '.kt', '.r', '.m']
        
        if file_ext in text_extensions:
            # Metin dosyası - içeriği oku ve mesaja ekle
            try:
                with open(image_path, 'r', encoding='utf-8', errors='ignore') as f:
                    file_content = f.read()
                file_name = os.path.basename(image_path)
                user_message += f"\n\n--- Uploaded File: {file_name} ---\n```\n{file_content}\n```\n"
                user_message += "\nAnalyze the uploaded file content and answer the user's question about it."
                messages.append({"role": "user", "content": user_message})
                print(f"GPT: Text file read successfully: {file_name} ({len(file_content)} chars)")
            except Exception as e:
                print(f"GPT: Text file read error: {e}")
                messages.append({"role": "user", "content": user_message})
        
        elif file_ext == '.pdf':
            # PDF dosyası - pypdf ile oku
            try:
                from pypdf import PdfReader
                reader = PdfReader(image_path)
                pdf_text = ""
                for page in reader.pages:
                    pdf_text += page.extract_text() or ""
                file_name = os.path.basename(image_path)
                user_message += f"\n\n--- Uploaded PDF: {file_name} ---\n{pdf_text[:10000]}\n"
                user_message += "\nAnalyze the uploaded PDF content and answer the user's question about it."
                messages.append({"role": "user", "content": user_message})
                print(f"GPT: PDF file read successfully: {file_name} ({len(pdf_text)} chars)")
            except Exception as e:
                print(f"GPT: PDF read error: {e}")
                messages.append({"role": "user", "content": user_message})
        
        elif file_ext in ['.doc', '.docx']:
            # Word dosyası - python-docx ile oku
            try:
                from docx import Document
                doc = Document(image_path)
                doc_text = "\n".join([para.text for para in doc.paragraphs])
                file_name = os.path.basename(image_path)
                user_message += f"\n\n--- Uploaded Word Document: {file_name} ---\n{doc_text[:10000]}\n"
                user_message += "\nAnalyze the uploaded Word document content and answer the user's question about it."
                messages.append({"role": "user", "content": user_message})
                print(f"GPT: Word file read successfully: {file_name} ({len(doc_text)} chars)")
            except Exception as e:
                print(f"GPT: Word read error: {e}")
                messages.append({"role": "user", "content": user_message})
        
        else:
            # Resim dosyası - mevcut mantık
            try:
                with open(image_path, "rb") as image_file:
                    encoded_string = base64.b64encode(image_file.read()).decode('utf-8')
                    mime_type, _ = mimetypes.guess_type(image_path)
                    if not mime_type: mime_type = 'image/jpeg'

                    messages.append({
                        "role": "user",
                        "content": [
                            {"type": "text", "text": user_message},
                            {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{encoded_string}"}}
                        ]
                    })
            except Exception as e:
                print(f"GPT image error: {e}")
                messages.append({"role": "user", "content": user_message})
    else:
        messages.append({"role": "user", "content": user_message})

    try:
        stream = openai_client.chat.completions.create(
            model=target_model, 
            messages=messages,
            temperature=0.7,
            max_tokens=2000,
            stream=True
        )
        for chunk in stream:
            if chunk.choices[0].delta.content:
                yield chunk.choices[0].delta.content

    except Exception as e:
        print("OPENAI ERROR:", repr(e))
        error_kind = _classify_openai_error(e)
        print(f"[GPT DEBUG] model={target_model} kind={error_kind} error={e}")
        traceback.print_exc()
        yield (
            f"\n\n*> [System]: OpenAI Error ({target_model}) [{error_kind}]: {e}. "
            f"GPT request stopped because the explicit GPT model failed.*\n\n"
        )
        return


def generate_conversation_title(question: str, answer: str = None):
    """Sohbet için kısa ve öz bir başlık üretir."""
    if not question:
        return "New Chat"
        
    prompt = f"""
    Create a very short, catchy title (max 5-6 words) for the following user question.
    Use the same language as the user.
    Do NOT use quotes or special characters.
    
    Question: {question}
    Answer Summary: {answer[:200] if answer else "None"}
    
    Title:"""
    
    # Try Gemini -> GPT -> Claude
    try:
        if GEMINI_API_KEY:
            # Try multiple gemini models
            for m_name in ['models/gemini-2.5-flash-lite', 'models/gemini-2.5-flash']:
                try:
                    model = genai.GenerativeModel(m_name)
                    response = model.generate_content(prompt)
                    title = response.text.strip().replace('"', '').replace("'", "")
                    if title: return title
                except:
                    continue

        if openai_client:
            try:
                response = openai_client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=20
                )
                title = response.choices[0].message.content.strip().replace('"', '').replace("'", "")
                if title: return title
            except:
                pass
            
        return question[:30] + "..."
    except:
        return question[:30] + "..."


def summarize_answer(answer: str) -> str:
    if not answer:
        return ""

    prompt = (
        "Summarize the following AI answer in one short sentence. "
        "Focus on the main technical solution or advice. "
        "Keep it under 15 words:\n\n"
        f"{answer[:1000]}"
    )

    # Try Gemini -> GPT -> Claude
    try:
        if GEMINI_API_KEY:
            for m_name in ['models/gemini-2.5-flash-lite', 'models/gemini-2.5-flash']:
                try:
                    model = genai.GenerativeModel(m_name)
                    response = model.generate_content(prompt)
                    summary = response.text.strip()
                    if summary: return summary
                except: continue

        if openai_client:
            try:
                response = openai_client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=30
                )
                summary = response.choices[0].message.content.strip()
                if summary: return summary
            except: pass
            
        return answer[:100] + "..."
    except:
        return answer[:100] + "..."


def _parse_bool(value):
    if isinstance(value, bool):
        return value
    if value is None:
        return False
    return str(value).strip().lower() in {'1', 'true', 'yes', 'on'}


def _safe_json_list(value):
    if not value:
        return []
    if isinstance(value, list):
        return value
    try:
        parsed = json.loads(value)
        return parsed if isinstance(parsed, list) else []
    except Exception:
        return []


def _resolve_include_previous_modules(user, payload, no_save=False):
    """Enable retrieval-based continuity for authenticated chat unless explicitly disabled."""
    if not user or no_save:
        return False

    raw_value = None
    if payload and isinstance(payload, dict):
        raw_value = payload.get('include_previous_modules')

    # Explicit user preference always wins when provided.
    if raw_value is not None:
        return _parse_bool(raw_value)

    return True


def _is_context_reset_intent(question_text):
    text = (question_text or '').lower()
    reset_signals = [
        'from scratch',
        'start over',
        'ignore previous',
        'yeni bastan',
        'sifirdan',
        'eskiyi yok say',
        'fikri degistir',
        'karar degisti',
    ]
    return any(signal in text for signal in reset_signals)


def _load_previous_memory_context(user, question, conversation=None, include_previous_modules=False):
    if not user or not include_previous_modules:
        return {'text': '', 'hit_count': 0, 'hits': []}

    if _is_context_reset_intent(question):
        return {
            'text': '',
            'hit_count': 0,
            'hits': [],
            'retrieval_mode': 'reset-intent-skip',
            'focus_module': None,
        }

    memory_rows = (
        MemoryItem.query.filter_by(user_id=user.id)
        .order_by(
            MemoryItem.importance.desc(),
            MemoryItem.last_used_at.desc().nullslast(),
            MemoryItem.created_at.desc(),
        )
        .limit(50)
        .all()
    )

    # Backfill-style fallback: older chats may predate newer memory extraction
    # rules, so derive ephemeral memory candidates from recent saved turns.
    history_memory_candidates = []
    try:
        recent_history_query = (
            History.query
            .join(Conversation, History.conversation_id == Conversation.id)
            .filter(Conversation.user_id == user.id, History.is_deleted == False)  # noqa: E712
        )
        if conversation and conversation.id:
            recent_history_query = recent_history_query.filter(History.conversation_id != conversation.id)

        recent_history_rows = (
            recent_history_query
            .order_by(History.timestamp.desc())
            .limit(80)
            .all()
        )

        seen_history_candidates = set()
        for history_row in recent_history_rows:
            for item in extract_memory_candidates(history_row.user_question or '', history_row.ai_response or ''):
                content = (item.get('content') or '').strip()
                module_key = item.get('module_key') or 'general'
                key = (module_key, content.lower())
                if not content or key in seen_history_candidates:
                    continue
                seen_history_candidates.add(key)
                history_memory_candidates.append({
                    'id': None,
                    'memory_type': item.get('memory_type') or 'preference',
                    'module_key': module_key,
                    'content': content,
                    'importance': item.get('importance') or 1,
                    'last_used_at': history_row.timestamp,
                    'created_at': history_row.timestamp,
                    'updated_at': history_row.timestamp,
                })
    except Exception as history_memory_err:
        print(f"[MEMORY] Recent history fallback skipped: {history_memory_err}")

    memory_sources = list(memory_rows) + history_memory_candidates

    summary_query = ConversationSummary.query.filter_by(user_id=user.id)
    if conversation and conversation.id:
        summary_query = summary_query.filter(ConversationSummary.conversation_id != conversation.id)

    summary_rows = (
        summary_query.order_by(ConversationSummary.updated_at.desc())
        .limit(1)
        .all()
    )

    memory_context = build_minimum_continuation_capsule(
        question,
        memory_sources,
        summary_rows,
        char_budget=420,
        max_lines=5,
        min_confidence=0.42,
    )

    conflict_result = detect_memory_conflicts(question, memory_context.get('hits', []))
    if conflict_result.get('has_conflict'):
        drop_module_keys = set(conflict_result.get('drop_module_keys', []))
        filtered_hits = [
            hit for hit in memory_context.get('hits', [])
            if (hit.get('module_key') or 'general') not in drop_module_keys
        ]

        memory_context['hits'] = filtered_hits
        memory_context['hit_count'] = len(filtered_hits)
        memory_context['conflicts'] = conflict_result.get('conflicts', [])
        if filtered_hits:
            memory_context['text'] = '[Memory Capsule]\n' + '\n'.join(
                f"- {(hit.get('module_key') or 'general').title()}: {(hit.get('content') or hit.get('summary_text') or '').strip()} (score={float(hit.get('score') or 0.0):.2f})"
                for hit in filtered_hits[:5]
            )
            memory_context['focus_module'] = filtered_hits[0].get('module_key')
        else:
            memory_context['text'] = ''
            memory_context['focus_module'] = None

    retrieval_plan = build_memory_retrieval_plan(
        question,
        memory_sources,
        summary_rows,
        top_k=5,
        min_confidence=0.42,
    )
    if retrieval_plan.get('text'):
        memory_context['text'] = retrieval_plan.get('text', '')
        memory_context['hit_count'] = retrieval_plan.get('hit_count', 0)
        memory_context['hits'] = retrieval_plan.get('hits', [])
        memory_context['focus_module'] = retrieval_plan.get('focus_module')

    memory_context['structured_capsule'] = retrieval_plan.get('structured_capsule')
    memory_context['memory_graph'] = retrieval_plan.get('memory_graph')
    memory_context['memory_transitions'] = retrieval_plan.get('transitions')
    memory_context['retrieval_plan_steps'] = retrieval_plan.get('plan_steps', [])

    used_memory_ids = [hit.get('source_id') for hit in memory_context.get('hits', []) if hit.get('source_type') == 'memory' and hit.get('source_id')]
    if used_memory_ids:
        MemoryItem.query.filter(MemoryItem.id.in_(used_memory_ids)).update(
            {MemoryItem.last_used_at: _utcnow()},
            synchronize_session=False,
        )

    return memory_context


def _upsert_conversation_summary(current_conv, history_id, user_id, summary_text, extracted_memory_items):
    module_keys = sorted({item.get('module_key') for item in extracted_memory_items if item.get('module_key')})
    summary_row = ConversationSummary.query.filter_by(conversation_id=current_conv.id).first()

    if not summary_row:
        summary_row = ConversationSummary(
            user_id=user_id,
            conversation_id=current_conv.id,
            project_id=current_conv.project_id,
        )
        db.session.add(summary_row)

    existing_modules = _safe_json_list(summary_row.modules_json)
    merged_modules = sorted({*existing_modules, *module_keys})

    summary_row.user_id = user_id
    summary_row.conversation_id = current_conv.id
    summary_row.project_id = current_conv.project_id
    summary_row.summary_text = summary_text
    summary_row.modules_json = json.dumps(merged_modules, ensure_ascii=False)
    summary_row.last_history_id = history_id
    summary_row.updated_at = _utcnow()

    return summary_row


def _store_memory_items(user_id, conversation_id, extracted_memory_items):
    stored_items = []
    now = _utcnow()

    for item in extracted_memory_items:
        content = (item.get('content') or '').strip()
        if not content:
            continue

        memory_type = item.get('memory_type') or 'preference'
        module_key = item.get('module_key') or 'general'
        importance = int(item.get('importance') or 1)

        existing_item = MemoryItem.query.filter_by(
            user_id=user_id,
            source_conversation_id=conversation_id,
            memory_type=memory_type,
            module_key=module_key,
            content=content,
        ).first()

        if existing_item:
            existing_item.importance = max(int(existing_item.importance or 1), importance)
            existing_item.last_used_at = now
            stored_items.append(existing_item)
            continue

        memory_item = MemoryItem(
            user_id=user_id,
            source_conversation_id=conversation_id,
            memory_type=memory_type,
            module_key=module_key,
            content=content,
            importance=importance,
            last_used_at=now,
        )
        db.session.add(memory_item)
        stored_items.append(memory_item)

    return stored_items


def _memory_node_uid_from_entry(entry):
    source_type = entry.get('source_type') or 'memory'
    source_id = entry.get('source_id')
    module_key = entry.get('module_key') or 'general'
    summary = (entry.get('summary') or entry.get('content') or '')[:120]
    if source_id is not None:
        return f'{source_type}:{source_id}'
    digest = hashlib.sha1(f'{source_type}:{module_key}:{summary}'.encode('utf-8')).hexdigest()[:16]
    return f'{source_type}:{digest}'


def _safe_float(value, default=0.0):
    try:
        return float(value)
    except Exception:
        return default


def _memory_guardrail(user_id, module_key, node_uid, transition_action=None):
    recent_same_module = MemoryNode.query.filter_by(user_id=user_id, module_key=module_key).order_by(
        MemoryNode.last_accessed_at.desc(),
        MemoryNode.updated_at.desc(),
    ).limit(8).all()

    repeat_signals = sum(
        1 for item in recent_same_module
        if item.node_uid != node_uid and not item.is_deleted and item.validity_state == 'active'
    )

    dampening = 1.0
    anti_loop = repeat_signals >= 3
    if anti_loop:
        dampening *= 0.65

    if transition_action == 'drop':
        dampening *= 0.35
    elif transition_action == 'update':
        dampening *= 0.75
    elif transition_action == 'merge':
        dampening *= 0.85

    return {
        'anti_loop': anti_loop,
        'repeat_signals': repeat_signals,
        'dampening': dampening,
    }


def _prune_memory_graph_state(user_id, protected_uids=None, max_active_nodes=220):
    protected_uids = set(protected_uids or [])
    active_count = MemoryNode.query.filter_by(user_id=user_id, validity_state='active', is_deleted=False).count()
    if active_count <= max_active_nodes:
        return {
            'triggered': False,
            'pruned_nodes': 0,
            'active_count': active_count,
            'max_active_nodes': max_active_nodes,
        }

    prune_target = active_count - max_active_nodes
    candidates = MemoryNode.query.filter_by(user_id=user_id, validity_state='active', is_deleted=False).order_by(
        MemoryNode.reinforcement_score.asc(),
        MemoryNode.decay_score.desc(),
        MemoryNode.importance.asc(),
        MemoryNode.updated_at.asc(),
    ).all()

    pruned = 0
    for node in candidates:
        if pruned >= prune_target:
            break
        if node.node_uid in protected_uids:
            continue
        node.validity_state = 'archived'
        node.is_deleted = True
        pruned += 1

    return {
        'triggered': True,
        'pruned_nodes': pruned,
        'active_count': active_count,
        'max_active_nodes': max_active_nodes,
    }


def _write_back_memory_graph(user_id, conversation_id, history_id, memory_context):
    structured = memory_context.get('structured_capsule') or {}
    graph = memory_context.get('memory_graph') or {}
    transitions = memory_context.get('memory_transitions') or {}
    entries = list(structured.get('entries') or [])
    transition_map = {item.get('module_key'): item for item in transitions.get('transitions', []) if item.get('module_key')}
    now = _utcnow()

    node_uids = []
    persisted_nodes = []
    guardrail_snapshots = []

    for entry in entries:
        module_key = entry.get('module_key') or 'general'
        node_uid = _memory_node_uid_from_entry(entry)
        node_uids.append(node_uid)
        transition = transition_map.get(module_key)
        guardrail = _memory_guardrail(user_id, module_key, node_uid, transition.get('action') if transition else None)
        guardrail_snapshots.append({
            'module_key': module_key,
            'node_uid': node_uid,
            **guardrail,
        })

        existing = MemoryNode.query.filter_by(user_id=user_id, node_uid=node_uid).first()
        next_state = 'active'
        if transition and transition.get('action') == 'drop':
            next_state = 'deprecated'
        elif existing and existing.validity_state == 'deprecated' and transition and transition.get('action') == 'update':
            next_state = 'active'

        depends_on = []
        conflicts_with = []
        for graph_node in (graph.get('nodes') or []):
            if _memory_node_uid_from_entry(entry) == graph_node.get('id'):
                depends_on = list(graph_node.get('depends_on') or [])
                conflicts_with = list(graph_node.get('conflicts_with') or [])
                break

        summary_text = entry.get('summary') or ''
        content_text = summary_text
        reinforcement = _safe_float(entry.get('relevance_score'), 0.0)
        decay = 0.0
        reinforcement_delta = 0.12 * guardrail['dampening']

        if existing:
            previous_summary = existing.summary_text or ''
            existing.module_key = module_key
            existing.node_type = entry.get('type') or existing.node_type or 'fact'
            existing.summary_text = summary_text
            existing.content = content_text
            existing.depends_on_json = json.dumps(depends_on, ensure_ascii=False)
            existing.conflicts_with_json = json.dumps(conflicts_with, ensure_ascii=False)
            existing.validity_state = next_state
            existing.importance = _safe_float(entry.get('importance'), existing.importance or 0.0)
            existing.relevance_score = max(_safe_float(existing.relevance_score, 0.0), reinforcement)
            existing.semantic_similarity = _safe_float(entry.get('semantic_similarity'), existing.semantic_similarity or 0.0) if entry.get('semantic_similarity') is not None else existing.semantic_similarity
            existing.task_alignment = _safe_float(entry.get('task_alignment'), existing.task_alignment or 0.0) if entry.get('task_alignment') is not None else existing.task_alignment
            existing.reinforcement_score = min(1.0, _safe_float(existing.reinforcement_score, 0.0) + reinforcement_delta)
            existing.decay_score = min(1.0, _safe_float(existing.decay_score, 0.0) * 0.96 + decay)
            existing.usage_count = int(existing.usage_count or 0) + 1
            existing.version = int(existing.version or 1) + (1 if previous_summary != summary_text else 0)
            existing.last_accessed_at = now
            existing.source_history_id = history_id
            existing.conversation_id = conversation_id
            existing.is_deleted = next_state == 'deprecated'
            persisted_nodes.append(existing)
        else:
            node = MemoryNode(
                node_uid=node_uid,
                user_id=user_id,
                conversation_id=conversation_id,
                source_history_id=history_id,
                node_type=entry.get('type') or 'fact',
                module_key=module_key,
                summary_text=summary_text,
                content=content_text,
                depends_on_json=json.dumps(depends_on, ensure_ascii=False),
                conflicts_with_json=json.dumps(conflicts_with, ensure_ascii=False),
                validity_state=next_state,
                importance=_safe_float(entry.get('importance'), 0.0),
                relevance_score=reinforcement,
                semantic_similarity=_safe_float(entry.get('semantic_similarity'), 0.0) if entry.get('semantic_similarity') is not None else None,
                task_alignment=_safe_float(entry.get('task_alignment'), 0.0) if entry.get('task_alignment') is not None else None,
                reinforcement_score=min(1.0, reinforcement + 0.1),
                decay_score=0.0,
                usage_count=1 if not guardrail['anti_loop'] else 0,
                version=1,
                last_accessed_at=now,
                is_deleted=next_state == 'deprecated',
            )
            db.session.add(node)
            persisted_nodes.append(node)

    # Graph edges persist as audit trail.
    persisted_edges = 0
    for graph_edge in (graph.get('edges') or []):
        edge = MemoryEdge.query.filter_by(
            user_id=user_id,
            source_node_uid=graph_edge.get('source_id'),
            target_node_uid=graph_edge.get('target_id'),
            relation_type=graph_edge.get('relation_type'),
        ).first()
        if edge:
            edge.weight = max(_safe_float(edge.weight, 0.0), _safe_float(graph_edge.get('weight'), 1.0))
            edge.created_at = now
        else:
            db.session.add(MemoryEdge(
                user_id=user_id,
                source_node_uid=graph_edge.get('source_id'),
                target_node_uid=graph_edge.get('target_id'),
                relation_type=graph_edge.get('relation_type'),
                weight=_safe_float(graph_edge.get('weight'), 1.0),
                source_module_key=(graph_edge.get('source_id') or '').split(':', 1)[0] if graph_edge.get('source_id') else None,
                target_module_key=(graph_edge.get('target_id') or '').split(':', 1)[0] if graph_edge.get('target_id') else None,
                created_at=now,
            ))
            persisted_edges += 1

    # Reinforce nodes that were actively injected.
    for node in persisted_nodes:
        node.reinforcement_score = min(1.0, _safe_float(node.reinforcement_score, 0.0) + (0.03 if node.validity_state == 'active' else 0.0))
        node.decay_score = max(0.0, _safe_float(node.decay_score, 0.0) * 0.995)

    # Apply gentle decay to older active nodes not used in this turn.
    stale_query = MemoryNode.query.filter(
        MemoryNode.user_id == user_id,
        MemoryNode.validity_state == 'active',
        MemoryNode.is_deleted.is_(False),
    )
    if node_uids:
        stale_query = stale_query.filter(~MemoryNode.node_uid.in_(node_uids))

    stale_nodes = stale_query.limit(50).all()
    for node in stale_nodes:
        node.reinforcement_score = max(0.0, _safe_float(node.reinforcement_score, 0.0) * 0.985)
        node.decay_score = min(1.0, _safe_float(node.decay_score, 0.0) + 0.01)
        if node.decay_score >= 0.7 and node.reinforcement_score < 0.2:
            node.validity_state = 'deprecated'
            node.is_deleted = True

    prune_snapshot = _prune_memory_graph_state(user_id, protected_uids=node_uids, max_active_nodes=220)

    return {
        'nodes_upserted': len(persisted_nodes),
        'edges_upserted': persisted_edges,
        'graph_nodes': len(graph.get('nodes') or []),
        'graph_edges': len(graph.get('edges') or []),
        'planned_steps': list(memory_context.get('retrieval_plan_steps') or []),
        'compaction_audit': memory_context.get('graph_compaction_audit'),
        'debug_signals': memory_context.get('debug_signals'),
        'learning_signals': memory_context.get('learning_signals'),
        'guardrail': {
            'anti_loop': any(item['anti_loop'] for item in guardrail_snapshots),
            'repeat_signals': max([item['repeat_signals'] for item in guardrail_snapshots] or [0]),
            'dampening_floor': min([item['dampening'] for item in guardrail_snapshots] or [1.0]),
        },
        'pruning': prune_snapshot,
    }


# --- GAMIFICATION SYSTEM ---

BADGES = {
    'first_question': {'icon': '🌱', 'name': 'İlk Adım', 'desc': 'İlk soruyu sordun'},
    '10_questions': {'icon': '💬', 'name': 'Meraklı', 'desc': '10 soru sordun'},
    '100_questions': {'icon': '🧠', 'name': 'Bilge', 'desc': '100 soru sordun'},
    'first_share': {'icon': '📢', 'name': 'Paylaşımcı', 'desc': 'Toplulukta ilk paylaşım'},
    '10_shares': {'icon': '🌟', 'name': 'İçerik Üretici', 'desc': '10 paylaşım yaptın'},
    'first_answer': {'icon': '🤝', 'name': 'Yardımsever', 'desc': 'Toplulukta ilk cevap'},
    'streak_7': {'icon': '🔥', 'name': 'Haftalık Seri', 'desc': '7 gün üst üste aktif'},
    'streak_30': {'icon': '⚡', 'name': 'Aylık Seri', 'desc': '30 gün üst üste aktif'},
    '10_likes': {'icon': '❤️', 'name': 'Beğenilen', 'desc': '10 beğeni aldın'},
    'multi_model': {'icon': '⚗️', 'name': 'Simyacı', 'desc': 'Birden fazla AI modeli kullandın'},
    'level_10': {'icon': '👑', 'name': 'Usta', 'desc': '10. Seviyeye ulaştın'}
}

XP_REWARDS = {
    'ask_question': 50,
    'share_solution': 30,
    'community_post': 100,
    'daily_login': 20,
}

def calculate_level(xp):
    """XP'ye göre seviye hesaplar (Seviye = Kök(XP/100) + 1)
    Not: total_xp_earned kullanılarak level hiçbir zaman düşmez"""
    safe_xp = max(int(xp or 0), 0)
    return math.floor(math.sqrt(safe_xp / 100)) + 1

def resolve_effective_progress(user):
    """Kullanıcı için tutarlı total_xp ve level döndürür.
    total_xp_earned eski kayıtlarda boş kalmışsa xp ile tamamlanır.
    """
    safe_xp = max(int((user.xp or 0)), 0)
    safe_total_xp = max(int((user.total_xp_earned or 0)), 0)
    effective_total_xp = max(safe_total_xp, safe_xp)
    calculated_level = calculate_level(effective_total_xp)
    stored_level = max(int((user.level or 1)), 1)
    effective_level = max(stored_level, calculated_level)
    return effective_total_xp, effective_level

def get_level_xp_bounds(level):
    """Verilen seviye için [min_xp, next_level_xp) aralığını döner."""
    safe_level = max(int(level or 1), 1)
    current_level_min_xp = 100 * ((safe_level - 1) ** 2)
    next_level_xp = 100 * (safe_level ** 2)
    return current_level_min_xp, next_level_xp

def get_rank_title(level):
    safe_level = max(int(level or 1), 1)
    if safe_level >= 20:
        return 'Grand Archmage'
    if safe_level >= 15:
        return 'Master Alchemist'
    if safe_level >= 10:
        return 'Arcane Engineer'
    if safe_level >= 6:
        return 'Code Adept'
    if safe_level >= 3:
        return 'Junior Alchemist'
    return 'Novice Alchemist'

def check_and_award_badges(user):
    """Kullanıcının statülerine göre hak ettiği rozetleri kontrol eder ve verir."""
    if not user:
        return []
    
    new_badges = []
    
    # Mevcut rozetleri al
    existing_badges = [b.badge_id for b in user.badges.all()]
    
    stats = {
        'questions': History.query.join(Conversation).filter(Conversation.user_id == user.id).count(),
        'shares': XPEvent.query.filter(
            XPEvent.user_id == user.id,
            XPEvent.source.in_(['share_solution', 'community_post'])
        ).count(),
        'answers': Answer.query.filter_by(author_id=user.id).count(),
        'likes': db.session.query(db.func.sum(History.likes)).join(Conversation).filter(Conversation.user_id == user.id).scalar() or 0
    }
    
    # Soru Badge'leri
    if stats['questions'] >= 1 and 'first_question' not in existing_badges:
        new_badges.append('first_question')
    if stats['questions'] >= 10 and '10_questions' not in existing_badges:
        new_badges.append('10_questions')
    if stats['questions'] >= 100 and '100_questions' not in existing_badges:
        new_badges.append('100_questions')
        
    # Cevap Badge'i
    if stats['answers'] >= 1 and 'first_answer' not in existing_badges:
        new_badges.append('first_answer')

    # Paylaşım Badge'leri
    if stats['shares'] >= 1 and 'first_share' not in existing_badges:
        new_badges.append('first_share')
    if stats['shares'] >= 10 and '10_shares' not in existing_badges:
        new_badges.append('10_shares')
        
    # Beğeni Badge'i
    if stats['likes'] >= 10 and '10_likes' not in existing_badges:
        new_badges.append('10_likes')
        
    # Streak Badge'leri
    if user.streak_days >= 7 and 'streak_7' not in existing_badges:
        new_badges.append('streak_7')
    if user.streak_days >= 30 and 'streak_30' not in existing_badges:
        new_badges.append('streak_30')
        
    # Seviye Badge'i
    if user.level >= 10 and 'level_10' not in existing_badges:
        new_badges.append('level_10')
        
    # Multi-model Badge (Farklı modeller kullanmış mı?)
    if 'multi_model' not in existing_badges:
        prefs = get_user_preferences(user)
        usage = prefs.get('usage_stats', {})
        used_models_count = sum(1 for m, count in usage.items() if count > 0)
        if used_models_count > 1:
            new_badges.append('multi_model')
            
    # Rozetleri veritabanına ekle
    for badge_id in new_badges:
        badge = UserBadge(user_id=user.id, badge_id=badge_id)
        db.session.add(badge)
        
    # Kullanıcıya bildirim gönderilebilir (opsiyonel)
    if new_badges:
        db.session.commit()
        
    return new_badges

def update_streak(user, activity_date=None):
    """Kullanıcının streak (seri) günlerini günceller."""
    if isinstance(activity_date, datetime.date):
        today = activity_date
    elif isinstance(activity_date, str) and activity_date:
        today = datetime.datetime.strptime(activity_date[:10], '%Y-%m-%d').date()
    else:
        today = _utcnow().date()
    
    if user.last_active_date == today:
        return False # Bugün zaten güncellenmiş
        
    if user.last_active_date == today - timedelta(days=1):
        # Dün aktifmiş, streak artar
        user.streak_days += 1
    else:
        # Arada gün boşluk var veya ilk defa, sıfırla/başlat
        user.streak_days = 1
        
    if user.streak_days > user.longest_streak:
        user.longest_streak = user.streak_days
        
    user.last_active_date = today
    return True

def award_xp(user_id, amount, reason="", source="generic", metadata=None, activity_date=None):
    """Kullanıcıya XP verir ve seviyesini günceller.
    Not: total_xp_earned her zaman artar, xp harcama mekanizmalarından etkilenmez."""
    user = db.session.get(User, user_id)
    if not user:
        return None

    old_xp = user.xp or 0
    old_total_xp = user.total_xp_earned or 0
    
    # XP ve total_xp_earned'i güncelle
    user.xp = old_xp + amount
    user.total_xp_earned = old_total_xp + amount
    
    awarded_total = amount
    metadata_json = json.dumps(metadata) if metadata else None

    db.session.add(XPEvent(
        user_id=user.id,
        amount=amount,
        source=source,
        reason=reason,
        metadata_json=metadata_json
    ))
    
    # Streak güncellemesi
    streak_updated = update_streak(user, activity_date=activity_date)
    if streak_updated and user.streak_days > 1:
        streak_bonus = 20
        user.xp += streak_bonus # Streak bonusu
        user.total_xp_earned += streak_bonus  # Total'a da ekle
        awarded_total += streak_bonus
        db.session.add(XPEvent(
            user_id=user.id,
            amount=streak_bonus,
            source='streak_bonus',
            reason='Streak Bonus',
            metadata_json=None
        ))
        
    # Seviye hesaplama - total_xp_earned'e göre (hiçbir zaman düşmez)
    old_level = user.level
    new_level = calculate_level(user.total_xp_earned)
    
    level_up = False
    if new_level > old_level:
        user.level = new_level
        level_up = True
    
    # 💰 COIN KAZANMA MANTIGI - Sadık kullanıcılara bonus
    coins_earned = 0
    if source == 'ask_question':
        coins_earned = 1  # Her soru sorunca 1 coin
    elif source == 'share_solution':
        coins_earned = 3  # Çözüm paylaşınca 3 coin
    elif source == 'community_post':
        coins_earned = 2  # Community post için 2 coin
    elif source in ('received_like_post', 'received_like_answer'):
        coins_earned = 1  # Beğeni için 1 coin
    
    # Streak bonusu - sadık kullanıcılara ek coin
    if user.streak_days >= 14:
        coins_earned += 5
    elif user.streak_days >= 7:
        coins_earned += 2
    elif user.streak_days >= 3:
        coins_earned += 1
    
    user.coins += coins_earned
        
    db.session.commit()
    
    # Rozet kontrolü yap
    earned_badges = check_and_award_badges(user)
    
    # Rozet aldığında ek bonus coin
    if earned_badges:
        bonus_coins = len(earned_badges) * 10
        user.coins += bonus_coins
        db.session.commit()
    
    return {
        'old_xp': old_xp,
        'new_xp': user.xp,
        'awarded_xp': awarded_total,
        'coins_earned': coins_earned,
        'total_coins': user.coins,
        'old_level': old_level,
        'new_level': new_level,
        'level_up': level_up,
        'streak_days': user.streak_days,
        'earned_badges': earned_badges,
        'reason': reason
    }


# --- GAMIFICATION API ROUTES ---

@app.route('/api/gamification/profile', methods=['GET'])
@jwt_required()
def get_gamification_profile():
    user = get_current_user()
    if not user:
        return jsonify({'error': 'User not found'}), 404

    # Level ve total_xp değerlerini tutarlı hale getir
    total_xp_earned, user_level = resolve_effective_progress(user)
    needs_commit = False
    if (user.total_xp_earned or 0) != total_xp_earned:
        user.total_xp_earned = total_xp_earned
        needs_commit = True
    if (user.level or 1) != user_level:
        user.level = user_level
        needs_commit = True
    if needs_commit:
        db.session.commit()

    # Progress hesaplaması da total_xp_earned'e göre
    current_level_min_xp, next_level_xp = get_level_xp_bounds(user_level)
    xp_to_next_level = max(next_level_xp - total_xp_earned, 0)
    level_span = max(next_level_xp - current_level_min_xp, 1)
    progress_percent = round(((total_xp_earned - current_level_min_xp) / level_span) * 100, 1)
    progress_percent = max(0, min(progress_percent, 100))
    # Ensure locale-independent JSON encoding (always use dot, not comma)
    progress_percent = float(progress_percent)

    base_name = (user.display_name or '').strip()
    safe_handle = re.sub(r'[^a-z0-9_]', '', re.sub(r'\s+', '_', base_name.lower()))
    if not safe_handle:
        safe_handle = f'user_{user.id}'

    badges = []
    for b in user.badges.all():
        info = BADGES.get(b.badge_id, {})
        badges.append({
            'id': b.badge_id,
            'earned_at': b.earned_at.isoformat(),
            'icon': info.get('icon', '🏅'),
            'name': info.get('name', b.badge_id),
            'description': info.get('desc', '')
        })
    
    return jsonify({
        'user_id': user.id,
        'display_name': user.display_name,
        'username': safe_handle,
        'xp': user.xp or 0,
        'total_xp_earned': total_xp_earned,
        'coins': user.coins or 0,
        'level': user_level,
        'rank_title': get_rank_title(user_level),
        'current_level_min_xp': current_level_min_xp,
        'next_level_xp': next_level_xp,
        'xp_to_next_level': xp_to_next_level,
        'progress_percent': progress_percent,
        'streak_days': user.streak_days,
        'longest_streak': user.longest_streak,
        'last_active_date': user.last_active_date.isoformat() if user.last_active_date else None,
        'badges': badges,
        'all_badges_info': BADGES
    })

@app.route('/api/gamification/leaderboard', methods=['GET'])
@jwt_required()
def get_leaderboard():
    # En yüksek total_xp_earned'e sahip aktif kullanıcıları getir (0 XP doldurma yapma)
    top_users = (
        User.query
        .filter(db.func.coalesce(User.total_xp_earned, 0) > 0)
        .order_by(db.func.coalesce(User.total_xp_earned, 0).desc(), User.id.asc())
        .limit(20)
        .all()
    )
    
    leaderboard = []
    needs_commit = False
    for idx, u in enumerate(top_users):
        total_xp, user_level = resolve_effective_progress(u)
        if (u.total_xp_earned or 0) != total_xp:
            u.total_xp_earned = total_xp
            needs_commit = True
        if (u.level or 1) != user_level:
            u.level = user_level
            needs_commit = True
        leaderboard.append({
            'user_id': u.id,
            'display_name': u.display_name,
            'profile_image': _serialize_profile_image(u.profile_image),
            'level': user_level,
            'total_xp': total_xp,
            'rank': idx + 1,
            'badges_count': u.badges.count()
        })
    if needs_commit:
        db.session.commit()
        
    return jsonify({'leaderboard': leaderboard})

@app.route('/api/gamification/sync', methods=['POST'])
@jwt_required()
def sync_gamification():
    """Client tarafından tetiklenen günlük giriş XP'si kazanımı."""
    user = get_current_user()
    if not user:
         return jsonify({'error': 'User not found'}), 404

    data = request.get_json(silent=True) or {}
    activity_date = data.get('activity_date')

    if XPEvent.query.filter(
        XPEvent.user_id == user.id,
        XPEvent.source == 'daily_login',
        XPEvent.reason == 'Daily Login',
        db.func.date(XPEvent.created_at) == (activity_date[:10] if isinstance(activity_date, str) and activity_date else _utcnow().date().isoformat())
    ).first():
        return jsonify({'status': 'noop', 'message': 'Daily login already synced today.'})
         
    # update_streak zaten award_xp içinde çalışıyor
    result = award_xp(
        user.id,
        XP_REWARDS['daily_login'],
        "Daily Login",
        source='daily_login',
        activity_date=activity_date
    )
    
    return jsonify(result)

def get_user_preferences(user):
    """Kullanıcının AI Taste Profile bilgilerini JSON olarak döner."""
    if not user or not user.preferences:
        return {
            "preferred_model": "auto",
            "response_style": "balanced",
            "fav_language": "natural",
            "usage_stats": {"claude": 0, "gemini": 0, "gpt": 0},
            "subscription_plan": "free",
            "persona": "General User",
            "expertise": "Mid-level",
            "interests": []
        }
    try:
        return json.loads(user.preferences)
    except:
        return {
            "preferred_model": "auto",
            "response_style": "balanced",
            "fav_language": "natural",
            "usage_stats": {"claude": 0, "gemini": 0, "gpt": 0},
            "subscription_plan": "free",
            "persona": "General User",
            "expertise": "Mid-level",
            "interests": []
        }


def _normalize_subscription_plan(raw_plan):
    return 'premium' if str(raw_plan or '').strip().lower() == 'premium' else 'free'


def _estimate_tokens_for_request(question, code=''):
    chars = len(question or '') + len(code or '')
    return max(1, math.ceil(chars / 4))


def _build_usage_payload(plan, daily_count, month_tokens):
    limits = PLAN_LIMITS.get(plan, PLAN_LIMITS['free'])
    daily_remaining = max(0, limits['daily_requests'] - daily_count)
    monthly_remaining = max(0, limits['monthly_tokens'] - month_tokens)
    return {
        'plan': plan,
        'limits': limits,
        'usage': {
            'daily_requests_used': daily_count,
            'monthly_tokens_used': month_tokens,
            'daily_requests_remaining': daily_remaining,
            'monthly_tokens_remaining': monthly_remaining,
        }
    }


def _get_or_init_usage_counters(user):
    prefs = get_user_preferences(user)
    plan = _normalize_subscription_plan(prefs.get('subscription_plan', 'free'))

    usage_limits = prefs.get('usage_limits') or {}
    today = _utcnow().date().isoformat()
    month_key = _utcnow().strftime('%Y-%m')

    daily = usage_limits.get('daily') or {}
    monthly = usage_limits.get('monthly') or {}

    if daily.get('date') != today:
        daily = {'date': today, 'count': 0}

    if monthly.get('month') != month_key:
        monthly = {'month': month_key, 'tokens': 0}

    usage_limits['daily'] = daily
    usage_limits['monthly'] = monthly
    prefs['usage_limits'] = usage_limits
    prefs['subscription_plan'] = plan

    return prefs, plan, usage_limits


def consume_plan_quota(user, estimated_tokens, request_weight=1):
    """Consume request/token quota for authenticated users."""
    if not user:
        return {
            'allowed': True,
            'plan': 'guest',
            'limits': None,
            'usage': None,
            'reason': None,
        }

    prefs, plan, usage_limits = _get_or_init_usage_counters(user)
    limits = PLAN_LIMITS.get(plan, PLAN_LIMITS['free'])

    daily_count = int(usage_limits['daily'].get('count', 0) or 0)
    month_tokens = int(usage_limits['monthly'].get('tokens', 0) or 0)

    projected_daily = daily_count + max(1, int(request_weight or 1))
    projected_monthly_tokens = month_tokens + max(1, int(estimated_tokens or 1))

    if projected_daily > limits['daily_requests']:
        payload = _build_usage_payload(plan, daily_count, month_tokens)
        return {
            'allowed': False,
            'plan': plan,
            'limits': limits,
            'usage': payload['usage'],
            'reason': 'daily_requests_exceeded',
        }

    if projected_monthly_tokens > limits['monthly_tokens']:
        payload = _build_usage_payload(plan, daily_count, month_tokens)
        return {
            'allowed': False,
            'plan': plan,
            'limits': limits,
            'usage': payload['usage'],
            'reason': 'monthly_tokens_exceeded',
        }

    usage_limits['daily']['count'] = projected_daily
    usage_limits['monthly']['tokens'] = projected_monthly_tokens
    prefs['usage_limits'] = usage_limits
    user.preferences = json.dumps(prefs)
    db.session.commit()

    payload = _build_usage_payload(plan, projected_daily, projected_monthly_tokens)
    return {
        'allowed': True,
        'plan': plan,
        'limits': limits,
        'usage': payload['usage'],
        'reason': None,
    }


@app.route('/api/billing/usage', methods=['GET'])
@jwt_required()
def get_billing_usage():
    user = get_current_user()
    if not user:
        return jsonify({'error': 'User not found'}), 404

    prefs, plan, usage_limits = _get_or_init_usage_counters(user)
    daily_count = int(usage_limits['daily'].get('count', 0) or 0)
    month_tokens = int(usage_limits['monthly'].get('tokens', 0) or 0)
    user.preferences = json.dumps(prefs)
    db.session.commit()

    payload = _build_usage_payload(plan, daily_count, month_tokens)
    return jsonify(payload)


@app.route('/api/billing/plan', methods=['POST'])
@jwt_required()
def update_billing_plan():
    user = get_current_user()
    if not user:
        return jsonify({'error': 'User not found'}), 404

    data = request.json or {}
    requested_plan = _normalize_subscription_plan(data.get('plan', 'free'))

    prefs, plan, usage_limits = _get_or_init_usage_counters(user)
    prefs['subscription_plan'] = requested_plan
    user.preferences = json.dumps(prefs)
    db.session.commit()

    daily_count = int(usage_limits['daily'].get('count', 0) or 0)
    month_tokens = int(usage_limits['monthly'].get('tokens', 0) or 0)
    payload = _build_usage_payload(requested_plan, daily_count, month_tokens)
    payload['message'] = f"Plan updated to {requested_plan}."
    return jsonify(payload)

def update_user_taste(user, model_used, answer_text, user_question=""):
    """Kullanıcının tercihlerini ve personasını analiz ederek otomatik günceller."""
    if not user:
        return
    
    # DetachedInstanceError fix: Re-fetch or merge user in the current session
    try:
        user = db.session.merge(user)
    except Exception as e:
        print(f"User merge error: {e}")
        return

    prefs = get_user_preferences(user)
    
    # 1. Model kullanımını takip et
    if 'usage_stats' not in prefs:
        prefs['usage_stats'] = {"claude": 0, "gemini": 0, "gpt": 0}
    
    model_type = "gemini"
    if "claude" in model_used.lower(): model_type = "claude"
    elif "gpt" in model_used.lower(): model_type = "gpt"
    
    prefs['usage_stats'][model_type] = prefs['usage_stats'].get(model_type, 0) + 1
    
    # 2. En çok kullanılan modeli tespit et (Sadece istatistik amaçlı, tercihi otomatik ezmiyoruz)
    max_usage = 0
    best_model = "gemini"
    for m, count in prefs['usage_stats'].items():
        if count > max_usage:
            max_usage = count
            best_model = m
            
    # NOT: preferred_model artık otomatik güncellenmiyor, 
    # 'auto' modunda kalması akıllı yönlendirmenin çalışması için kritik.
    # Kullanıcı Manuel Ayarlar sayfasından isterse değiştirebilir.
        
    # Yanıt tarzını analiz et (kısa/uzun)
    if answer_text:
        char_count = len(answer_text)
        current_style = prefs.get('response_style', 'balanced')
        
        if char_count < 400:
            if current_style == 'balanced': prefs['response_style'] = 'concise'
            elif current_style == 'detailed': prefs['response_style'] = 'balanced'
        elif char_count > 1500:
            if current_style == 'balanced': prefs['response_style'] = 'detailed'
            elif current_style == 'concise': prefs['response_style'] = 'balanced'

    # 3. Persona Analizi (Gemini ile derin analiz)
    if user_question and GEMINI_API_KEY:
        try:
            # Sadece her 5 mesajda bir veya persona yoksa analiz yap (Token tasarrufu)
            total_usage = sum(prefs.get('usage_stats', {}).values())
            if total_usage % 5 == 0 or not prefs.get('persona') or prefs.get('persona') == "General User":
                persona_prompt = f"""
                Analyze the user's interaction style and expertise based on this question: "{user_question}"
                Return a JSON object with:
                - "persona": One word describing the user (e.g. Developer, Student, Artist, Curious, Professional)
                - "expertise": (Beginner, Intermediate, Advanced)
                - "tone": (Formal, Casual, Technical, Creative)
                - "interests": [List of 2-3 keywords]
                
                Respond ONLY with JSON.
                """
                model = genai.GenerativeModel(GEMINI_MODEL)
                response = model.generate_content(persona_prompt)
                analysis = json.loads(response.text.strip().strip('```json').strip('```'))
                
                prefs['persona'] = analysis.get('persona', prefs.get('persona', 'General User'))
                prefs['expertise'] = analysis.get('expertise', prefs.get('expertise', 'Intermediate'))
                prefs['tone_preference'] = analysis.get('tone', 'Balanced')
                
                new_interests = analysis.get('interests', [])
                current_interests = set(prefs.get('interests', []))
                current_interests.update(new_interests)
                prefs['interests'] = list(current_interests)[:5] # Max 5 ilgi alanı
        except Exception as e:
            print(f"Persona analizi hatası: {e}")

    user.preferences = json.dumps(prefs)
    db.session.commit()

def post_process_response(text: str) -> str:
    """Yapay zeka yanıtlarını temizler ve hataları düzeltir (Markdown, Parantez vs)."""
    if not text:
        return ""
    
    # 1. Eksik Markdown bloklarını kapat
    code_block_count = text.count("```")
    if code_block_count % 2 != 0:
        text += "\n```"
    
    # 2. Basit parantez eşleştirme (Eksikse kapatmaya çalış)
    pairs = {"(": ")", "[": "]", "{": "}"}
    for open_char, close_char in pairs.items():
        if text.count(open_char) > text.count(close_char):
            diff = text.count(open_char) - text.count(close_char)
            if diff <= 2: # Çok fazla hata varsa dokunma, bozabiliriz
                text += close_char * diff

    # 3. Gereksiz başlangıç/bitiş temizliği
    text = text.strip()
    
    # 4. Code Standardization
    try:
        text = CodeStandardizer.standardize(text)
    except Exception as e:
        print(f"Standardization error: {e}")

    return text


def detect_intent(question: str, code: str = "") -> str:
    """
    Detects the user's intent from their question to route to the right model.
    
    Returns one of: simple | simple_code | debug | explain | architecture | creative | general | image_generation
    """
    intent_prompt = f"""Analyze the user's question and classify it into EXACTLY ONE intent category.

Categories (pick the BEST match):
- 'simple':       Quick factual questions, short summaries, one-liner answers. 
                  Examples: "What is REST?", "Summarize this text", "What does API stand for?"
- 'simple_code':  Basic/trivial code requests. Small functions, simple loops, basic syntax.
                  Examples: "Write a loop to print 1-10", "Reverse a string", "Sort a list"
- 'debug':        Finding bugs, fixing errors, reading tracebacks or stack traces.
                  Examples: "Why does my code crash?", "Fix this error: TypeError", "Debug this"  
- 'explain':      Conceptual explanations, how-something-works, principle overviews.
                  Examples: "Explain SOLID principles", "How does JWT work?", "What is polymorphism?"
- 'architecture': Complex system design, large-scale refactoring, multi-component patterns.
                  Examples: "Design a microservice architecture", "How to structure a clean architecture project"
- 'creative':     Writing, storytelling, poetry, brainstorming, non-technical creativity.
                  Examples: "Write a story about...", "Generate a tagline for..."
- 'general':      Greetings, off-topic chat, general conversation.
                  Examples: "Hello", "How are you?", "What's the weather?"
- 'image_generation': Explicit requests to generate/draw/create an image or picture.

IMPORTANT RULES:
- If the question starts with "kısaca", "briefly", "in short", "özetle" → prefer 'simple' or 'explain' (NOT 'architecture')
- Only use 'architecture' for truly complex multi-component system design questions
- Questions about explaining concepts (even code concepts) → 'explain', NOT 'architecture'
- Simple one-off code snippets → 'simple_code', NOT 'explain'

User Question: {question}
Related Code: {code if code else "None"}

Respond with ONLY the category name (one word, no punctuation):"""

    valid_intents = ['simple', 'simple_code', 'debug', 'explain', 'architecture', 'creative', 'general', 'image_generation']

    # Strategy: Try Gemini 2.5 Flash Lite -> 2.5 Flash (both have free quota)
    try:
        if GEMINI_API_KEY:
            for m_name in ['models/gemini-2.5-flash-lite', 'models/gemini-2.5-flash']:
                try:
                    model = genai.GenerativeModel(m_name)
                    result = model.generate_content(intent_prompt)
                    intent = getattr(result, "text", "general").strip().lower().replace("'", "").replace('"', '').replace('.', '')
                    if intent in valid_intents:
                        print(f"[IntentDetect] '{intent}' detected by {m_name}")
                        return intent
                except Exception as ge:
                    if "429" in str(ge) or "quota" in str(ge).lower():
                        continue # Try next model
                    break # Other error, try another provider

        # Fallback to GPT-4o-mini
        if openai_client:
            try:
                response = openai_client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": intent_prompt}],
                    max_tokens=15
                )
                intent = response.choices[0].message.content.strip().lower().replace("'", "").replace('.', '')
                if intent in valid_intents:
                    print(f"[IntentDetect] '{intent}' detected by gpt-4o-mini")
                    return intent
            except: pass

        # Fallback to Claude Haiku (lightweight)
        if claude_client:
            try:
                response = claude_client.messages.create(
                    model="claude-3-5-haiku-20241022",
                    max_tokens=15,
                    messages=[{"role": "user", "content": intent_prompt}]
                )
                intent = response.content[0].text.strip().lower().replace("'", "").replace('.', '')
                if intent in valid_intents:
                    print(f"[IntentDetect] '{intent}' detected by claude-haiku")
                    return intent
            except: pass
            
        return "general"
    except:
        return "general"





# --- SERİALİZASYON VE YARDIMCI FONKSİYONLAR ---

def _serialize_profile_image(image_value):
    if not image_value:
        return None
    image_str = str(image_value).strip()
    if image_str.startswith('http://') or image_str.startswith('https://'):
        return image_str
    return f"/api/files/{os.path.basename(image_str)}"

def serialize_history(item: History) -> dict:
    # Kullanıcı bilgisini conversation üzerinden al
    author_name = None
    author_id = None
    author_image = None
    if item.conversation and item.conversation.user:
        author_name = item.conversation.user.display_name
        author_id = item.conversation.user.id
        if item.conversation.user.profile_image:
            author_image = _serialize_profile_image(item.conversation.user.profile_image)
    
    data = {
        'id': item.id,
        'conversation_id': item.conversation_id,
        'user_question': item.user_question,
        'ai_response': item.ai_response,
        'selected_model': item.selected_model,
        'timestamp': item.timestamp.strftime('%Y-%m-%d %H:%M'),
        'summary': item.summary or "",
        'likes': item.likes or 0,
        'answer_count': item.answers.count() if hasattr(item, 'answers') else 0,
        'image_url': f"/api/files/{os.path.basename(item.image_path)}" if item.image_path else None,
        'author_name': author_name,
        'author_id': author_id,
        'author_image': author_image,
        'reasoning': item.reasoning or "",
        'routing_reason': item.routing_reason or "",
        'persona': item.persona or ""
    }

    # Eğer ai_response bir JSON string ise ve isComparison içeriyorsa
    try:
        import json
        if item.ai_response and item.ai_response.strip().startswith('{'):
            parsed = json.loads(item.ai_response)
            if parsed.get('isComparison'):
                data.update(parsed)
                # ai_response'u override et ki frontend tek kolonda json görmesin (geriye uyumluluk)
                # data['ai_response'] = parsed.get('ai_response', item.ai_response) 
                # Üstteki satır yerine parsed içindeki ai_response zaten doğru metni taşıyorsa onu kullanırız
                # Ancak frontend isComparison: true görünce zaten response1/response2 kullanacak.
    except:
        pass
        
    return data

def serialize_conversation(conv: Conversation) -> dict:
    return {
        'id': conv.id,
        'title': conv.title,
        'created_at': conv.created_at.strftime('%Y-%m-%d %H:%M'),
        'user_id': conv.user_id,
        'is_pinned': conv.is_pinned if hasattr(conv, 'is_pinned') else False,
        'is_archived': conv.is_archived if hasattr(conv, 'is_archived') else False,
        'linked_repo': conv.linked_repo,
        'repo_branch': conv.repo_branch
    }

def serialize_answer(answer: Answer) -> dict:
    return {
        'id': answer.id,
        'history_id': answer.history_id,
        'author': answer.author,
        'author_id': answer.author_id,
        'body': answer.body,
        'code_snippet': answer.code_snippet,
        'likes': answer.likes or 0,
        'image_url': f"/api/files/{os.path.basename(answer.image_path)}" if answer.image_path else None,
        'created_at': answer.created_at.strftime('%Y-%m-%d %H:%M'),
    }

def serialize_user(user: User) -> dict:
    prefs = {}
    if user.preferences:
        try:
            prefs = json.loads(user.preferences)
        except:
            pass
            
    token_wallet = get_or_create_token_balance(user)
    tokens = token_wallet.balance if token_wallet else SIGNUP_GRANT_TOKENS

    return {
        'id': user.id,
        'email': user.email,
        'display_name': user.display_name,
        'bio': prefs.get('bio', '') if isinstance(prefs, dict) else '',
        'is_admin': user.is_admin,
        'profile_image': _serialize_profile_image(user.profile_image),
        'created_at': user.created_at.strftime('%Y-%m-%d %H:%M') if user.created_at else None,
        'preferences': prefs,
        'xp': getattr(user, 'xp', 0),
        'total_xp_earned': getattr(user, 'total_xp_earned', 0),
        'coins': getattr(user, 'coins', 0),
        'level': getattr(user, 'level', 1),
        'streak_days': getattr(user, 'streak_days', 0),
        'last_active_date': user.last_active_date.isoformat() if hasattr(user.last_active_date, 'isoformat') else None,
        'longest_streak': getattr(user, 'longest_streak', 0),
        'tokens': tokens

    }


def _normalize_google_display_name(name: str, email: str) -> str:
    base_name = (name or email.split('@')[0] or 'google_user').strip()
    safe_name = re.sub(r'[^a-zA-Z0-9._-]+', '_', base_name).strip('._-') or 'google_user'
    candidate = safe_name[:120]
    suffix = 1

    while User.query.filter_by(display_name=candidate).first():
        suffix += 1
        candidate = f"{safe_name[:110]}_{suffix}"

    return candidate


def _exchange_google_credential(credential: str) -> tuple[dict, int]:
    google_client_id = os.getenv('GOOGLE_CLIENT_ID') or os.getenv('VITE_GOOGLE_CLIENT_ID')
    if not google_client_id:
        return {'error': 'GOOGLE_CLIENT_ID is not configured on the server.'}, 500

    try:
        import requests

        response = requests.get(
            'https://oauth2.googleapis.com/tokeninfo',
            params={'id_token': credential},
            timeout=10,
        )
        info = response.json()
    except Exception as exc:
        print(f"[AUTH] Google token verification error: {exc}")
        return {'error': 'Google token verification failed.'}, 400

    if not isinstance(info, dict) or info.get('error'):
        return {'error': info.get('error_description') or 'Invalid Google credential.'}, 401

    if info.get('aud') != google_client_id:
        return {'error': 'Google credential audience mismatch.'}, 401

    if info.get('email_verified') not in (True, 'true', 'True'):
        return {'error': 'Google account email is not verified.'}, 401

    email = (info.get('email') or '').strip().lower()
    if not email:
        return {'error': 'Google credential did not include an email address.'}, 401

    user = User.query.filter_by(email=email).first()
    created = False

    if not user:
        created = True
        display_name = _normalize_google_display_name(info.get('name') or '', email)
        user = User(
            email=email,
            display_name=display_name,
            password_hash=hash_password(secrets.token_urlsafe(32)),
            profile_image=info.get('picture') or None,
        )
        db.session.add(user)
        db.session.flush()
        get_or_create_token_balance(user)
    elif not user.profile_image and info.get('picture'):
        user.profile_image = info.get('picture')

    db.session.commit()

    token = create_access_token(identity=str(user.id))
    status_code = 201 if created else 200
    return {'token': token, 'user': serialize_user(user)}, status_code

def hash_password(password: str) -> str:
    return pbkdf2_sha256.hash(password)

def verify_password(password: str, hashed: str) -> bool:
    return pbkdf2_sha256.verify(password, hashed)


def _get_safe_jwt_identity():
    """
    Safely get JWT identity.
    Handles encoding errors from malformed Authorization headers.
    Returns None if any error occurs (encoding, missing token, invalid token).
    """
    try:
        # Get Authorization header safely - never call encode/decode on it
        auth_header = request.headers.get('Authorization')
        if not auth_header:
            return None
        
        # Ensure it's a string (should always be from Flask)
        auth_header = str(auth_header)
        
        # Log auth header details only when explicitly enabled for debugging.
        if _parse_bool(os.getenv('SAFE_JWT_DEBUG_HEADERS')):
            print(f"DEBUG: Safe JWT identity check, auth header (first 80): {repr(str(auth_header)[:80])}")
        
        # Try to get JWT identity through library
        # Library will handle validation
        verify_jwt_in_request(optional=True)
        return get_jwt_identity()
    except Exception as e:
        error_msg = str(e)
        # Log encoding errors distinctly
        if 'utf-8' in error_msg.lower() or 'codec' in error_msg.lower():
            print(f"DEBUG: UTF-8 ENCODING ERROR in _get_safe_jwt_identity: {error_msg}")
        return None


def get_current_user():
    """
    Get current authenticated user.
    Uses safe JWT extraction that handles encoding errors.
    Returns None if not authenticated or on any error.
    """
    try:
        identity = _get_safe_jwt_identity()
        if identity:
            user = db.session.get(User, identity)
            return user
        return None
    except Exception as e:
        error_msg = str(e)
        if 'utf-8' in error_msg.lower() or 'codec' in error_msg.lower():
            print(f"DEBUG: ENCODING ERROR in get_current_user: {error_msg}")
        return None


# ============================================================
# 💰 TOKEN EKONOMİSİ — YARDIMCI FONKSİYONLAR
# ============================================================

def get_or_create_token_balance(user: User) -> TokenBalance:
    """Kullanıcının TokenBalance kaydını döndürür. Yoksa signup_grant ile oluşturur."""
    wallet = TokenBalance.query.filter_by(user_id=user.id).first()
    if not wallet:
        wallet = TokenBalance(user_id=user.id, balance=SIGNUP_GRANT_TOKENS, total_spent=0)
        db.session.add(wallet)
        # İlk grant işlemini logla
        grant_tx = TokenTransaction(
            user_id=user.id,
            amount=SIGNUP_GRANT_TOKENS,
            type='signup_grant',
            description=f'Yeni kullanıcı hoş geldin bonusu — {SIGNUP_GRANT_TOKENS} token',
        )
        db.session.add(grant_tx)
        db.session.commit()
    # Not: Aylık otomatik grant artık yoktur. Sadece paket satın alma sırasında renewal yapılır.
    return wallet



# ============================================================
# 🔧 ADMIN & RENEWAL MANAGEMENT
# ============================================================

def check_and_apply_monthly_renewal(user_id: int) -> bool:
    """Eğer renewal zamanı geldiyse, satın alınan token'ları yenile."""
    purchase = (
        TokenPurchase.query
        .filter_by(user_id=user_id, status='completed', auto_renew=True)
        .order_by(TokenPurchase.completed_at.desc())
        .first()
    )
    if not purchase or not purchase.renewal_day:
        return False
    
    wallet = TokenBalance.query.filter_by(user_id=user_id).first()
    if not wallet:
        return False
    
    from datetime import datetime, timezone
    now = datetime.now(timezone.utc).replace(tzinfo=None)
    
    # Renewal zamanı kontrol et
    if purchase.last_renewal_at is None:
        # İlk kez, oluşturulma tarihine bak
        purchase.last_renewal_at = purchase.completed_at
    
    last_renewal = purchase.last_renewal_at
    next_renewal_date = last_renewal.replace(
        month=(last_renewal.month % 12) + 1 if last_renewal.month < 12 else 1,
        year=last_renewal.year if last_renewal.month < 12 else last_renewal.year + 1,
        day=min(purchase.renewal_day, 28)  # 29-31 taarihler için 28 kullan
    )
    
    if now >= next_renewal_date:
        # Yenile
        wallet.balance += purchase.tokens_granted
        purchase.last_renewal_at = now
        
        tx = TokenTransaction(
            user_id=user_id,
            amount=purchase.tokens_granted,
            type='monthly_renewal',
            description=f'{purchase.package_name} paketi otomatik yenileme — {purchase.tokens_granted} token',
            reference_id=str(purchase.id)
        )
        db.session.add(tx)
        db.session.commit()
        return True
    
    return False


def ensure_default_token_packages() -> None:
    """Seed a small set of active packages the first time billing is enabled."""
    existing_names = {pkg.name for pkg in TokenPackage.query.all()}
    created_any = False

    for package_data in DEFAULT_TOKEN_PACKAGES:
        if package_data['name'] in existing_names:
            continue
        db.session.add(TokenPackage(**package_data, is_active=True))
        created_any = True

    if created_any:
        db.session.commit()


def _serialize_token_package(package: TokenPackage) -> dict:
    return {
        'id': package.id,
        'name': package.name,
        'description': package.description,
        'tokens': package.tokens,
        'price_usd': float(package.price_usd),
        'price_cents': int(round((package.price_usd or 0) * 100)),
        'bonus_pct': package.bonus_pct or 0,
        'is_active': package.is_active,
        'stripe_price_id': package.stripe_price_id,
        'created_at': package.created_at.isoformat() if package.created_at else None,
    }


# ============================================================
# 📊 HAFTALIK/GÜNLÜK KOTA SİSTEMİ
# ============================================================

def calculate_weekly_reset_time():
    """Pazar akşamı (UTC+0) 3:00 AM'deki reset zamanını döndür."""
    from datetime import timedelta, datetime, timezone
    now = datetime.now(timezone.utc)
    days_since_sunday = (now.weekday() + 1) % 7  # Pazara göre
    if days_since_sunday == 0:  # Eğer bugün Pazar
        next_sunday = now + timedelta(weeks=1)
    else:
        next_sunday = now + timedelta(days=7 - days_since_sunday)
    reset_time = next_sunday.replace(hour=3, minute=0, second=0, microsecond=0)
    return reset_time


def calculate_daily_reset_time():
    """Yarın saat 00:00 UTC'deki reset zamanını döndür."""
    from datetime import timedelta, datetime, timezone
    now = datetime.now(timezone.utc)
    tomorrow = now + timedelta(days=1)
    reset_time = tomorrow.replace(hour=0, minute=0, second=0, microsecond=0)
    return reset_time


def reset_quota_if_needed(wallet: TokenBalance) -> bool:
    """Eğer reset zamanı geçtiyse, kotaları sıfırla. True döndür = reset yapıldı."""
    from datetime import datetime, timezone
    now = datetime.now(timezone.utc).replace(tzinfo=None)
    
    reset_needed = False
    
    # Haftalık reset kontrolü
    if wallet.weekly_reset_at and now >= wallet.weekly_reset_at:
        wallet.weekly_used = 0
        wallet.weekly_reset_at = calculate_weekly_reset_time()
        reset_needed = True
    
    # Günlük reset kontrolü
    if wallet.daily_reset_at and now >= wallet.daily_reset_at:
        wallet.daily_used = 0
        wallet.daily_reset_at = calculate_daily_reset_time()
        reset_needed = True
    
    if reset_needed:
        db.session.commit()
    
    return reset_needed


def init_quota_for_new_user(wallet: TokenBalance) -> None:
    """Yeni kullanıcı için kota zamanlarını başlat."""
    if not wallet.weekly_reset_at:
        wallet.weekly_reset_at = calculate_weekly_reset_time()
    if not wallet.daily_reset_at:
        wallet.daily_reset_at = calculate_daily_reset_time()
    db.session.commit()


def check_quota_available(wallet: TokenBalance, tokens_needed: int = 1) -> tuple[bool, str]:
    """
    Kota kontrolü yapıp, bool ve mesaj döndür.
    Returns: (allowed: bool, reason: str)
    """
    reset_quota_if_needed(wallet)
    
    # Günlük limit kontrolü
    if wallet.daily_used + tokens_needed > wallet.daily_limit:
        remaining = max(0, wallet.daily_limit - wallet.daily_used)
        return False, f"Günlük limit aşıldı. Kalan: {remaining} token. Reset: {wallet.daily_reset_at.isoformat() if wallet.daily_reset_at else 'N/A'}"
    
    # Haftalık limit kontrolü
    if wallet.weekly_used + tokens_needed > wallet.weekly_limit:
        remaining = max(0, wallet.weekly_limit - wallet.weekly_used)
        return False, f"Haftalık limit aşıldı. Kalan: {remaining} token. Reset: {wallet.weekly_reset_at.isoformat() if wallet.weekly_reset_at else 'N/A'}"
    
    # Cüzdan bakiyesi kontrolü
    if wallet.balance < tokens_needed:
        return False, f"Yetersiz token bakiyesi. Kalan: {wallet.balance} token"
    
    return True, "OK"


def deduct_tokens_and_update_quota(wallet: TokenBalance, amount: int) -> None:
    """Token harca ve kota bilgilerini güncelle."""
    reset_quota_if_needed(wallet)
    wallet.balance -= amount
    wallet.total_spent += amount
    wallet.weekly_used += amount
    wallet.daily_used += amount
    db.session.commit()


def get_quota_status(user_id: int) -> dict:
    """Kullanıcının kota durumunu JSON olarak döndür."""
    wallet = TokenBalance.query.filter_by(user_id=user_id).first()
    if not wallet:
        return None
    
    reset_quota_if_needed(wallet)
    
    weekly_used_pct = round((wallet.weekly_used / wallet.weekly_limit) * 100) if wallet.weekly_limit > 0 else 0
    daily_used_pct = round((wallet.daily_used / wallet.daily_limit) * 100) if wallet.daily_limit > 0 else 0
    
    return {
        'balance': wallet.balance,
        'total_spent': wallet.total_spent,
        'weekly': {
            'limit': wallet.weekly_limit,
            'used': wallet.weekly_used,
            'remaining': wallet.weekly_limit - wallet.weekly_used,
            'used_pct': weekly_used_pct,
            'reset_at': wallet.weekly_reset_at.isoformat() if wallet.weekly_reset_at else None,
        },
        'daily': {
            'limit': wallet.daily_limit,
            'used': wallet.daily_used,
            'remaining': wallet.daily_limit - wallet.daily_used,
            'used_pct': daily_used_pct,
            'reset_at': wallet.daily_reset_at.isoformat() if wallet.daily_reset_at else None,
        },
    }


def _payload_value(payload_obj, key, default=None):
    if payload_obj is None:
        return default
    if isinstance(payload_obj, dict):
        return payload_obj.get(key, default)
    return getattr(payload_obj, key, default)


def _get_stripe_client():
    """Return an initialized Stripe client or a (None, reason) tuple when unavailable."""
    stripe_secret_key = os.getenv('STRIPE_SECRET_KEY')
    if not stripe_secret_key:
        return None, 'Stripe is not configured. Set STRIPE_SECRET_KEY.'

    try:
        import stripe as stripe_lib
    except ImportError:
        return None, 'Stripe package is not installed.'

    stripe_lib.api_key = stripe_secret_key
    return stripe_lib, None


def _get_checkout_base_url() -> str:
    return (os.getenv('FRONTEND_URL') or os.getenv('APP_FRONTEND_URL') or request.host_url.rstrip('/')).rstrip('/')


def _checkout_success_url() -> str:
    custom_url = os.getenv('STRIPE_SUCCESS_URL')
    if custom_url:
        return custom_url
    return f"{_get_checkout_base_url()}/?billing=success&session_id={{CHECKOUT_SESSION_ID}}"


def _checkout_cancel_url() -> str:
    custom_url = os.getenv('STRIPE_CANCEL_URL')
    if custom_url:
        return custom_url
    return f"{_get_checkout_base_url()}/?billing=cancelled"


def _build_checkout_purchase(user: User, package: TokenPackage, session_obj, payment_intent_id: str | None = None):
    """Credit tokens once a Stripe checkout has completed."""
    session_id = _payload_value(session_obj, 'id')
    if not session_id:
        raise ValueError('Stripe checkout session id is missing.')

    existing_purchase = TokenPurchase.query.filter_by(stripe_checkout_session_id=session_id).first()
    if existing_purchase and existing_purchase.status == 'completed':
        wallet = get_or_create_token_balance(user)
        return existing_purchase, wallet.balance, False

    metadata = _payload_value(session_obj, 'metadata', None) or {}
    amount_cents = int(_payload_value(session_obj, 'amount_total', 0) or round((package.price_usd or 0) * 100))
    tokens_granted = package.tokens
    purchase = existing_purchase or TokenPurchase(
        user_id=user.id,
        package_id=package.id,
        package_name=package.name,
        tokens_granted=tokens_granted,
        amount_cents=amount_cents,
        currency=(_payload_value(session_obj, 'currency', None) or os.getenv('STRIPE_CURRENCY', 'usd')).lower(),
        stripe_checkout_session_id=session_id,
        stripe_payment_intent_id=payment_intent_id or _payload_value(session_obj, 'payment_intent', None),
        stripe_customer_id=_payload_value(session_obj, 'customer', None),
        status='pending',
        metadata_json=json.dumps(metadata),
    )

    purchase.user_id = user.id
    purchase.package_id = package.id
    purchase.package_name = package.name
    purchase.tokens_granted = tokens_granted
    purchase.amount_cents = amount_cents
    purchase.currency = (_payload_value(session_obj, 'currency', None) or os.getenv('STRIPE_CURRENCY', 'usd')).lower()
    purchase.stripe_checkout_session_id = session_id
    purchase.stripe_payment_intent_id = payment_intent_id or _payload_value(session_obj, 'payment_intent', None)
    purchase.stripe_customer_id = _payload_value(session_obj, 'customer', None)
    purchase.metadata_json = json.dumps(metadata)

    wallet = get_or_create_token_balance(user)
    wallet.balance += tokens_granted

    purchase.status = 'completed'
    purchase.completed_at = _utcnow()
    db.session.add(purchase)
    db.session.add(TokenTransaction(
        user_id=user.id,
        amount=tokens_granted,
        type='purchase',
        description=f'{package.name} token paketi satın alındı',
        reference_id=session_id,
    ))

    try:
        db.session.commit()
    except IntegrityError:
        db.session.rollback()
        existing_purchase = TokenPurchase.query.filter_by(stripe_checkout_session_id=session_id).first()
        wallet = get_or_create_token_balance(user)
        return existing_purchase, wallet.balance, False

    return purchase, wallet.balance, True


def _handle_checkout_session_completed(session_obj):
    metadata = _payload_value(session_obj, 'metadata', None) or {}
    user_id = metadata.get('user_id')
    package_id = metadata.get('package_id')

    if not user_id or not package_id:
        raise ValueError('Checkout session metadata is incomplete.')

    user = db.session.get(User, int(user_id))
    package = db.session.get(TokenPackage, int(package_id))
    if not user:
        raise ValueError('User not found for checkout session.')
    if not package:
        raise ValueError('Token package not found for checkout session.')

    return _build_checkout_purchase(user, package, session_obj)


with app.app_context():
    try:
        ensure_default_token_packages()
    except Exception as e:
        print(f"Warning: default token packages could not be seeded: {e}")


def _resolve_token_cost(model_name: str) -> int:
    """Model adını TOKEN_COSTS'a göre eşleştirir, en spesifik eşleşmeyi döndürür."""
    if not model_name:
        return TOKEN_COSTS['default']
    model_lower = model_name.lower().replace('models/', '')
    # Önce tam eşleşme
    if model_lower in TOKEN_COSTS:
        return TOKEN_COSTS[model_lower]
    # Sonra prefix eşleşme (en uzun önce)
    sorted_keys = sorted(TOKEN_COSTS.keys(), key=len, reverse=True)
    for key in sorted_keys:
        if key in model_lower:
            return TOKEN_COSTS[key]
    return TOKEN_COSTS['default']


def check_tokens(user: User, model_name: str = 'default') -> tuple[bool, int, int]:
    """Kullanıcının belirtilen model için yeterli token'ı var mı kontrol eder.
    
    Returns:
        (yeterli_mi: bool, mevcut_bakiye: int, gerekli_token: int)
    """
    cost = _resolve_token_cost(model_name)
    wallet = get_or_create_token_balance(user)
    return wallet.balance >= cost, wallet.balance, cost


def deduct_tokens(user: User, model_name: str = "default", description: str = None, reference_id: str = None) -> tuple[bool, int]:
    """Kullanıcının cüzdanından token düşer ve işlemi loglar.
    
    Returns:
        (başarılı_mı: bool, yeni_bakiye: int)
    """
    cost = _resolve_token_cost(model_name)
    wallet = get_or_create_token_balance(user)

    print(f"[TOKEN] Deduction attempt for user {user.id} ({user.email}). Current: {wallet.balance}, Cost: {cost}, Model: {model_name}")

    # Note: We allow balance to go negative if authorized at the start of request, 
    # to ensure the user is correctly penalized and blocked on the next turn.
    if wallet.balance < cost:
        print(f"[TOKEN] Deduction notice: Insufficient balance for user {user.id}. Required: {cost}, Found: {wallet.balance}. Allowing debt to prevent bypass loop.")

    wallet.balance -= cost
    wallet.total_spent += cost

    tx = TokenTransaction(
        user_id=user.id,
        amount=-cost,      # negatif = harcama
        type="usage",
        description=description or f"AI sorgu — {model_name}",
        reference_id=str(reference_id) if reference_id else None,
    )
    db.session.add(tx)
    try:
        db.session.commit()
        print(f"[TOKEN] SUCCESS: User {user.email} -{cost} tokens. Final Balance: {wallet.balance}")
        print(f"[TOKEN] Deduction success for user {user.id}. New balance: {wallet.balance}")
    except Exception as e:
        db.session.rollback()
        print(f"[TOKEN] DATABASE ERROR during deduction for user {user.id}: {e}")
        return False, wallet.balance + cost # Balance might be inconsistent here, but rollback should help

    return True, wallet.balance



# --- TOKEN & BILLING ROUTES DEPRECATED (Moved to end of file) ---


# --- API ROTALARI ---

@app.route('/api/auth/register', methods=['POST'])
def register():
    data = request.json or {}
    email = (data.get('email') or '').strip().lower()
    password = data.get('password') or ''
    display_name = (data.get('display_name') or '').strip()

    if not email or not password or not display_name:
        return jsonify({'error': 'Email, password and display name are required.'}), 400

    if User.query.filter_by(email=email).first():
        return jsonify({'error': 'This email is already registered.'}), 409

    if User.query.filter_by(display_name=display_name).first():
        return jsonify({'error': 'This username is already taken.'}), 409

    user = User(
        email=email,
        display_name=display_name,
        password_hash=hash_password(password)
    )
    db.session.add(user)
    db.session.commit()

    # Otomatik 100 token yüklemesini yap
    get_or_create_token_balance(user)

    token = create_access_token(identity=str(user.id))
    return jsonify({'token': token, 'user': serialize_user(user)}), 201


@app.route('/api/auth/google', methods=['POST'])
def google_login():
    data = request.json or {}
    credential = (data.get('credential') or '').strip()

    if not credential:
        return jsonify({'error': 'Google credential is required.'}), 400

    payload, status_code = _exchange_google_credential(credential)
    return jsonify(payload), status_code


@app.route('/api/auth/login', methods=['POST'])
def login():
    data = request.json or {}
    email = (data.get('email') or '').strip().lower()
    password = data.get('password') or ''

    user = User.query.filter_by(email=email).first()
    if not user or not verify_password(password, user.password_hash):
        return jsonify({'error': 'Incorrect email or password.'}), 401

    token = create_access_token(identity=str(user.id))
    return jsonify({'token': token, 'user': serialize_user(user)})


@app.route('/api/auth/me', methods=['GET'])
@jwt_required()
def me():
    user = get_current_user()
    return jsonify({'user': serialize_user(user)})


@app.route('/api/user/preferences', methods=['GET'])
@jwt_required()
def get_preferences_api():
    user = get_current_user()
    if not user:
        return jsonify({'error': 'User not found'}), 404
    prefs = get_user_preferences(user)
    return jsonify({'preferences': prefs})


@app.route('/api/user/preferences', methods=['PUT'])
@jwt_required()
def update_preferences_api():
    user = get_current_user()
    if not user:
        return jsonify({'error': 'User not found'}), 404
    
    data = request.json or {}
    current_prefs = get_user_preferences(user)
    
    # Allow manual override for specific fields
    if 'preferred_model' in data:
        current_prefs['preferred_model'] = data['preferred_model']
    if 'response_style' in data:
        current_prefs['response_style'] = data['response_style']
    if 'persona' in data:
        current_prefs['persona'] = data['persona']
    if 'expertise' in data:
        current_prefs['expertise'] = data['expertise']
    if 'interests' in data:
        # Expected list of strings
        current_prefs['interests'] = data['interests'] if isinstance(data['interests'], list) else current_prefs.get('interests', [])
        
    user.preferences = json.dumps(current_prefs)
    db.session.commit()
    
    return jsonify({'message': 'Preferences updated successfully', 'preferences': current_prefs})


@app.route('/api/auth/profile', methods=['PUT'])
@jwt_required()
def update_profile():
    """Kullanıcı profilini güncelle (display_name, bio ve şifre)."""
    user = get_current_user()
    data = request.json or {}
    
    new_display_name = (data.get('display_name') or '').strip()
    has_bio_field = 'bio' in data
    new_bio = (data.get('bio') or '').strip()
    new_password = data.get('new_password') or ''
    current_password = data.get('current_password') or ''
    
    # Display name güncelleme
    if new_display_name and new_display_name != user.display_name:
        # Benzersizlik kontrolü
        existing = User.query.filter_by(display_name=new_display_name).first()
        if existing and existing.id != user.id:
            return jsonify({'error': 'This username is already taken.'}), 409
        user.display_name = new_display_name

    # Bio güncelleme (preferences JSON içinde saklanır)
    try:
        prefs = json.loads(user.preferences) if user.preferences else {}
        if not isinstance(prefs, dict):
            prefs = {}
    except Exception:
        prefs = {}
    if has_bio_field:
        prefs['bio'] = new_bio[:500]
        user.preferences = json.dumps(prefs)
    
    # Şifre güncelleme
    if new_password:
        if not current_password:
            return jsonify({'error': 'Enter your current password.'}), 400
        if not verify_password(current_password, user.password_hash):
            return jsonify({'error': 'Incorrect current password.'}), 401
        user.password_hash = hash_password(new_password)
    
    db.session.commit()
    return jsonify({'user': serialize_user(user), 'message': 'Profile updated.'})


@app.route('/api/auth/profile/image', methods=['POST', 'DELETE'])
@jwt_required()
def upload_profile_image():
    """Profil fotoğrafı yükle veya kaldır."""
    user = get_current_user()

    if request.method == 'DELETE':
        if user.profile_image and os.path.exists(user.profile_image):
            try:
                os.remove(user.profile_image)
            except:
                pass

        user.profile_image = None
        db.session.commit()

        return jsonify({
            'user': serialize_user(user),
            'message': 'Profile picture removed.'
        })
    
    if 'image' not in request.files:
        return jsonify({'error': 'Image file required.'}), 400
    
    image_file = request.files['image']
    if not image_file.filename:
        return jsonify({'error': 'No file selected.'}), 400
    
    # Dosya uzantısı kontrolü
    allowed_extensions = {'png', 'jpg', 'jpeg', 'gif', 'webp', 'jfif'}
    file_ext = image_file.filename.rsplit('.', 1)[-1].lower() if '.' in image_file.filename else ''
    if file_ext not in allowed_extensions:
        return jsonify({'error': 'Invalid file type. Upload PNG, JPG, JPEG, GIF, WEBP or JFIF.'}), 400
    
    # Dosyayı kaydet
    filename = secure_filename(f"profile_{user.id}_{int(time.time())}.{file_ext}")
    image_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    image_file.save(image_path)
    
    # Eski profil fotoğrafını sil (varsa)
    if user.profile_image and os.path.exists(user.profile_image):
        try:
            os.remove(user.profile_image)
        except:
            pass
    
    # Kullanıcı kaydını güncelle
    user.profile_image = image_path
    db.session.commit()
    
    return jsonify({
        'user': serialize_user(user),
        'message': 'Profile picture updated.'
    })


@app.route('/api/auth/profile/analyze', methods=['POST'])
@jwt_required()
def analyze_profile():
    """Analyzes user history to generate an adaptive AI profile."""
    user = get_current_user()
    
    # 1. Fetch recent history (Last 20 interactions)
    recent_history = History.query.filter_by(conversation_id=None).first() # Fallback logic check
    # Actually we need all history for this user, possibly across conversations
    # Join with Conversation to filter by user_id
    
    history_items = db.session.query(History).join(Conversation).filter(
        Conversation.user_id == user.id
    ).order_by(History.timestamp.desc()).limit(20).all()
    
    if not history_items:
        return jsonify({'message': 'Not enough history to analyze. Chat more!'}), 200
        
    conversation_text = ""
    for h in reversed(history_items): # Chronological order
        conversation_text += f"User: {h.user_question}\n"
        if h.code_snippet:
            conversation_text += f"User Code: {h.code_snippet}\n"
            
    # 2. Construct Analysis Prompt
    prompt = f"""
    Analyze the following user's conversation history with a coding assistant.
    Determine the following profile attributes based on their questions and code:
    
    1. 'expertise': "Beginner", "Intermediate", or "Advanced".
    2. 'interests': A list of top 3 technical topics they are interested in (e.g., "Python", "React", "Algorithms").
    3. 'persona': A short title for this user (e.g., "Frontend Learner", "Data Scientist", "System Architect").
    4. 'response_style': "concise" (if they ask for quick fixes) or "detailed" (if they ask for explanations).
    
    Output ONLY valid JSON in this format:
    {{
        "expertise": "...",
        "interests": ["...", "..."],
        "persona": "...",
        "response_style": "..."
    }}
    
    User History:
    {conversation_text}
    """
    
    # 3. Call Gemini for Analysis
    try:
        # Strategy: Gemini 2.5 Flash Lite (10 RPM) -> Gemini 2.5 Flash (5 RPM)
        model_candidates = [
            'models/gemini-2.5-flash-lite',
            'gemini-2.5-flash-lite',
            'models/gemini-2.5-flash',
            'gemini-2.5-flash',
        ]
        
        response = None
        last_error = None
        
        for m_name in model_candidates:
            try:
                print(f"Analyzing profile with model: {m_name}")
                model = genai.GenerativeModel(m_name)
                response = model.generate_content(prompt)
                if response:
                    break
            except Exception as e:
                print(f"Model {m_name} failed: {e}")
                last_error = e
                # Prepare for next candidate
                if "429" in str(e) or "quota" in str(e).lower():
                    time.sleep(1) # Backoff for quota errors
                continue
        
        # If Gemini fails, try Claude as final resort
        text = ""
        if not response and claude_client:
            try:
                # User requested 4.5 specifically
                target_claude = ANTHROPIC_MODEL if ANTHROPIC_MODEL else "claude-sonnet-4-5-20250929"
                print(f"Gemini models failed, trying Claude ({target_claude}) as fallback...")
                
                cl_msg = claude_client.messages.create(
                    model=target_claude,
                    max_tokens=1000,
                    system="You are an expert user profiler. Respond ONLY with valid JSON.",
                    messages=[{"role": "user", "content": prompt}]
                )
                text = cl_msg.content[0].text
                print(f"Claude ({target_claude}) analysis successful.")
            except Exception as ce:
                print(f"Claude fallback failed: {ce}")
                # Keep the last gemini error as the main one unless this fail is more specific
                if not last_error: last_error = ce

        if not response and not text:
            raise last_error or Exception("All models failed (Gemini chain + Claude Opus)")

        if response:
            text = response.text
        
        # Improve JSON extraction
        if "```json" in text:
            text = text.split("```json")[1].split("```")[0]
        elif "```" in text:
            text = text.split("```")[1].split("```")[0]
            
        profile_data = json.loads(text.strip())
        
        # 4. Save to User Preferences
        user.preferences = json.dumps(profile_data)
        db.session.commit()
        
        return jsonify({
            'user': serialize_user(user),
            'message': 'Profile analyzed! (Updated based on history)'
        })
        
    except Exception as e:
        print(f"Profile analysis failed: {e}")
        error_msg = str(e)
        if "429" in error_msg or "quota" in error_msg.lower():
            return jsonify({'error': 'AI is busy (Rate Limit). Please try again in 1 minute.'}), 429
        return jsonify({'error': 'Failed to analyze profile.', 'details': str(e)}), 500


@app.route('/api/auth/delete-account', methods=['DELETE'])
@jwt_required()
def delete_account():
    """Kullanıcı hesabını ve ilgili tüm verileri siler."""
    import sys
    sys.stderr.write('>>> DELETE_ACCOUNT: İstek alındı.\n')
    try:
        user = get_current_user()
        if not user:
            sys.stderr.write('>>> DELETE_ACCOUNT: Kullanıcı bulunamadı.\n')
            return jsonify({'error': 'User not found.'}), 404

        data = request.json or {}
        password = (data.get('password') or '').strip()
        
        # Şifre sağlandıysa kontrol et (Normal user'lar için)
        # Şifre sağlanmadıysa skip et (Google OAuth user'lar için)
        if password:
            if not verify_password(password, user.password_hash):
                return jsonify({'error': 'Incorrect password.'}), 401

        # 1. Genel tablolar
        UserFollow.query.filter(db.or_(UserFollow.follower_id == user.id, UserFollow.following_id == user.id)).delete(synchronize_session=False)
        Snippet.query.filter_by(user_id=user.id).delete(synchronize_session=False)
        PasswordResetToken.query.filter_by(user_id=user.id).delete(synchronize_session=False)
        ApiKey.query.filter_by(user_id=user.id).delete(synchronize_session=False)
        VSCodeLoginState.query.filter_by(user_id=user.id).delete(synchronize_session=False)
        VSCodeOTP.query.filter_by(user_id=user.id).delete(synchronize_session=False)
        XPEvent.query.filter_by(user_id=user.id).delete(synchronize_session=False)
        UserBadge.query.filter_by(user_id=user.id).delete(synchronize_session=False)
        UserTheme.query.filter_by(user_id=user.id).delete(synchronize_session=False)
        MemoryEdge.query.filter_by(user_id=user.id).delete(synchronize_session=False)
        MemoryNode.query.filter_by(user_id=user.id).delete(synchronize_session=False)
        MemoryItem.query.filter_by(user_id=user.id).delete(synchronize_session=False)

        # 2. Sohbetler
        conversations = Conversation.query.filter_by(user_id=user.id).all()
        conv_ids = [c.id for c in conversations]
        history_ids = []
        if conv_ids:
            histories = History.query.filter(History.conversation_id.in_(conv_ids)).all()
            history_ids = [h.id for h in histories]

        if history_ids:
            Notification.query.filter(Notification.related_post_id.in_(history_ids)).delete(synchronize_session=False)
            Favorite.query.filter(Favorite.history_id.in_(history_ids)).delete(synchronize_session=False)
            Feedback.query.filter(Feedback.history_id.in_(history_ids)).delete(synchronize_session=False)
            FeedbackDetail.query.filter(FeedbackDetail.history_id.in_(history_ids)).delete(synchronize_session=False)
            ConversationSummary.query.filter(ConversationSummary.last_history_id.in_(history_ids)).delete(synchronize_session=False)
            PostLike.query.filter(PostLike.history_id.in_(history_ids)).delete(synchronize_session=False)
            
            answers = Answer.query.filter(Answer.history_id.in_(history_ids)).all()
            ans_ids = [a.id for a in answers]
            if ans_ids:
                AnswerLike.query.filter(AnswerLike.answer_id.in_(ans_ids)).delete(synchronize_session=False)
            Answer.query.filter(Answer.history_id.in_(history_ids)).delete(synchronize_session=False)

        Notification.query.filter(db.or_(Notification.user_id == user.id, Notification.related_user_id == user.id)).delete(synchronize_session=False)
        NotificationRead.query.filter_by(user_id=user.id).delete(synchronize_session=False)
        NotificationHidden.query.filter_by(user_id=user.id).delete(synchronize_session=False)
        Favorite.query.filter_by(user_id=user.id).delete(synchronize_session=False)
        Feedback.query.filter_by(user_id=user.id).delete(synchronize_session=False)
        FeedbackDetail.query.filter_by(user_id=user.id).delete(synchronize_session=False)
        PostLike.query.filter_by(user_id=user.id).delete(synchronize_session=False)
        AnswerLike.query.filter_by(user_id=user.id).delete(synchronize_session=False)
        Answer.query.filter_by(author_id=user.id).delete(synchronize_session=False)

        if conv_ids:
            ConversationSummary.query.filter(ConversationSummary.conversation_id.in_(conv_ids)).delete(synchronize_session=False)
        if history_ids:
            History.query.filter(History.id.in_(history_ids)).delete(synchronize_session=False)

        if conv_ids:
            # SharedSession'lara bağlı yorum ve incelemeleri sil (Foreign Key hatasını önlemek için)
            shared_sessions = SharedSession.query.filter(SharedSession.conversation_id.in_(conv_ids)).all()
            session_ids = [ss.id for ss in shared_sessions]
            if session_ids:
                CollaborationComment.query.filter(CollaborationComment.session_id.in_(session_ids)).delete(synchronize_session=False)
                CollaborationReview.query.filter(CollaborationReview.session_id.in_(session_ids)).delete(synchronize_session=False)
            
            SharedSession.query.filter(SharedSession.conversation_id.in_(conv_ids)).delete(synchronize_session=False)
            Conversation.query.filter(Conversation.id.in_(conv_ids)).delete(synchronize_session=False)

        # Kullanıcının sahip olduğu diğer paylaşımlar için de aynı temizliği yap
        owned_shared_sessions = SharedSession.query.filter_by(owner_id=user.id).all()
        owned_session_ids = [ss.id for ss in owned_shared_sessions]
        if owned_session_ids:
            CollaborationComment.query.filter(CollaborationComment.session_id.in_(owned_session_ids)).delete(synchronize_session=False)
            CollaborationReview.query.filter(CollaborationReview.session_id.in_(owned_session_ids)).delete(synchronize_session=False)
            SharedSession.query.filter(SharedSession.id.in_(owned_session_ids)).delete(synchronize_session=False)

        CollaborationReview.query.filter_by(updated_by_user_id=user.id).delete(synchronize_session=False)
        CollaborationComment.query.filter_by(author_user_id=user.id).delete(synchronize_session=False)
        
        projects = Project.query.filter_by(user_id=user.id).all()
        for p in projects:
            db.session.delete(p)
            
        TokenBalance.query.filter_by(user_id=user.id).delete(synchronize_session=False)
        TokenTransaction.query.filter_by(user_id=user.id).delete(synchronize_session=False)
        TokenPurchase.query.filter_by(user_id=user.id).delete(synchronize_session=False)

        if user.profile_image and os.path.exists(user.profile_image):
            try: os.remove(user.profile_image)
            except: pass

        db.session.delete(user)
        db.session.commit()
        return jsonify({'message': 'Account deleted.'})
    except Exception as e:
        db.session.rollback()
        sys.stderr.write(f'ERROR: {str(e)}\n')
        return jsonify({'error': str(e)}), 500

# --- ŞİFRE SIFIRLAMA ---

def send_reset_email(to_email, reset_code):
    """Resend API veya SMTP ile şifre sıfırlama kodu gönderir.
    
    Production: RESEND_API_KEY kullanılır (önerilen)
    Development: SMTP ayarları kullanılabilir (fallback)
    """
    resend_api_key = os.getenv('RESEND_API_KEY')
    mail_from = os.getenv('MAIL_FROM', 'CodeAlchemist <onboarding@resend.dev>')
    
    html_content = f"""
    <html>
    <body style="font-family: Arial, sans-serif; background-color: #1a1a2e; color: #eee; padding: 20px;">
        <div style="max-width: 500px; margin: 0 auto; background: #16213e; border-radius: 12px; padding: 30px;">
            <h2 style="color: #a855f7; margin-bottom: 20px;">🔐 Password Reset</h2>
            <p>Hello,</p>
            <p>Password reset request received.</p>
            <div style="background: #0f0f23; border: 2px solid #a855f7; border-radius: 8px; padding: 20px; text-align: center; margin: 20px 0;">
                <span style="font-size: 32px; font-weight: bold; color: #a855f7; letter-spacing: 8px;">{reset_code}</span>
            </div>
            <p style="color: #888; font-size: 14px;">This code will expire in 15 minutes.</p>
            <hr style="border: none; border-top: 1px solid #333; margin: 20px 0;">
            <p style="color: #666; font-size: 12px;">If you did not request this, please ignore this email.</p>
        </div>
    </body>
    </html>
    """
    
    text_content = f"""Hello,

Password reset request received.

Your Verification Code: {reset_code}

This code will expire in 15 minutes.

If you did not request this, please ignore this email.

CodeAlchemist Team"""
    
    # Öncelik 1: Resend API (Production için önerilen)
    if resend_api_key:
        try:
            resend.api_key = resend_api_key
            params = {
                "from": mail_from,
                "to": [to_email],
                "subject": "CodeAlchemist - Password Reset Code",
                "html": html_content,
                "text": text_content
            }
            email_response = resend.Emails.send(params)
            print(f"Password reset code sent (Resend): {to_email}, ID: {email_response.get('id')}")
            return True
        except Exception as e:
            print(f"Resend email sending error: {e}")
            return False
    
    # Öncelik 2: SMTP (Development/Fallback)
    mail_server = os.getenv('MAIL_SERVER', 'smtp.gmail.com')
    mail_port = int(os.getenv('MAIL_PORT', 587))
    mail_username = os.getenv('MAIL_USERNAME')
    mail_password = os.getenv('MAIL_PASSWORD')
    
    if not mail_username or not mail_password:
        # Development mode: Print code to console instead of sending email
        print("=" * 50)
        print("📧 DEVELOPMENT MODE - Email would be sent to:", to_email)
        print(f"🔐 PASSWORD RESET CODE: {reset_code}")
        print("=" * 50)
        return True  # Return success so user can use the code from console
    
    try:
        msg = MIMEMultipart('alternative')
        msg['Subject'] = 'CodeAlchemist - Password Reset Code'
        msg['From'] = mail_username
        msg['To'] = to_email
        
        part1 = MIMEText(text_content, 'plain')
        part2 = MIMEText(html_content, 'html')
        msg.attach(part1)
        msg.attach(part2)
        
        with smtplib.SMTP(mail_server, mail_port) as server:
            server.starttls()
            server.login(mail_username, mail_password)
            server.sendmail(mail_username, to_email, msg.as_string())
        
        print(f"Password reset code sent (SMTP): {to_email}")
        return True
        
    except Exception as e:
        print(f"SMTP email sending error: {e}")
        return False


@app.route('/api/auth/forgot-password', methods=['POST'])
def forgot_password():
    """Şifre sıfırlama kodu gönderir."""
    data = request.json or {}
    email = data.get('email', '').strip().lower()
    
    if not email:
        return jsonify({'error': 'Email address required.'}), 400
    
    user = User.query.filter_by(email=email).first()
    if not user:
        # Güvenlik: Kullanıcı olmasa bile başarılı mesajı göster
        return jsonify({'message': 'If this email is registered, a password reset code has been sent.'})
    
    # 6 haneli rastgele kod oluştur
    reset_code = ''.join([str(random.randint(0, 9)) for _ in range(6)])
    
    # Eski tokenları sil
    PasswordResetToken.query.filter_by(user_id=user.id, used=False).delete()
    
    # Yeni token oluştur (15 dakika geçerli)
    token = PasswordResetToken(
        user_id=user.id,
        token=reset_code,
        expires_at=_utcnow() + timedelta(minutes=15)
    )
    db.session.add(token)
    db.session.commit()
    
    # Email gönder
    if send_reset_email(email, reset_code):
        return jsonify({'message': 'Password reset code sent to your email address.'})
    else:
        return jsonify({'error': 'Failed to send email. Please try again later.'}), 500


@app.route('/api/auth/reset-password', methods=['POST'])
def reset_password():
    """Kod ile şifreyi sıfırlar."""
    data = request.json or {}
    email = data.get('email', '').strip().lower()
    code = data.get('code', '').strip()
    new_password = data.get('new_password', '')
    
    if not email or not code or not new_password:
        return jsonify({'error': 'Email, code and new password are required.'}), 400
    
    if len(new_password) < 6:
        return jsonify({'error': 'Password must be at least 6 characters.'}), 400
    
    user = User.query.filter_by(email=email).first()
    if not user:
        return jsonify({'error': 'Invalid code or email.'}), 400
    
    # Token'ı kontrol et
    token = PasswordResetToken.query.filter_by(
        user_id=user.id,
        token=code,
        used=False
    ).first()
    
    if not token:
        return jsonify({'error': 'Invalid or expired code.'}), 400
    
    if token.expires_at < _utcnow():
        return jsonify({'error': 'Code expired. Request a new code.'}), 400
    
    # Şifreyi güncelle
    user.password_hash = hash_password(new_password)
    token.used = True
    db.session.commit()
    
    return jsonify({'message': 'Your password has been successfully updated. You can log in.'})


@app.route('/api/github/link', methods=['POST'])
@jwt_required()
def link_github_repo():
    user = get_current_user()
    data = request.json or {}
    repo_name = data.get('repo', '').strip()
    branch = data.get('branch', 'main').strip()
    conversation_id = data.get('conversation_id')
    
    if not repo_name:
        return jsonify({'error': 'Repo name is required.'}), 400
        
    parser = GitHubParser()
    tree = parser.get_repo_tree(repo_name, branch)
    
    if tree is None:
        return jsonify({'error': 'Could not fetch repo tree. Check repo name and permissions.'}), 400

    # If conversation_id is provided and valid, link it in the DB
    if conversation_id and conversation_id != 'null' and conversation_id != 'undefined':
        conversation = db.session.get(Conversation, conversation_id)
        if not conversation or conversation.user_id != user.id:
            return jsonify({'error': 'Conversation not found or unauthorized.'}), 404
        
        conversation.linked_repo = repo_name
        conversation.repo_branch = branch
        db.session.commit()
    
    return jsonify({
        'message': f'Successfully verified {repo_name} ({branch}).' if not conversation_id or conversation_id == 'null' else f'Successfully linked {repo_name} ({branch}).',
        'tree_size': len(tree),
        'repo': repo_name,
        'branch': branch
    })

@app.route('/api/github/tree', methods=['GET'])
@jwt_required()
def get_github_tree():
    user = get_current_user()
    conversation_id = request.args.get('conversation_id')
    repo_param = request.args.get('repo')
    branch_param = request.args.get('branch', 'main')
    
    # 1. Option: Direct repo/branch (works without a conversation)
    if repo_param:
        parser = GitHubParser()
        tree = parser.get_repo_tree(repo_param, branch_param)
        if tree is None:
            return jsonify({'error': 'Could not fetch repo tree for this repository.'}), 400
        return jsonify({
            'repo': repo_param,
            'branch': branch_param,
            'tree': tree
        })

    # 2. Option: Conversation ID (legacy support / associated data)
    if not conversation_id or conversation_id == 'null' or conversation_id == 'undefined':
        return jsonify({'error': 'Conversation ID or Repo name is required.'}), 400
        
    conversation = db.session.get(Conversation, conversation_id)
    if not conversation or conversation.user_id != user.id:
        return jsonify({'error': 'Conversation not found or unauthorized.'}), 404
        
    if not conversation.linked_repo:
        return jsonify({'error': 'No repository linked to this conversation.'}), 400
        
    parser = GitHubParser()
    tree = parser.get_repo_tree(conversation.linked_repo, conversation.repo_branch)
    
    if tree is None:
        return jsonify({'error': 'Could not fetch repo tree.'}), 400
        
    return jsonify({
        'repo': conversation.linked_repo,
        'branch': conversation.repo_branch,
        'tree': tree
    })

@app.route('/api/github/file', methods=['GET'])
@jwt_required()
def get_github_file():
    user = get_current_user()
    repo_param = request.args.get('repo')
    branch_param = request.args.get('branch', 'main')
    path_param = request.args.get('path')
    
    if not repo_param or not path_param:
        return jsonify({'error': 'Repo and path parameters are required.'}), 400
        
    parser = GitHubParser()
    content = parser.get_file_content(repo_param, path_param, branch_param)
    
    if content is None:
        return jsonify({'error': 'Could not fetch file content from GitHub.'}), 400
        
    return jsonify({
        'repo': repo_param,
        'branch': branch_param,
        'path': path_param,
        'content': content
    })

@app.route('/api/generate_tests', methods=['POST'])
@jwt_required()
def generate_tests():
    data = request.json or {}
    code = data.get('code', '').strip()
    language = data.get('language', 'javascript').strip()
    
    if not code:
        return jsonify({'error': 'No code provided.'}), 400
        
    prompt = f"You are an expert QA Engineer. Generate robust, complete unit tests for the following {language} code snippet.\n"
    if language.lower() in ['javascript', 'typescript', 'jsx', 'tsx']:
        prompt += "Use Jest or Vitest framework. Assume standard imports.\n"
    elif language.lower() in ['python']:
        prompt += "Use Pytest framework.\n"
    elif language.lower() in ['java']:
        prompt += "Use JUnit 5 framework.\n"
    else:
        prompt += f"Use the standard/most popular testing framework for {language}.\n"
    
    prompt += "Output ONLY the raw testing code without any markdown formatting or explanation, just the code itself.\n\nCode:\n"
    prompt += code
    
    try:
        model = genai.GenerativeModel('models/gemini-2.5-flash')
        response = model.generate_content(prompt)
        test_code = response.text.replace('```' + language, '').replace('```', '').strip()
        return jsonify({'tests': test_code})
    except Exception as e:
        print(f"Test Gen Error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/github/audit_pr', methods=['POST'])
@jwt_required()
def audit_github_pr():
    user = get_current_user()
    data = request.json or {}
    file_changes = data.get('file_changes', [])
    
    if not file_changes:
        return jsonify({'error': 'No file changes provided for audit.'}), 400
        
    audit_prompt = "You are a Senior AI Security and Performance Auditor.\n"
    audit_prompt += "Review the following incoming Pull Request file changes for security vulnerabilities (e.g., SQL Injection, XSS, leaked sensitive tokens/passwords) and severe performance bottlenecks.\n\n"
    audit_prompt += "Respond ONLY with a valid JSON object in the following format:\n"
    audit_prompt += "{\n  \"passed\": true/false,\n  \"issues\": [\"Issue 1\", \"Issue 2\"] // empty array if passed\n}\n\n"
    audit_prompt += "File Changes to Review:\n"
    
    for f in file_changes:
        audit_prompt += f"\n--- File: {f.get('path')} ---\n{f.get('content')}\n"
        
    try:
        model = genai.GenerativeModel('models/gemini-2.5-flash')
        response = model.generate_content(audit_prompt)
        text = response.text.replace('```json', '').replace('```', '').strip()
        result = json.loads(text)
        return jsonify(result)
    except Exception as e:
        print(f"Audit PR error: {e}")
        # Gracefully pass if AI fails
        return jsonify({'passed': True, 'issues': [f"Audit could not complete due to AI service error: {str(e)}"]})

@app.route('/api/github/pr', methods=['POST'])
@jwt_required()
def create_github_pr():
    user = get_current_user()
    data = request.json or {}
    repo_name = data.get('repo', '').strip()
    base_branch = data.get('base_branch', 'main').strip()
    new_branch = data.get('new_branch', 'code-alchemist-fix').strip()
    title = data.get('title', 'AI Generated Refactor').strip()
    body = data.get('body', 'This PR includes changes generated by Code Alchemist.')
    file_changes = data.get('file_changes', [])
    
    if not repo_name or not file_changes:
        return jsonify({'error': 'Repo name and file changes are required.'}), 400
        
    parser = GitHubParser()
    result = parser.create_pull_request(repo_name, base_branch, new_branch, title, body, file_changes)
    
    if 'error' in result:
        return jsonify(result), 400
        
    return jsonify(result)

# Legacy endpoint kept for reference. Disabled to avoid route collision.
# @app.route('/api/github/blueprint', methods=['GET'])
def get_github_blueprint():
    repo_param = request.args.get('repo')
    branch_param = request.args.get('branch', 'main')
    
    if not repo_param:
        return jsonify({'error': 'Repo parameter is required.'}), 400
        
    parser = GitHubParser()
    tree = parser.get_repo_tree(repo_param, branch_param)
    if tree is None:
        return jsonify({'error': 'Could not fetch repo tree.'}), 400
    
    tree_str = parser.format_tree_for_prompt(tree)
    
    prompt = f"You are a Senior Software Architect. Generate a comprehensive project blueprint for the following repository structure:\n\n{tree_str}\n\n"
    prompt += "Include:\n1. A high-level summary of the project architecture.\n"
    prompt += "2. A Mermaid diagram representing the folder and file relationships.\n"
    prompt += "3. Key technology stack identifies (if possible).\n"
    prompt += "Output ONLY the markdown content."
    
    try:
        model = genai.GenerativeModel(GEMINI_MODEL)
        response = model.generate_content(prompt)
        content = response.text
        # Handle cases where AI might wrap output in markdown code blocks
        if '```' in content:
            content = content.replace('```markdown', '').replace('```', '').strip()
        return jsonify({'markdown': content})
    except Exception as e:
        print(f"Blueprint Error: {e}")
        return jsonify({'error': str(e)}), 500

# DUPLICATE ENDPOINT #2 - DEVRE DIŞI (Gerçek endpoint satır ~3550'de)
# @app.route('/api/github/health', methods=['GET'])
# def get_github_health():
#     ...eski kod kaldırıldı...

@app.route('/api/refactor/bulk', methods=['POST'])
@jwt_required()
def bulk_refactor():
    user = get_current_user()
    data = request.json or {}
    repo_name = data.get('repo', '').strip()
    branch = data.get('branch', 'main').strip()
    instructions = data.get('instructions', '').strip()
    
    if not repo_name or not instructions:
        return jsonify({'error': 'Repo name and instructions are required.'}), 400
        
    # In a full implementation, this would fetch files, send to LLM, etc.
    return jsonify({
        'message': 'Bulk refactoring initiated',
        'repo': repo_name,
        'instructions_received': instructions
    })

@app.route('/api/ask', methods=['POST'])
def ask():
    # Debug logging
    print(f"DEBUG: /api/ask called. Content-Type: {request.content_type}")

    payload = {}
    include_previous_modules = None
    payload_project_id = None
    workspace_files = []
    agent_mode = False
    allow_write_tools = False
    source_header = request.headers.get("X-Client-Source") or "web"
    repo_param = None
    branch_param = "main"
    
    # Handle multipart/form-data
    if request.content_type and 'multipart/form-data' in request.content_type:
        question = request.form.get('question', '')
        code = request.form.get('code', '')
        model = request.form.get('model', 'auto')
        conversation_id = request.form.get('conversation_id')
        include_previous_modules = request.form.get('include_previous_modules')
        payload_project_id = request.form.get('project_id')
        agent_mode = _parse_bool(request.form.get('agent_mode'))
        allow_write_tools = _parse_bool(request.form.get('allow_write_tools'))
        workspace_files = _parse_workspace_files_payload(request.form.get('workspace_files'))
        if conversation_id == 'null' or conversation_id == 'undefined':
            conversation_id = None

        payload = request.form.to_dict(flat=True)
        
        image_file = request.files.get('image')
        image_path = None
        if image_file and image_file.filename:
            filename = secure_filename(f"{int(time.time())}_{image_file.filename}")
            image_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            image_file.save(image_path)
    else:
        # Handle JSON (or attempt to parse as JSON)
        payload = request.get_json(silent=True) or {}
        question = payload.get('question', '')
        code = payload.get('code', '')
        model = payload.get('model', 'auto')
        conversation_id = payload.get('conversation_id')
        include_previous_modules = payload.get('include_previous_modules')
        payload_project_id = payload.get('project_id')
        agent_mode = _parse_bool(payload.get('agent_mode'))
        allow_write_tools = _parse_bool(payload.get('allow_write_tools'))
        workspace_files = _parse_workspace_files_payload(payload.get('workspace_files'))
        repo_param = payload.get('repo')
        branch_param = payload.get('branch', 'main')
        image_path = None

    no_save = _parse_bool(payload.get('no_save'))
    
    # Kullanıcı tespiti
    user = get_current_user()
    user_id = user.id if user else None
    print(f"DEBUG: agent_mode value received: {agent_mode} (type: {type(agent_mode)})")
    include_previous_modules = _resolve_include_previous_modules(
        user,
        {'include_previous_modules': include_previous_modules} if include_previous_modules is not None else payload,
        no_save=no_save,
    )
    
    # AI Taste Profile al
    prefs = get_user_preferences(user)

    # 💰 TOKEN EKONOMİSİ — Bakiye Kontrolü
    if user:
        allowed, balance, req_cost = check_tokens(user, model)
        if not allowed:
            return jsonify({
                'error': f'Yetersiz Token! Bu işlem için {req_cost} token gerekiyor, mevcut bakiyeniz: {balance}.',
                'insufficient_tokens': True,
                'required': req_cost,
                'balance': balance
            }), 402 # Payment Required

    # Client source (Web vs Extension)
    source_header = request.headers.get("X-Client-Source") or payload.get("source") or source_header or "web"
    if source_header == 'vscode': # Legacy compatibility
        source_header = 'extension'

    # Legacy Free/Premium quota gate (opsiyonel, paralel çalışabilir)
    try:
        estimated_tokens = _estimate_tokens_for_request(question, code)
        quota_result = consume_plan_quota(user, estimated_tokens, request_weight=1)
        if not quota_result.get('allowed', True):
            # ... existing logic ...
            pass 
    except Exception as e:
        print(f"Quota check failed (continuing): {e}")

    # Konuşma Yönetimi
    conversation = None
    history_context = []
    agent_history_context = []
    github_context = ""
    resolved_agent_project = None
    agent_project_source = None

    if not no_save:
        if conversation_id:
            conversation = db.session.get(Conversation, conversation_id)
            # Eğer conversation varsa ve kullanıcı geçmişi istiyorsa, geçmişi çek
            if conversation and include_previous_modules:
                # Sadece bu konuşmaya ait son 5 mesajı al (Token tasarrufu için)
                prev_items = History.query.filter_by(conversation_id=conversation.id)\
                    .order_by(History.timestamp.desc())\
                    .limit(5)\
                    .all()
                # Descending aldığımız için kronolojik sıraya geri çeviriyoruz
                prev_items.reverse()
                
                for item in prev_items:
                    history_context.append({'user': item.user_question, 'ai': item.ai_response})
        
        if not conversation:
            # Yeni konuşma başlat
            conversation = Conversation(user_id=user_id, title=question[:50], source=source_header)
            
            # If repo was pre-verified and passed here, link it to the new conversation
            if 'repo_param' in locals() and repo_param:
                conversation.linked_repo = repo_param
                conversation.repo_branch = branch_param
                print(f"DEBUG: Pre-linking repo {repo_param} to new conversation {conversation.id}")
                
            db.session.add(conversation)
            db.session.commit()

        resolved_agent_project, agent_project_source = _resolve_agent_project(user, conversation, payload_project_id)
        if resolved_agent_project and conversation and not conversation.project_id:
            # Persist project binding for new conversations created from project workspace requests.
            conversation.project_id = resolved_agent_project.id
            db.session.add(conversation)
            db.session.commit()

        # Check for linked GitHub repo to inject context
        if conversation.linked_repo:
            repo = conversation.linked_repo
            branch = conversation.repo_branch
            parser = GitHubParser()
            
            tree = parser.get_repo_tree(repo, branch)
            if tree:
                tree_str = parser.format_tree_for_prompt(tree)
                github_context = f"\n\n[System: This conversation is linked to GitHub repository '{repo}' (branch '{branch}').\nRepository Structure:\n{tree_str[:3000]}\n]"
                print(f"Injected GitHub context for {repo}")
                
                # Diff Formatting Instruction for UI 
                diff_instruction = "\n\n[CRITICAL SYSTEM INSTRUCTION: For EVERY code change you propose, you MUST use the Side-by-Side Diff format. DO NOT use standard `-` or `+` lines. You MUST follow this EXACT structure for EACH file change:\n\nFile: `path/to/file` (Relative to project root)\n```diff\n<<<OLD>>>\n[Insert the EXACT block of old code being replaced - must be a perfect match for the original file code]\n<<<NEW>>>\n[Insert the full new code block that replaces the OLD section]\n```\n\nFAILURE TO USE THIS EXACT <<<OLD>>> / <<<NEW>>> FORMAT WILL BREAK THE USER'S INTERFACE. DO NOT SKIP THIS.]"

                # Check for Magic Fix intent based on common error trace keywords
                error_pattern = re.compile(r'(Traceback \(most recent call last\):|Error:|Exception:|TypeError|ValueError|ReferenceError|SyntaxError|IndexError|KeyError|ModuleNotFoundError)\b', re.IGNORECASE)
                if error_pattern.search(question):
                    print("Magic Fix triggered based on error pattern.")
                    magic_fix_prompt = "\n\n[System Magic Fix Instruction: The user has provided an error trace. Analyze the error based on the linked repository context and provide a direct, root-cause solution. Output the required file changes.]"
                    github_context += magic_fix_prompt

                # Always append diff instruction if linked to repo
                github_context += diff_instruction

                # REMOVED: question = f"{question}{github_context}"
                # Instead, we pass github_context separately to the generation functions
                pass

        print(f"Model İsteği: {model}, ConvID: {conversation.id}, Image: {image_path}")

        # Inject project file context if conversation belongs to a project
        if conversation.project_id:
            proj = resolved_agent_project or db.session.get(Project, conversation.project_id)
            if proj:
                resolved_agent_project = proj
                # Prefer relevance-ranked embedding context. Fallback to static context if embeddings fail.
                project_context = build_project_context_for_question(proj, question)

                if not project_context:
                    files = proj.files.order_by(ProjectFile.name).all()
                    if files:
                        ctx_parts = [f"[System: Bu sohbet '{proj.name}' projesine aittir. Aşağıdaki proje dosyaları bağlam olarak sağlanmıştır:"]
                        if proj.description:
                            ctx_parts.append(f"Proje açıklaması: {proj.description}")
                        total_chars = 0
                        for pf in files:
                            file_content = pf.content[:3000]  # her dosya max 3000 karakter
                            total_chars += len(file_content)
                            if total_chars > 12000:
                                break
                            ctx_parts.append(f"\n## Dosya: {pf.name} ({pf.language})\n```{pf.language}\n{file_content}\n```")
                        ctx_parts.append("]")
                        project_context = '\n'.join(ctx_parts)

                if project_context:
                    github_context = project_context + "\n" + github_context
                    print(f"Injected project context for project {proj.id} (embedding-aware)")

    else:
        print(f"Model İsteği (no_save): {model}, Image: {image_path}")

    # Agent mode state should carry full conversation turns when conversation_id is provided.
    if agent_mode and conversation_id:
        if conversation is None:
            conversation = db.session.get(Conversation, conversation_id)
        if conversation:
            full_prev_items = History.query.filter_by(conversation_id=conversation.id)\
                .order_by(History.timestamp.asc())\
                .all()
            for item in full_prev_items:
                agent_history_context.append({
                    'user': item.user_question or '',
                    'ai': item.ai_response or '',
                })

    if resolved_agent_project is None and agent_mode:
        resolved_agent_project, agent_project_source = _resolve_agent_project(user, conversation, payload_project_id)

    memory_context = {'text': '', 'hit_count': 0, 'hits': []}
    if include_previous_modules and user:
        print(f"[MEMORY] Loading previous context for user {user.id}")
        memory_context = _load_previous_memory_context(user, question, conversation, include_previous_modules=True)
        if memory_context.get('text'):
            print(f"[MEMORY] Successfully retrieved {memory_context.get('hit_count', 0)} memory hits ({len(memory_context['text'])} chars)")
            github_context = f"{memory_context['text']}\n\n{github_context}".strip() if github_context else memory_context['text']
        else:
            print("[MEMORY] No relevant memory context found for this query")

    # --- Akıllı Model Routing (Smart Routing) ---
    original_model = model
    routing_reason = None
    detected_lang = 'unknown'
    detected_intent = 'general'
    
    # Ses dosyası kontrolü - Sadece Gemini ses desteği sağlıyor
    audio_extensions = ['.mp3', '.wav', '.webm', '.m4a', '.ogg', '.aac', '.flac']
    is_audio_file = image_path and any(image_path.lower().endswith(ext) for ext in audio_extensions)
    
    # Model fonksiyonlarına gönderilecek image_path (Varsayılan: orijinal path)
    model_image_path = image_path 
    
    if is_audio_file:
        # Eğer seçilen model Gemini değilse, sesi metne çevir ve öyle gönder
        if 'gemini' not in model:
            print(f"DEBUG: Audio detected for non-Gemini model ({model}). Transcribing...")
            transcription = transcribe_audio_with_gemini(image_path)
            
            if transcription:
                # Transkripti soruya ekleyelim ve modele özel talimat verelim
                instruction = (
                    f"\n\n[System: The user sent a voice message. Here is the transcription:]\n"
                    f"\"{transcription}\"\n\n"
                    f"[Instruction: The user's input was transcribed from audio. Directly provide your answer based on this transcription. DO NOT repeat or quote the transcription in your response.]"
                )
                
                if question:
                    question = f"{question}{instruction}"
                else:
                    question = instruction
                
                model_image_path = None 
                
                routing_reason = f"🎤 Ses mesajı metne çevrildi ve **{model}** modeline iletildi."
            else:
                # Transkripsiyon başarısızsa direkt döngüyü sonlandırmak yerine kibar bir hata göster
                def generate_error():
                    yield "data: {\"chunk\": \"⚠️ Sesli mesajınız işlenemedi. Lütfen tekrar deneyin veya farklı bir model seçin.\"}\n\n"
                    yield "data: {\"done\": true}\n\n"
                return Response(stream_with_context(generate_error()), mimetype='text/event-stream')
        else:
            # Gemini zaten seçili
            if model != GEMINI_MODEL.replace('models/', '') and 'flash' not in model:
                 model = GEMINI_MODEL.replace('models/', '')
                 routing_reason = f"🎤 Ses mesajları için **{model}** optimize edildi."
    
    # Görsel Oluşturma İsteği Kontrolü (Image Generation Intent)
    q_lower = question.lower()
    
    # Daha esnek Türkçe kontrolü
    creation_verbs = ['çiz', 'oluştur', 'yarat', 'yap', 'hazırla', 'generate', 'create', 'draw', 'make', 'tasarla', 'üret', 'çizsene', 'yaparmısın', 'çizer misin', 'istiyorum', 'gönder', 'yolla', 'boya', 'canlandır']
    image_nouns = ['resim', 'görsel', 'fotoğraf', 'image', 'picture', 'photo', 'drawing', 'art', 'logo', 'ikon', 'icon', 'sketch', 'tasarım', 'resmini', 'gorselini', 'fotografini', 'resmi', 'gorseli', 'fotografi', 'çizim', 'cizim', 'png', 'jpg', 'karikatür', 'illüstrasyon', 'poster', 'afiş', 'kapak', 'banner', 'avatar', 'manzara', 'portre', 'karakter']
    
    # Basit anahtar kelime öbekleri
    exact_phrases = [
        'create image', 'generate image', 'draw a picture', 'resim çiz', 'görsel oluştur', 
        'resim yap', 'görsel yarat', 'fotoğraf oluştur', 'resim istiyorum', 'görsel istiyorum',
        'çizgi film', 'logo yap', 'ikon yap', 'resmi yap', 'görseli yap', 'resmini oluştur', 'görselini çiz'
    ]
    
    # Kelime bazlı kontrol (Hem 'resim' hem 'çiz' geçiyorsa)
    has_noun = any(noun in q_lower for noun in image_nouns)
    has_verb = any(verb in q_lower for verb in creation_verbs)
    
    # Kodlama ile çizim isteği (matplotlib, turtle vs.) var mı?
    code_keywords = ['python', 'kod', 'code', 'script', 'matplotlib', 'turtle', 'grafik', 'plot', 'chart', 'pandas', 'seaborn', 'html', 'css', 'react', 'component']
    has_code_intent = any(k in q_lower for k in code_keywords)
    
    # Logic update: If strong phrases match, ignore code check. If noun+verb, check code.
    is_image_scope = (has_noun and has_verb) or any(phrase in q_lower for phrase in exact_phrases)
    
    # Resim ve fiil varsa, kod isteği yoksa DALL-E varsay.
    # image_path olsa bile (belki referans resimdir), DALL-E'yi deniyoruz (API desteklemese bile prompt ile deneriz).
    is_image_request = is_image_scope and (not has_code_intent) and len(question) < 1000
    
    if is_image_request:
        model = 'dall-e-3'
        routing_reason = "🎨 Görsel oluşturma isteği algılandı (resim+fiil), **DALL-E 3** seçildi."
        print(f"DEBUG: Image generation request detected: {model}")
        sys.stdout.flush()

    elif model == 'auto':
        intent = detect_intent(question, code)
        detected_lang = language_detector.detect(question, code)
        detected_intent = intent
        
        # Upgrade intent if linked to GitHub
        has_github = conversation and conversation.linked_repo
        if has_github and intent == 'general':
            intent = 'code'
            print("DEBUG: GitHub repository linked. Upgraded intent to 'code'.")

        # Eğer dil tespit edildiyse, intent 'general' olsa bile 'code' olarak işlem yap
        if detected_lang != 'unknown':
            intent = 'code'
            
        model, _router_reason = model_router.route(detected_lang, intent, prefs)
        
        # Build a clear routing message: show detected language + responding model
        lang_display = detected_lang if detected_lang != 'unknown' else 'general'
        routing_reason = (
            f"🔍 **Detected Language**: `{lang_display}` | "
            f"🤖 **Responding Model**: `{model}`\n\n"
            f"_{_router_reason}_"
        )

        # Debug print
        print(f"DEBUG: Smart Routing -> Intent: {intent}, Lang: {detected_lang} -> {model}")

    # Manual model selection: do NOT set a routing_reason.
    # The AI Reasoning Layer should only appear when auto-routing,
    # audio transcription, or image generation overrides the model.
    # (routing_reason stays None for explicit user selections)

    # Auto-routing sonrası model değişebildiği için nihai modele göre tekrar doğrula.
    if user:
        allowed_final, balance_final, req_cost_final = check_tokens(user, model)
        if not allowed_final:
            return jsonify({
                'error': f'Yetersiz Token! Bu işlem için {req_cost_final} token gerekiyor, mevcut bakiyeniz: {balance_final}.',
                'insufficient_tokens': True,
                'required': req_cost_final,
                'balance': balance_final
            }), 402

    # ── ADVANCED HEURISTIC GUARDRAIL FOR AGENT MODE ──────────────
    routing_audit = "Standard Mode"
    if agent_mode:
        q_clean = question.lower().strip()
        words = q_clean.split()
        
        # Heuristics for skipping Agent Mode even if enabled
        is_trivial = False
        chat_keywords = {'selam', 'merhaba', 'nasılsın', 'teşekkürler', 'sa', 'as', 'teşekkür', 'günaydın', 'iyi akşamlar', 'hi', 'hello', 'selamlar', 'naber'}
        action_keywords = {'araştır', 'search', 'bul', 'find', 'getir', 'fetch', 'oku', 'read', 'yaz', 'write', 'değiştir', 'edit', 'sil', 'delete', 'listele', 'list', 'check', 'kontrol', 'dosya', 'file', 'proje', 'project', 'repo', 'github'}
        
        has_action = any(kw in q_clean for kw in action_keywords)
        is_greeting = len(words) < 3 and any(w in chat_keywords for w in words)
        is_short_general = len(words) < 8 and not has_action
        
        if is_greeting:
            is_trivial = True
            routing_audit = "Skipped Agent Mode (Greeting detected)"
        elif is_short_general:
            is_trivial = True
            routing_audit = "Skipped Agent Mode (Short general query)"
        else:
            routing_audit = "Active Agent Mode (Complex/Actionable query)"

        if is_trivial:
            agent_mode = False
            routing_reason = f"🛡️ {routing_audit}, Normal Mod kullanılıyor."
            print(f"DEBUG: [RoutingAudit] {routing_audit} for query: '{question[:50]}...'")
        else:
            print(f"DEBUG: [RoutingAudit] {routing_audit} triggered for query: '{question[:50]}...'")
    # ─────────────────────────────────────────────────────────────

    answer = ""
    agent_trace = []
    agent_changed_files = []
    agent_tool_capable = False
    agent_provider = None
    agent_effective_model = None

    # Capture primitive IDs for thread-safe/session-safe streaming
    final_user_id = user.id if user else None
    final_conv_id = conversation.id if conversation else None
    final_proj_id = resolved_agent_project.id if resolved_agent_project else (conversation.project_id if conversation else None)

    # --- Model Yönlendirme Mantığı ---
    def generate_stream(u_id, c_id, p_id, source):
        nonlocal answer # Outer scope answer variable updating
        nonlocal agent_trace, agent_changed_files, agent_tool_capable, agent_provider, agent_effective_model
        full_answer = ""

        # Send routing metadata as an early event so UI can show model/language even if stream ends early.
        early_meta = {
            'meta': True,
            'routing_reason': routing_reason,
            'selected_model': model,
            'detected_language': detected_lang,
            'detected_intent': detected_intent,
            'agent_mode': bool(agent_mode),
            'agent_project_id': p_id,
            'agent_project_source': agent_project_source,
            'agent_workspace_file_count': len(workspace_files or []),
        }
        yield f"data: {json.dumps(early_meta)}\n\n"
        
        # For agent mode, send an immediate placeholder to show the system is working
        if agent_mode and model != 'dall-e-3':
            placeholder = {
                'type': 'status',
                'message': '🔄 Agent is processing your request...',
                'step': 0,
            }
            yield f"data: {json.dumps(placeholder)}\n\n"
        
        generator = None
        
        # Generator seçimi
        if model == 'dall-e-3':
             # DALL-E streaming desteklemez, senkron çağırıp yield ediyoruz
             img_response = generate_image_with_dalle(question)
             full_answer = img_response
             json_data = json.dumps({'chunk': img_response})
             yield f"data: {json_data}\n\n"
             # Continue to allow DB saving logic below to run
             generator = None # No further generation needed


        agent_executed = False
        if agent_mode and model != 'dall-e-3' and not model_image_path:
            print(f"DEBUG: Attempting Agent Mode for model {model}")
            provider_key, provider_model = _resolve_agent_provider_model(model)
            print(f"DEBUG: Resolved provider: {provider_key}, model: {provider_model}")
            if provider_key:
                if not _is_agent_model_supported(provider_key, provider_model):
                    blocked_msg = (
                        f"Agent Mode bu modelde desteklenmiyor: {provider_model}. "
                        "Agent Mode için Gemini tarafında gemini-2.5-flash, gemini-2.5-flash-lite veya bir Gemma modeli seçin."
                    )
                    json_data = json.dumps({'chunk': blocked_msg})
                    yield f"data: {json_data}\n\n"
                    full_answer = blocked_msg
                    agent_executed = True
                    agent_provider = provider_key
                    agent_effective_model = provider_model
                    agent_tool_capable = False
                    print(f"DEBUG: Agent Mode blocked for unsupported model {provider_model}")
                else:
                    print("DEBUG: Calling stream_agent_bridge...")
                    from services.agent_bridge import stream_agent_bridge

                    agent_provider = provider_key
                    agent_effective_model = provider_model

                    agent_messages = []
                    effective_history_context = agent_history_context if agent_history_context else history_context
                    for turn in effective_history_context:
                        u_text = (turn.get('user') or '').strip()
                        a_text = (turn.get('ai') or '').strip()
                        if u_text:
                            agent_messages.append({"role": "user", "content": u_text})
                        if a_text:
                            agent_messages.append({"role": "assistant", "content": a_text})
                    agent_messages.append({"role": "user", "content": question.strip() or "Unspecified"})
                    print("DEBUG Agent messages:", agent_messages)

                    for chunk_sse in stream_agent_bridge(
                        question=question,
                        code=code,
                        model=provider_model,
                        project=p_id,
                        user=u_id,
                        conversation=c_id,
                        workspace_files=workspace_files,
                        prefs=prefs,
                        messages=agent_messages[:-1],
                        history_context=effective_history_context,
                        github_context=github_context,
                        allow_write_tools=allow_write_tools,
                        search_project_callback=_agent_project_search,
                        db_read_callback=_agent_db_read,
                        invalidate_project_cache=invalidate_project_embedding_cache,
                    ):
                        yield chunk_sse

                        # Capture state for DB persistence at the end of generator_stream
                        if chunk_sse.startswith("data: "):
                            try:
                                data = json.loads(chunk_sse[6:].strip())
                                etype = data.get("type")
                                if etype == "message" or "chunk" in data:
                                    # Standard chunk or mapped agent message
                                    text_chunk = data.get("chunk") or data.get("text") or ""
                                    full_answer += text_chunk
                                elif etype == "done":
                                    agent_trace = data.get("trace") or []
                                    agent_changed_files = data.get("changed_files") or []
                                    agent_tool_capable = bool(data.get("agent_tool_capable", True))
                                    # If done event has final text, ensure it's captured
                                    if data.get("text"):
                                        # Depending on runtime, text might be empty if already streamed
                                        # but we check just in case.
                                        pass 
                                elif etype == "error":
                                    err_text = data.get("message") or data.get("error") or "Agent Mode failed before producing a response."
                                    if err_text and not full_answer:
                                        full_answer = f"[Agent error]: {err_text}"
                            except Exception:
                                pass

                    agent_executed = True
                    if not full_answer.strip():
                        full_answer = "Agent Mode finished without producing a response."
                    full_answer = post_process_response(full_answer)


        if not agent_executed:
            if 'claude' in model:
                generator = generate_claude_answer(question, code, history_context, model, model_image_path, prefs, github_context)
            elif 'gpt' in model or 'o1' in model:
                 generator = generate_gpt_answer(question, code, history_context, model, model_image_path, prefs, github_context)
            elif 'gemini' in model or 'gemma' in model:
                generator = generate_gemini_answer(question, code, history_context, model, model_image_path, prefs, github_context)

        # Ortak Generator Döngüsü
        if generator and not agent_executed:
            try:
                for chunk in generator:
                    if chunk:
                        full_answer += chunk
                        json_data = json.dumps({'chunk': chunk})
                        yield f"data: {json_data}\n\n"
                        
                # 0. Post-Processing Layer (İşlem Sonrası Katmanı)
                full_answer = post_process_response(full_answer)
                
            except Exception as e:
                err_msg = f"\n[Model Error]: {str(e)}"
                full_answer += err_msg
                json_data = json.dumps({'chunk': err_msg})
                yield f"data: {json_data}\n\n"

        # Bitiş işlemleri (Veritabanı kayıt)
        with app.app_context():
            clipped_agent_meta = _clip_agent_metadata(agent_trace, agent_changed_files)
            # Keep these available even if DB save fails; client can still close stream cleanly.
            final_data = {
                'done': True,
                'answer': full_answer,
                'routing_reason': routing_reason,
                'selected_model': model,
                'detected_language': detected_lang,
                'detected_intent': detected_intent,
                'agent_mode': bool(agent_mode),
                'agent_provider': agent_provider,
                'agent_model': agent_effective_model,
                'agent_tool_capable': agent_tool_capable,
                'agent_project_id': p_id,
                'agent_project_source': agent_project_source,
                'agent_trace': clipped_agent_meta['trace'],
                'agent_changed_files': clipped_agent_meta['changed_files'],
                'agent_trace_total': clipped_agent_meta['trace_total'],
                'agent_changed_total': clipped_agent_meta['changed_total'],
                'agent_trace_truncated': clipped_agent_meta['trace_truncated'],
                'agent_changed_truncated': clipped_agent_meta['changed_truncated'],
                'persona': prefs.get('persona', 'General User') if u_id else 'General User',
                'memory_used': bool(memory_context.get('text')),
                'memory_hits': int(memory_context.get('hit_count') or 0),
                'carryover': include_previous_modules,
                'memory_capsules': memory_context.get('hits', []),
            }
            
            # Resolve current_user early for token deduction
            current_user = db.session.get(User, u_id) if u_id else None
            
            # Only save to database if not a no_save request
            if not no_save and c_id:
                try:
                    # Session'a conversation'ı tekrar bağla/getir
                    current_conv = db.session.get(Conversation, c_id)
                    # (already resolved above)
                    
                    if not current_conv:
                        # Should not happen normally
                        current_conv = Conversation(id=c_id, user_id=u_id, title=question[:50], source=source)
                        db.session.add(current_conv)

                    # Özetleme ve kalıcı hafıza çıkarımı
                    summary = summarize_answer(full_answer)
                    extracted_memory_items = extract_memory_candidates(question, full_answer) if current_user else []

                    # Başlık güncelleme (ilk mesajsa) - kısa ve öz başlık üret
                    if not history_context:
                        current_conv.title = generate_conversation_title(question, full_answer)
                        db.session.add(current_conv)

                    history = History(
                        conversation_id=current_conv.id,
                        user_question=question,
                        code_snippet=code,
                        ai_response=full_answer,
                        selected_model=model,
                        summary=summary,
                        image_path=image_path,
                        routing_reason=routing_reason,
                        persona=prefs.get('persona', 'General User') if current_user else 'General User'
                    )
                    db.session.add(history)

                    # Flush first so the new history id can be attached to summary rows.
                    db.session.flush()

                    if current_user:
                        _upsert_conversation_summary(current_conv, history.id, current_user.id, summary, extracted_memory_items)
                        _store_memory_items(current_user.id, current_conv.id, extracted_memory_items)
                        learning_snapshot = _write_back_memory_graph(current_user.id, current_conv.id, history.id, memory_context)
                        final_data['memory_learning'] = learning_snapshot

                    db.session.commit()

                    # 4. Update final_data for frontend (only when saved)
                    final_data.update({
                        'history_id': history.id,
                        'conversation_id': current_conv.id,
                        'summary': summary,
                        'persona': history.persona
                    })
                except Exception as save_err:
                    db.session.rollback()
                    print(f"WARN: Final save failed in /api/ask stream: {save_err}")
                    sys.stdout.flush()
                    # Keep stream contract stable so frontend does not append generic error.
                    final_data['warning'] = 'response_saved_with_warning'

            # 3. Handle Token Deduction & Taste Profile (Always if user exists)
            if current_user:
                try:
                    # 💰 TOKEN EKONOMİSİ — Harcamayı düş (Sadece başarılı yanıtlarda)
                    if full_answer and len(full_answer.strip()) > 0:
                        # AI Taste Profile Güncelle (Öğrenme + Persona Analizi)
                        update_user_taste(current_user, model, full_answer, question)

                        is_compare = _parse_bool(payload.get('is_compare'))
                        if is_compare or no_save:
                            token_desc = f"Compare: {question[:30]}..."
                        else:
                            token_desc = f"Chat: {question[:30]}..."

                        # Reference ID is only available if we saved to history
                        ref_id = None
                        if 'history' in locals() and history:
                            ref_id = history.id

                        success, new_bal = deduct_tokens(
                            current_user,
                            model,
                            description=token_desc,
                            reference_id=ref_id
                        )
                        final_data['new_token_balance'] = new_bal
                        if not success:
                            final_data['token_warning'] = 'Token düşümü yapılamadı. Bakiye yetersiz olabilir.'
                    else:
                        print(f"DEBUG: Skipping token deduction for empty/failed response (User: {current_user.id})")

                    # Soru sorma XP ödülü (sadece yeni geçmiş oluşturuluyorsa veya karşılaştırma ise)
                    # Karşılaştırma için yarım XP verelim? Ya da tam verelim.
                    if not no_save:
                        xp_result = award_xp(current_user.id, XP_REWARDS['ask_question'], "Asking a Question", source='ask_question')
                        if xp_result:
                            final_data['xp_awarded'] = xp_result
                except Exception as token_err:
                    print(f"WARN: Token deduction/Taste update failed: {token_err}")

            yield f"data: {json.dumps(final_data)}\n\n"

    return Response(stream_with_context(generate_stream(final_user_id, final_conv_id, final_proj_id, source_header)), mimetype='text/event-stream')


# ==========================================
# MULTI-MODEL BLEND (ÇOKLU MODEL HARMANLAMA)
# ==========================================

from concurrent.futures import ThreadPoolExecutor, as_completed

def fetch_model_response_sync(model: str, question: str, code: str = '', prefs = None):
    """Tek bir modelden senkron yanıt al (thread içinde kullanılır)."""
    full_response = ""
    # prefs is passed directly now, no need to call get_user_preferences here
    
    try:
        if 'claude' in model:
            for chunk in generate_claude_answer(question=question, code=code, history_context=[], requested_model=model, prefs=prefs):
                full_response += chunk
        elif 'gpt' in model:
            for chunk in generate_gpt_answer(question=question, code=code, history_context=[], requested_model=model, prefs=prefs):
                full_response += chunk
        elif 'gemini' in model or 'gemma' in model:
            for chunk in generate_gemini_answer(question=question, code=code, history_context=[], requested_model=model, prefs=prefs):
                full_response += chunk
    except Exception as e:
        full_response = f"[{model} Error]: {str(e)}"
    return model, full_response


@app.route('/api/blend', methods=['POST'])
def blend_models():
    """Birden fazla modelden yanıt al ve Gemini ile harmanlayarak tek yanıt döndür."""
    data = request.get_json(silent=True) or {}
    question = data.get('question', '')
    code = data.get('code', '')
    models = data.get('models', [])  # List of model names
    conversation_id = data.get('conversation_id')  # Optional conversation ID
    
    if not question:
        return jsonify({'error': 'Question required'}), 400
    
    if not models or len(models) < 2:
        return jsonify({'error': 'At least 2 models must be selected'}), 400
    
    if len(models) > 4:
        return jsonify({'error': 'Maximum 4 models can be selected'}), 400
    
    print(f"BLEND İsteği: {len(models)} model - {models}")
    
    # Get current user for conversation saving and personalization
    user = get_current_user()
    prefs = get_user_preferences(user)
    persona = prefs.get('persona', 'General User')
    expertise = prefs.get('expertise', 'intermediate')

    # 💰 TOKEN EKONOMİSİ — Bakiye Kontrolü (Blend için özel maliyet)
    if user:
        allowed, balance, req_cost = check_tokens(user, 'model-blend')
        if not allowed:
            return jsonify({
                'error': f'Yetersiz Token! Blend işlemi için {req_cost} token gerekiyor, bakiyeniz: {balance}.',
                'insufficient_tokens': True
            }), 402

    # Legacy quota check
    try:
        blend_estimated_tokens = _estimate_tokens_for_request(question, code) * max(1, len(models))
        quota_result = consume_plan_quota(user, blend_estimated_tokens, request_weight=max(2, len(models)))
        if not quota_result.get('allowed', True):
            # ...
            pass
    except Exception as e:
        print(f"Blend quota check failed (continuing): {e}")
    
    # Create or get conversation
    if user:
        conversation = None
        if conversation_id:
            conversation = Conversation.query.filter_by(id=conversation_id, user_id=user.id).first()

        # If conversation_id is missing/stale, create a fresh conversation so save + token deduction never gets skipped.
        if not conversation:
            conversation = Conversation(
                user_id=user.id,
                title=question[:50] + ('...' if len(question) > 50 else ''),
                source='web',
                created_at=datetime.datetime.now()
            )

            # If repo was pre-verified and passed here, link it to the new conversation
            repo_param = data.get('repo')
            branch_param = data.get('branch', 'main')
            if repo_param:
                conversation.linked_repo = repo_param
                conversation.repo_branch = branch_param
                print(f"DEBUG: Pre-linking repo {repo_param} to new BLEND conversation")

            db.session.add(conversation)
            db.session.commit()

        conversation_id = conversation.id
    else:
        conversation = None
        conversation_id = None
    
    def generate_blend_stream():
        model_responses = {}
        xp_result = None
        new_token_balance = None
        history_entry = None
        
        # 1. Paralel olarak tüm modellerden yanıt al
        yield f"data: {json.dumps({'status': 'fetching', 'message': 'Sending query to selected models...'})}\n\n"
        
        with ThreadPoolExecutor(max_workers=4) as executor:
            futures = {
                executor.submit(fetch_model_response_sync, model, question, code, prefs): model 
                for model in models
            }
            
            completed = 0
            try:
                # Add 30s timeout to prevent infinite freezing
                for future in as_completed(futures, timeout=30):
                    model_name, response = future.result()
                    model_responses[model_name] = response
                    completed += 1
                    yield f"data: {json.dumps({'status': 'progress', 'completed': completed, 'total': len(models), 'model': model_name})}\n\n"
            except Exception as e:
                # Timeout or other error during execution
                print(f"Blending timeout/error: {e}")
                for f in futures:
                    f.cancel()
                # Fill missing models with error text so blending can continue
                for f, m_name in futures.items():
                    if m_name not in model_responses:
                        model_responses[m_name] = f"[{m_name} Error]: Timeout or execution failed."
        
        # 2. Tüm yanıtları Gemini ile harmanlama
        yield f"data: {json.dumps({'status': 'blending', 'message': 'Blending responses...'})}\n\n"
        
        # Harmanlama prompt'u oluştur
        blend_prompt = f"""Below are the answers given by different AI models to the same question.
Analyze these responses and combine the best parts of all of them to create a single, comprehensive and consistent "super response".

**USER CONTEXT:**
- Persona: {persona}
- Expertise Level: {expertise}
- Interests: {', '.join(prefs.get('interests', []))}

**QUESTION:** {question}

"""
        for model_name, response in model_responses.items():
            blend_prompt += f"**{model_name} Response:**\n{response}\n\n---\n\n"
        
        blend_prompt += """
**Your Task:**
1. Identify valuable information in each model's response
2. Compare conflicting information and choose the most accurate one
3. Combine all of these into a single, fluent and comprehensive response
4. Do not mention source models, only give the blended result
5. IMPORTANT: Respond in the same language as the user's question (e.g., if the question is in Turkish, respond in Turkish).
"""
        
        # Gemini ile harmanla (Fallback mekanizmalı)
        blended_response = ""
        blender_models = ['gemini-2.5-flash', 'gpt-4o-mini', 'claude-sonnet-4-5-20250929']
        blender_error = None
        
        success = False
        for blender_model in blender_models:
            if success: break
            
            try:
                # Modeli seç ve generate et
                generator = None
                if 'gemini' in blender_model:
                    generator = generate_gemini_answer(blend_prompt, '', [], blender_model)
                elif 'gpt' in blender_model:
                    generator = generate_gpt_answer(blend_prompt, '', [], blender_model)
                elif 'claude' in blender_model:
                    generator = generate_claude_answer(blend_prompt, '', [], blender_model)
                
                temp_response = ""
                error_in_stream = False
                
                if generator:
                    for chunk in generator:
                        # Kota hatası kontrolü (Gemini için)
                        if "[Error]: Quota limit exceeded" in chunk:
                            error_in_stream = True
                            break
                            
                        temp_response += chunk
                        yield f"data: {json.dumps({'status': 'streaming', 'chunk': chunk})}\n\n"
                
                if not error_in_stream and temp_response:
                    blended_response = temp_response
                    success = True
                else:
                    # Bu model başarısız oldu, bir sonrakine geç
                    continue
                    
            except Exception as e:
                blender_error = e
                continue

        if not success:
             error_msg = str(blender_error) if blender_error else "All blending models failed."
             blended_response = f"Blending error: {error_msg}"
             yield f"data: {json.dumps({'status': 'error', 'message': error_msg})}\n\n"
        
        # 2.5 Referee (Judge) Call for Explainable AI
        referee_reasoning = ""
        if success and blended_response:
            yield f"data: {json.dumps({'status': 'refereeing', 'message': 'AI Referee is evaluating the models...'})}\n\n"
            
            referee_prompt = f"""Compare the following AI model responses and the final blended result.
Provide a clear reasoning/justification for why this blended response was chosen and how the models performed.

**QUESTION:** {question}

"""
            for model_name, response in model_responses.items():
                referee_prompt += f"**{model_name} Response:**\n{response}\n\n"
            
            referee_prompt += f"\n**FINAL BLENDED RESPONSE:**\n{blended_response}\n\n"
            
            referee_prompt += """
**Referee Task:**
1. Briefly compare model performances and the final blended result.
2. Which model was strongest? Which had errors or omissions?
3. Justify the blended result.

**Constraints:**
- Keep the entire evaluation EXTREMELY CONCISE and bulleted.
- EXPLICITLY STATE why you used which model (e.g., "GPT-4o provided better code, so it was prioritized for the solution").
- Explain the contribution of each model to the final answer.
- IMPORTANT: Respond in the same language as the user's question (e.g., if the question is in Turkish, respond in Turkish).
"""
            try:
                # Use Gemini 2.5 Flash as referee
                referee_model = genai.GenerativeModel('models/gemini-2.5-flash')
                referee_result = referee_model.generate_content(referee_prompt)
                referee_reasoning = getattr(referee_result, "text", "").strip()
                yield f"data: {json.dumps({'status': 'referee_done', 'reasoning': referee_reasoning})}\n\n"
            except Exception as ref_err:
                print(f"Referee error: {ref_err}")
                referee_reasoning = f"Referee failed: {str(ref_err)}"
        
        # 3. Save to database if user is logged in
        if user and conversation and blended_response:
            try:
                history_entry = History(
                    conversation_id=conversation.id,
                    user_question=question,
                    ai_response=blended_response,
                    code_snippet=code if code else None,
                    selected_model=f"Blend: {', '.join(models[:2])}{'...' if len(models) > 2 else ''}",
                    timestamp=_utcnow(),
                    reasoning=referee_reasoning,
                    routing_reason=f"Blended {', '.join(models)} for enhanced accuracy",
                    persona=persona
                )
                db.session.add(history_entry)
                db.session.commit()
                print(f"DEBUG: Saved Blend History item {history_entry.id}")
                
                # Update Taste Profile
                charge_user = db.session.get(User, user.id)
                update_user_taste(charge_user, "blend", blended_response, question)

                # Blend modu için token düşümü
                success, new_bal = deduct_tokens(
                    charge_user,
                    'model-blend',
                    description=f"Blend: {question[:30]}...",
                    reference_id=history_entry.id
                )
                if success:
                    new_token_balance = new_bal
                else:
                    print(f"WARN: Token deduction failed for blend user_id={user.id}, balance={new_bal}")

                # Keep blend and ask flows consistent for gamification rewards.
                xp_result = award_xp(user.id, XP_REWARDS['ask_question'], "Asking a Question", source='ask_question')
            except Exception as db_err:
                print(f"Database save error (blend): {db_err}")
        
        # 4. Return final result
        final_data = {
            'done': True,
            'blended_response': blended_response,
            'source_models': list(model_responses.keys()),
            'individual_responses': model_responses,
            'conversation_id': conversation_id,
            'history_id': history_entry.id if history_entry else None,
            'persona': persona,
            'routing_reason': f"Blended {', '.join(models)} for enhanced accuracy"
        }
        if xp_result:
            final_data['xp_awarded'] = xp_result
        if new_token_balance is not None:
            final_data['new_token_balance'] = new_token_balance
        yield f"data: {json.dumps(final_data)}\n\n"
    
    return Response(stream_with_context(generate_blend_stream()), mimetype='text/event-stream')


@app.route('/api/conversations', methods=['GET'])
def list_conversations():
    user = get_current_user()
    project_id = request.args.get('project_id')
    
    if user:
        # Kullanıcının konuşmaları
        query = Conversation.query.filter_by(user_id=user.id)
        
        if project_id:
            query = query.filter_by(project_id=project_id)
        else:
            # Default list should show only general chats.
            # Project-linked conversations are shown only in Project Workspace.
            query = query.filter(Conversation.project_id.is_(None))
            
        # Filter by source (Web vs Extension)
        source_filter = request.headers.get("X-Client-Source", "web")
        if source_filter == 'vscode': source_filter = 'extension'
        query = query.filter(Conversation.source == source_filter)

        # Exclude community posts and include archive filter logic if needed
        convs = query.filter(Conversation.is_deleted == False)\
            .filter(db.or_(Conversation.is_archived == False, Conversation.is_archived == None))\
            .order_by(Conversation.is_pinned.desc(), Conversation.created_at.desc())\
            .all()
    else:
        convs = []
    
    return jsonify({'conversations': [serialize_conversation(c) for c in convs]})


@app.route('/api/conversations', methods=['POST'])
@jwt_required()
def create_conversation():
    user = get_current_user()
    data = request.json or {}
    title = data.get('title', 'New Chat')
    project_id = data.get('project_id')
    
    conversation = Conversation(
        user_id=user.id,
        title=title,
        project_id=project_id,
        source='web',
        created_at=datetime.datetime.now()
    )
    db.session.add(conversation)
    db.session.commit()
    
    return jsonify({'conversation': serialize_conversation(conversation)})


# --- CONVERSATION MANAGEMENT ENDPOINTS ---

@app.route('/api/conversations/<int:conversation_id>/rename', methods=['PUT'])
@jwt_required()
def rename_conversation(conversation_id):
    """Konuşma başlığını yeniden adlandırır."""
    conversation = Conversation.query.get_or_404(conversation_id)
    user = get_current_user()
    
    if conversation.user_id != user.id:
        return jsonify({'error': 'Unauthorized access'}), 403
    
    data = request.json or {}
    new_title = data.get('title', '').strip()
    
    if not new_title:
        return jsonify({'error': 'Title cannot be empty'}), 400
    
    conversation.title = new_title[:255]  # Max 255 karakter
    db.session.commit()
    
    return jsonify({'conversation': serialize_conversation(conversation), 'message': 'Title updated'})


@app.route('/api/conversations/<int:conversation_id>/pin', methods=['PUT'])
@jwt_required()
def pin_conversation(conversation_id):
    """Konuşmayı sabitle veya sabitlemeden çıkar."""
    conversation = Conversation.query.get_or_404(conversation_id)
    user = get_current_user()
    
    if conversation.user_id != user.id:
        return jsonify({'error': 'Unauthorized access'}), 403
    
    conversation.is_pinned = not conversation.is_pinned
    db.session.commit()
    
    status = 'pinned' if conversation.is_pinned else 'unpinned'
    return jsonify({'conversation': serialize_conversation(conversation), 'message': f'Chat {status}'})


@app.route('/api/conversations/<int:conversation_id>/archive', methods=['PUT'])
@jwt_required()
def archive_conversation(conversation_id):
    """Konuşmayı arşivle veya arşivden çıkar."""
    conversation = Conversation.query.get_or_404(conversation_id)
    user = get_current_user()
    
    if conversation.user_id != user.id:
        return jsonify({'error': 'Unauthorized access'}), 403
    
    conversation.is_archived = not conversation.is_archived
    db.session.commit()
    
    status = 'archived' if conversation.is_archived else 'unarchived'
    return jsonify({'conversation': serialize_conversation(conversation), 'message': f'Chat {status}'})




@app.route('/api/conversations/archived', methods=['GET'])
@jwt_required()
def list_archived_conversations():
    """Arşivlenmiş konuşmaları listeler."""
    user = get_current_user()
    
    community_conv_ids = db.session.query(History.conversation_id)\
        .filter(History.selected_model == 'Community')\
        .subquery()

    convs = Conversation.query.filter_by(user_id=user.id)\
        .filter(Conversation.id.notin_(community_conv_ids))\
        .filter(Conversation.is_archived == True)\
        .order_by(Conversation.created_at.desc())\
        .all()
    
    return jsonify({'conversations': [serialize_conversation(c) for c in convs]})


@app.route('/api/community/my-posts', methods=['GET'])
@jwt_required()
def get_user_posts():
    user = get_current_user()
    # Kullanıcının 'Community' olarak işaretlenmiş, silinmemiş postlarını getir
    items = History.query.join(Conversation)\
        .filter(Conversation.user_id == user.id)\
        .filter(History.selected_model == 'Community')\
        .filter(History.is_deleted == False)\
        .order_by(History.timestamp.desc())\
        .all()
    return jsonify({'posts': [serialize_history(h) for h in items]})


@app.route('/api/conversations/<int:conversation_id>', methods=['GET'])
def get_conversation(conversation_id):
    conversation = Conversation.query.filter_by(id=conversation_id, is_deleted=False).first_or_404()
    # Güvenlik: Eğer kullanıcı giriş yapmışsa ve bu konuşma başkasınınsa erişimi engelle (admin hariç)
    user = get_current_user()
    if conversation.user_id and (not user or (user.id != conversation.user_id and not user.is_admin)):
         return jsonify({'error': 'Unauthorized access'}), 403

    items = History.query.filter_by(conversation_id=conversation_id, is_deleted=False).order_by(History.timestamp.asc()).all()
    return jsonify({
        'conversation': serialize_conversation(conversation),
        'history': [serialize_history(h) for h in items]
    })


@app.route('/api/conversations/<int:conversation_id>', methods=['DELETE'])
def delete_conversation(conversation_id):
    try:
        conversation = Conversation.query.filter_by(id=conversation_id, is_deleted=False).first_or_404()
        user = get_current_user()
        
        # Güvenlik kontrolü
        if not user:
            return jsonify({'error': 'Authentication required'}), 401
        
        if conversation.user_id and user.id != conversation.user_id and not user.is_admin:
             return jsonify({'error': 'Unauthorized access'}), 403
        
        # SaaS-Grade Deactivation Engine (Async)
        LifecycleOrchestrator.deactivate_conversation(conversation_id)
        
        return jsonify({'status': 'deactivation_initiated', 'message': 'Sohbet siliniyor...'})
    except Exception as e:
        print(f"Error in delete_conversation: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/conversations/<int:conversation_id>/restore', methods=['POST'])
@jwt_required()
def restore_conversation(conversation_id):
    """Geri yükleme motorunu tetikler."""
    user = get_current_user()
    conversation = Conversation.query.filter_by(id=conversation_id, user_id=user.id, is_deleted=True).first_or_404()
    
    # SaaS-Grade Restoration Engine (Async)
    LifecycleOrchestrator.restore_conversation(conversation_id)
    return jsonify({'status': 'restoration_initiated', 'message': 'Sohbet geri yükleniyor...'})
@app.route('/api/conversations/<int:conversation_id>/history', methods=['POST'])
@jwt_required()
def add_history_item(conversation_id):
    """Konuşmaya manuel olarak (generate etmeden) bir geçmiş öğesi ekler."""
    conversation = Conversation.query.get_or_404(conversation_id)
    user = get_current_user()
    
    if conversation.user_id != user.id:
        return jsonify({'error': 'Unauthorized access'}), 403
        
    data = request.json or {}
    user_question = data.get('user_question')
    ai_response = data.get('ai_response')
    selected_model = data.get('selected_model', 'Unknown')
    
    if not user_question or not ai_response:
        return jsonify({'error': 'Question and answer required'}), 400
        
    history = History(
        conversation_id=conversation.id,
        user_question=user_question,
        ai_response=ai_response,
        selected_model=selected_model,
        timestamp=datetime.datetime.now()
    )
    
    db.session.add(history)
    conversation.updated_at = datetime.datetime.now()
    db.session.commit()
    
    return jsonify(serialize_history(history))


@app.route('/api/posts/<int:post_id>', methods=['DELETE'])
@jwt_required()
def delete_post(post_id):
    """Kullanıcının kendi gönderisini (History kaydını) siler (deaktive eder)."""
    user = get_current_user()
    history = History.query.filter_by(id=post_id, is_deleted=False).first_or_404()
    
    # Ownership Check
    if not history.conversation or (history.conversation.user_id != user.id and not user.is_admin):
        return jsonify({'error': 'Unauthorized access'}), 403
        
    LifecycleOrchestrator.deactivate_history_item(post_id)
    return jsonify({'status': 'deleted', 'message': 'Gönderi silindi.'})


@app.route('/api/posts/<int:post_id>', methods=['PUT'])
@jwt_required()
def edit_post(post_id):
    """Kullanıcının kendi gönderisini düzenler."""
    user = get_current_user()
    
    # Gönderiyi (History kaydını) bul
    history = History.query.get_or_404(post_id)
    
    # Gönderi sahibini kontrol et
    if not history.conversation or not history.conversation.user:
        return jsonify({'error': 'Post not found.'}), 404
    
    if history.conversation.user_id != user.id and not user.is_admin:
        return jsonify({'error': 'You do not have permission to edit this post.'}), 403
    
    data = request.json or {}
    new_question = data.get('user_question', '').strip()
    new_summary = data.get('summary', '').strip()
    
    if not new_question:
        return jsonify({'error': 'Question field cannot be empty.'}), 400
    
    try:
        history.user_question = new_question
        if new_summary:
            history.summary = new_summary
        
        db.session.commit()
        
        return jsonify({
            'message': 'Post updated.',
            'post': serialize_history(history)
        })
    
    except Exception as e:
        db.session.rollback()
        print(f"Gönderi düzenleme hatası: {e}")
        return jsonify({'error': 'An error occurred while editing the post.'}), 500


@app.route('/api/notifications', methods=['GET'])
def get_notifications():
    user = get_current_user()
    if not user:
        return jsonify({'error': 'Unauthorized access'}), 401

    results = []

    # 1. Yorum Bildirimleri (Answer)
    # Kullanıcının sorularına (History) gelen cevaplar - Community postları hariç
    answers = db.session.query(Answer, History, Conversation)\
        .join(History, Answer.history_id == History.id)\
        .join(Conversation, History.conversation_id == Conversation.id)\
        .filter(Conversation.user_id == user.id)\
        .filter(Answer.author_id != user.id)\
        .filter(History.selected_model != 'Community')\
        .order_by(Answer.created_at.desc())\
        .limit(20)\
        .all()

    for answer, history, conversation in answers:
        results.append({
            'id': f"ans-{answer.id}",
            'type': 'comment',
            'author': answer.author,
            'message': f"{answer.author} added a solution to your question!",
            'question_title': conversation.title,
            'created_at': answer.created_at.strftime('%Y-%m-%d %H:%M'),
            'timestamp_obj': answer.created_at,
            'conversation_id': conversation.id,
            'history_id': history.id
        })

    # 2. Beğeni Bildirimleri (PostLike)
    # Kullanıcının sorularına (History) gelen beğeniler - Community postları hariç
    post_likes = db.session.query(PostLike, User, History, Conversation)\
        .join(User, PostLike.user_id == User.id)\
        .join(History, PostLike.history_id == History.id)\
        .join(Conversation, History.conversation_id == Conversation.id)\
        .filter(Conversation.user_id == user.id)\
        .filter(PostLike.user_id != user.id)\
        .filter(History.selected_model != 'Community')\
        .order_by(PostLike.timestamp.desc())\
        .limit(20)\
        .all()

    for like, liker, history, conversation in post_likes:
        results.append({
            'id': f"plike-{like.id}",
            'type': 'like',
            'author': liker.display_name,
            'message': f"{liker.display_name} liked your post!",
            'question_title': conversation.title,
            'created_at': like.timestamp.strftime('%Y-%m-%d %H:%M'),
            'timestamp_obj': like.timestamp,
            'conversation_id': conversation.id,
            'history_id': history.id
        })

    # 3. Yorum Beğeni Bildirimleri (AnswerLike)
    # Kullanıcının cevaplarına (Answer) gelen beğeniler
    answer_likes = db.session.query(AnswerLike, User, Answer, History, Conversation)\
        .join(User, AnswerLike.user_id == User.id)\
        .join(Answer, AnswerLike.answer_id == Answer.id)\
        .join(History, Answer.history_id == History.id)\
        .join(Conversation, History.conversation_id == Conversation.id)\
        .filter(Answer.author_id == user.id)\
        .filter(AnswerLike.user_id != user.id)\
        .order_by(AnswerLike.timestamp.desc())\
        .limit(20)\
        .all()

    for like, liker, answer, history, conversation in answer_likes:
        results.append({
            'id': f"alike-{like.id}",
            'type': 'like',
            'author': liker.display_name,
            'message': f"{liker.display_name} liked your comment!",
            'question_title': conversation.title,
            'created_at': like.timestamp.strftime('%Y-%m-%d %H:%M'),
            'timestamp_obj': like.timestamp,
            'conversation_id': conversation.id,
            'answer_id': answer.id,
            'history_id': history.id
        })

    # 4. Notification tablosundan tüm bildirimler (follow, like, comment)
    all_notifications = Notification.query.filter_by(user_id=user.id)\
        .order_by(Notification.created_at.desc())\
        .limit(30)\
        .all()
        
    for n in all_notifications:
        # Related user name bul
        related_user = db.session.get(User, n.related_user_id) if n.related_user_id else None
        author_name = related_user.display_name if related_user else "Birisi"
        
        # Emoji seç
        emoji = '🔔'
        if n.type == 'follow':
            emoji = '👥'
        elif n.type == 'like':
            emoji = '❤️'
        elif n.type == 'comment':
            emoji = '💬'
        
        results.append({
            'id': f"notif-{n.id}",
            'real_id': n.id,
            'type': n.type,
            'author': author_name,
            'message': n.message,
            'related_user_id': n.related_user_id,
            'question_title': '',
            'created_at': n.created_at.strftime('%Y-%m-%d %H:%M'),
            'timestamp_obj': n.created_at,
            'conversation_id': None,
            'history_id': n.related_post_id,
            'is_new_system': True
        })

    # Okunmuş bildirimleri al
    read_notifications = {n.notification_id for n in NotificationRead.query.filter_by(user_id=user.id).all()}
    # Silinmiş bildirimleri al
    hidden_notifications = {n.notification_id for n in NotificationHidden.query.filter_by(user_id=user.id).all()}

    # Tarihe göre sırala (En yeni en üstte)
    results.sort(key=lambda x: x['timestamp_obj'], reverse=True)
    
    final_results = []
    # Timestamp objesini JSON için kaldır ve is_read ekle
    for r in results:
        # Gizlenmiş bildirimleri atla
        if r['id'] in hidden_notifications:
            continue
            
        # Tüm bildirimler için NotificationRead tablosunu kullan
        r['is_read'] = r['id'] in read_notifications
              
        if 'timestamp_obj' in r:
            del r['timestamp_obj']
            
        final_results.append(r)

    return jsonify(final_results)


@app.route('/api/notifications/read', methods=['POST'])
@jwt_required()
def mark_notification_read():
    user = get_current_user()
    data = request.json
    notification_id = data.get('notification_id')
    
    if not notification_id:
        return jsonify({'error': 'Notification ID required'}), 400
        
    # Zaten okunmuş mu?
    existing = NotificationRead.query.filter_by(user_id=user.id, notification_id=notification_id).first()
    if not existing:
        new_read = NotificationRead(user_id=user.id, notification_id=notification_id)
        db.session.add(new_read)
        db.session.commit()
        
    return jsonify({'status': 'marked_read'})


@app.route('/api/notifications/delete', methods=['POST'])
@jwt_required()
def delete_notification():
    user = get_current_user()
    data = request.json
    notification_id = data.get('notification_id')
    
    if not notification_id:
        return jsonify({'error': 'Notification ID required'}), 400
        
    # Zaten silinmiş mi?
    existing = NotificationHidden.query.filter_by(user_id=user.id, notification_id=notification_id).first()
    if not existing:
        new_hidden = NotificationHidden(user_id=user.id, notification_id=notification_id)
        db.session.add(new_hidden)
        db.session.commit()
        
    return jsonify({'status': 'deleted'})
@app.route('/api/history', methods=['GET'])
def get_history():
    items = History.query.filter_by(is_deleted=False).order_by(History.timestamp.desc()).limit(20).all()
    return jsonify({'history':[serialize_history(h) for h in items]})


@app.route('/api/popular', methods=['GET'])
def get_popular():
    items = History.query.filter_by(is_deleted=False)\
        .order_by(History.likes.desc(), History.timestamp.desc())\
        .limit(5).all()
    return jsonify({'popular':[serialize_history(h) for h in items]})


@app.route('/api/stats/model-usage', methods=['GET'])
@jwt_required(optional=True)
def get_model_usage():
    """Returns model usage statistics and estimated costs for current user"""
    try:
        user_id = get_jwt_identity()
        if not user_id:
            return jsonify({'success': False, 'error': 'Not authenticated'}), 401
        
        # Mevcut kullanıcı için sorguyu filtrele (tüm zamanlar)
        results = db.session.query(
            History.selected_model,
            db.func.count(History.id)
        ).filter(
            History.conversation.has(user_id=user_id)
        ).group_by(History.selected_model).all()

        usage_data = []
        for model_name, count in results:
            if not model_name: 
                model_name = "Unknown"
            usage_data.append({
                'model': model_name,
                'count': count
            })

        return jsonify({
            'success': True,
            'stats': usage_data
        })
    except Exception as e:
        print(f"Stats Error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


from models import db, User, Conversation, History, Answer, PostLike, AnswerLike, NotificationRead, NotificationHidden, VSCodeLoginState, VSCodeOTP

# ... (existing imports)

@app.route('/api/history/<int:history_id>/like', methods=['POST'])
@jwt_required()
def like_history(history_id: int):
    history = History.query.get_or_404(history_id)
    user = get_current_user()
    
    # Check if already liked
    existing_like = PostLike.query.filter_by(user_id=user.id, history_id=history_id).first()
    if existing_like:
        # Unlike
        db.session.delete(existing_like)
        history.likes = max((history.likes or 0) - 1, 0)
        db.session.commit()
        return jsonify({'likes': history.likes, 'status': 'unliked'})

    # Add like record
    new_like = PostLike(user_id=user.id, history_id=history_id)
    db.session.add(new_like)
    
    # Increment counter
    history.likes = (history.likes or 0) + 1
    
    # Bildirim oluştur (gönderi sahibine)
    if history.conversation and history.conversation.user_id and history.conversation.user_id != user.id:
        notification = Notification(
            user_id=history.conversation.user_id,
            type='like',
            message=f'{user.display_name} gönderinizi beğendi',
            related_user_id=user.id,
            related_post_id=history_id
        )
        db.session.add(notification)
        
        # Beğeni alan kişiye XP ver
        award_xp(history.conversation.user_id, 5, "Received a Like", source='received_like_post')
    
    db.session.commit()
    return jsonify({'likes': history.likes, 'status': 'liked'})


@app.route('/api/answers/<int:answer_id>/like', methods=['POST'])
@jwt_required()
def like_answer(answer_id: int):
    answer = Answer.query.get_or_404(answer_id)
    user = get_current_user()

    # Check if already liked
    existing_like = AnswerLike.query.filter_by(user_id=user.id, answer_id=answer_id).first()
    if existing_like:
        # Unlike
        db.session.delete(existing_like)
        answer.likes = max((answer.likes or 0) - 1, 0)
        db.session.commit()
        return jsonify({'likes': answer.likes, 'status': 'unliked'})

    # Add like record
    new_like = AnswerLike(user_id=user.id, answer_id=answer_id)
    db.session.add(new_like)

    # Increment counter
    answer.likes = (answer.likes or 0) + 1
    
    # Bildirim oluştur (yorum sahibine)
    if answer.author_id and answer.author_id != user.id:
        notification = Notification(
            user_id=answer.author_id,
            type='like',
            message=f'{user.display_name} yorumunuzu beğendi',
            related_user_id=user.id,
            related_post_id=answer.history_id
        )
        db.session.add(notification)
        
        # Beğeni alan kişiye XP ver
        award_xp(answer.author_id, 5, "Received a Like on Comment", source='received_like_answer')
    
    db.session.commit()
    return jsonify({'likes': answer.likes, 'status': 'liked'})


@app.route('/api/history/<int:history_id>', methods=['DELETE'])
def delete_history(history_id: int):
    history = History.query.filter_by(id=history_id, is_deleted=False).first_or_404()
    # Güvenlik kontrolü genelde conversation üzerinden yapılır
    user = get_current_user()
    if history.conversation and history.conversation.user_id != user.id and not user.is_admin:
        return jsonify({'error': 'Unauthorized access'}), 403
        
    LifecycleOrchestrator.deactivate_history_item(history_id)
    return jsonify({'status': 'deleted'})


@app.route('/api/history/<int:history_id>/answers', methods=['GET'])
def list_answers(history_id: int):
    History.query.get_or_404(history_id)
    answers = Answer.query.filter_by(history_id=history_id)\
        .order_by(Answer.likes.desc(), Answer.created_at.desc()).all()
    return jsonify({'answers': [serialize_answer(a) for a in answers]})


@app.route('/api/github/blueprint', methods=['GET'])
@jwt_required(optional=True)
def generate_blueprint():
    repo_param = request.args.get('repo')
    branch_param = request.args.get('branch', 'main')

    if not repo_param:
        return jsonify({'error': 'Repository parameter is required'}), 400

    parser = GitHubParser()
    tree = parser.get_repo_tree(repo_param, branch_param)
    
    if not tree:
         return jsonify({'error': 'Failed to fetch repository tree or repository is empty.'}), 404

    tree_str = parser.format_tree_for_prompt(tree)

    prompt = f"""You are a senior enterprise software architect.

Analyze the repository structure below and generate a DETAILED and ACTIONABLE "Project Blueprint" in Markdown.

Output requirements:
1. Use concrete findings from folder and file names.
2. Do not use vague placeholders like "Identified project component".
3. If a detail cannot be fully proven, provide a best-effort inference and clearly label it as "Inferred".
4. Include meaningful Mermaid diagrams that show interactions, not only root-folder listings.

Your blueprint MUST include all sections below:

## 1. Executive Summary
- Project purpose, business/technical domain, and likely users.
- Overall architecture style (monolith, layered, service-based, etc.).
- Maturity/readiness signals from repo contents.

## 2. System Architecture
- Major layers (UI, API, domain, data, integrations) and responsibilities.
- Core runtime flow at a high level.
- Mermaid `graph TD` with real component relations.

## 3. Key Components And Responsibilities
For each important module/directory:
- Responsibility
- Key files
- Inputs/outputs
- Upstream/downstream dependencies

## 4. Technology Stack
- Backend (languages, frameworks, DB, ORM)
- Frontend (framework, state/data flow hints, build tool)
- Tooling (tests, linting, formatting, CI/CD, containerization)

## 5. API Surface
- Main endpoints/entry points inferred from routes/controllers.
- Group endpoints by feature area.
- Mention auth, request/response expectations where inferable.

## 6. Data Model And Persistence
- Database technology and schema clues.
- Core entities and relationships (inferred from model names/files).
- Migration/versioning approach if present.

## 7. Data Flow (Sequence)
- Provide a Mermaid `sequenceDiagram` for one critical user flow.

## 8. Build, Run, Deploy
- Local run strategy
- Build scripts/tasks
- Deployment/runtime artifacts (Dockerfile, Procfile, compose, etc.)
- Environment/configuration strategy

## 9. Risks And Improvement Opportunities
- Architectural risks
- Testing gaps
- Security/operability concerns
- Prioritized next improvements

## 10. Project Structure Snapshot
- Clean tree of key directories and representative files (not full dump).

Formatting rules:
- Write in professional technical English.
- Use concise but informative paragraphs and bullet points.
- Keep section headings exactly as provided above.
- Return ONLY Markdown content.

Repository structure input:
{tree_str[:12000]}
"""

    # Updated model fallback chain - Claude Sonnet first (best for documentation), then OpenAI, then Gemini
    fallback_chain = [
        {'type': 'anthropic', 'name': 'claude-sonnet-4-5-20250929'},
        {'type': 'gemini', 'name': 'gemini-2.5-flash-lite-preview-02-05'},
        {'type': 'gemini', 'name': 'gemini-3.1-flash-lite'},
        {'type': 'gemini', 'name': 'gemini-2.5-flash'},
    ]
    
    last_error = "Unknown error"
    for model_info in fallback_chain:
        try:
            model_type = model_info['type']
            model_name = model_info['name']
            
            if model_type == 'anthropic' and claude_client:
                print(f"⚡ Trying Blueprint generation with Claude: {model_name}")
                response = claude_client.messages.create(
                    model=model_name,
                    max_tokens=4000,
                    messages=[
                        {"role": "user", "content": prompt}
                    ]
                )
                if response.content and response.content[0].text:
                    content = response.content[0].text.strip()
                    if len(content) > 100:
                        print(f"✓ Blueprint generated successfully via Claude ({len(content)} chars)")
                        return jsonify({'markdown': content})
                        
            elif model_type == 'gemini':
                full_m_name = model_name if model_name.startswith('models/') else f"models/{model_name}"
                print(f"⚡ Trying Blueprint generation with Gemini: {full_m_name}")
                model = genai.GenerativeModel(full_m_name)
                response = model.generate_content(prompt, request_options={"timeout": 60})
                if response and response.text and len(response.text.strip()) > 100:
                    print(f"✓ Blueprint generated successfully ({len(response.text)} chars)")
                    return jsonify({'markdown': response.text})
                    
                if response.choices and response.choices[0].message.content:
                    content = response.choices[0].message.content.strip()
                    if len(content) > 100:
                        print(f"✓ Blueprint generated successfully via OpenAI ({len(content)} chars)")
                        return jsonify({'markdown': content})
                        
        except Exception as e:
            error_msg = str(e)
            print(f"❌ Blueprint trial with {model_info['name']} failed: {error_msg}")
            last_error = error_msg
            continue

    # If all models fail, return helpful error
    return jsonify({'error': f"Failed to generate blueprint after multiple attempts. Last error: {last_error}"}), 500


@app.route('/api/github/health', methods=['GET'])
def get_code_health():
    try:
        repo_param = request.args.get('repo')
        branch_param = request.args.get('branch', 'main')

        # Guard: repo parameter is required and cannot be null/empty
        if not repo_param or repo_param == 'null' or repo_param == 'undefined':
            return jsonify({'error': 'Repository parameter is required. Please link a repository first.'}), 400

        print(f"📊 Starting health check for: {repo_param} (branch: {branch_param})")

        # ===========================
        # GERÇEK REPO ANALİZİ
        # ===========================
        parser = GitHubParser()
        tree = parser.get_repo_tree(repo_param, branch_param)
        
        if not tree:
            print(f"❌ Failed to fetch tree for {repo_param}")
            return jsonify({'error': 'Failed to fetch repository tree'}), 404
        
        print(f"✓ Fetched tree with {len(tree)} items")
        
        # Dosya yollarını analiz et (safe extraction)
        try:
            paths = []
            for item in tree:
                if isinstance(item, dict) and item.get('type') == 'blob':
                    path = item.get('path', '')
                    if path:
                        paths.append(path.lower())
        except Exception as e:
            print(f"⚠️ Error parsing tree structure: {e}")
            paths = []
            
        total_files = len(paths)
        print(f"📁 Extracted {total_files} file paths")
        
        # 1. SECURITY SCORE (0-100)
        security_score = 100
        security_issues = []
        
        # Tehlikeli dosya/pattern'leri ara
        dangerous_patterns = {
            '.env': 15,
            'secret': 15,
            'password': 15,
            'api_key': 20,
            'private_key': 20,
            'credentials': 20,
            '.pem': 20,
            '.key': 15,
            'token': 10,
            'aws_access_key': 25,
            'database_url': 15
        }
        
        for path in paths:
            for pattern, penalty in dangerous_patterns.items():
                if pattern in path:
                    # .gitignore'da ise sorun yok
                    if 'gitignore' not in path:
                        security_score = max(0, security_score - penalty)
                        security_issues.append(f"Found '{pattern}' in {path}")
                        break
        
        # .gitignore varsa bonus
        if any('.gitignore' in p for p in paths):
            security_score = min(100, security_score + 5)
        
        # 2. TEST COVERAGE (0-100)
        test_files = 0
        code_files = 0
        
        code_extensions = {'.py', '.js', '.jsx', '.ts', '.tsx', '.java', '.go', '.rb', '.php', '.cs', '.cpp', '.c', '.h'}
        test_patterns = ['test_', '_test.', '.test.', '.spec.', '/test/', '/tests/', '__test__']
        
        for path in paths:
            ext = os.path.splitext(path)[1].lower()
            if ext in code_extensions:
                code_files += 1
                # Test dosyası mı?
                if any(pattern in path for pattern in test_patterns):
                    test_files += 1
        
        # Test coverage hesapla
        if code_files > 0:
            test_ratio = test_files / code_files
            test_coverage = min(100, int(test_ratio * 200))  # %50 test = 100 puan
        else:
            test_coverage = 0
        
        # Test klasörü varsa bonus
        if any('/test/' in p or '/tests/' in p for p in paths):
            test_coverage = min(100, test_coverage + 10)
        
        # 3. READABILITY SCORE (0-100)
        readability_grade = 40  # Base score
        
        # README check
        if any('readme.md' in p for p in paths):
            readability_grade += 25
        if any('readme.rst' in p or 'readme.txt' in p for p in paths):
            readability_grade += 15
        
        # Documentation check
        if any('contributing.md' in p for p in paths):
            readability_grade += 10
        
        if any('/docs/' in p or '/doc/' in p for p in paths):
            readability_grade += 15
        
        # LICENSE check
        if any('license' in p for p in paths):
            readability_grade += 5
        
        # Code comments (proxy: README + docs = iyi comment kültürü)
        readability_grade = min(100, readability_grade)
        
        # Cache key (gerçek metriklerle)
        cache_key = f"{repo_param}:{branch_param}:{security_score}:{test_coverage}:{readability_grade}"
        
        # Check cache first (kota tasarrufu!)
        current_time = time.time()
        if cache_key in health_narrative_cache:
            cached_data = health_narrative_cache[cache_key]
            if current_time - cached_data['timestamp'] < CACHE_TTL:
                print(f"✓ Cache HIT for {repo_param} (Quota saved!)")
                return jsonify({
                    'metrics': {
                        'security': security_score,
                        'test_coverage': test_coverage,
                        'readability': readability_grade
                    },
                    'narrative': cached_data['narrative']
                })
        
        print(f"⚡ Cache MISS for {repo_param} - calling AI...")
        print(f"📊 Real Analysis: Security={security_score}, Tests={test_coverage}, Readability={readability_grade}")
        print(f"   Files: {code_files} code, {test_files} test ({total_files} total)")

        # Static fallback narratives (kota aşımında kullanılacak)
        repo_name = repo_param.split('/')[-1]
        static_narratives = [
            f"Neural scan complete for {repo_name}. Security protocols operational. Test coverage needs enhancement. Code architecture shows solid foundation.",
            f"System diagnostics initialized. Repository {repo_name} shows moderate stability. Recommend increasing test coverage for mission-critical components.",
            f"Cyberdeck analysis complete. {repo_name} codebase functional but requires defensive programming enhancements. Deploy additional test frameworks.",
            f"Network scan of {repo_name} repository complete. Security: nominal. Coverage: suboptimal. Readability: acceptable. Proceed with caution, Netrunner.",
            f"OMNI-NET diagnostic: {repo_name} shows balanced architecture. Security measures holding. Test suite expansion recommended for zero-day protection."
        ]

        # AI Prompt - Gerçek analiz sonuçlarıyla
        analysis_context = f"""Repository: {repo_name}
Real Analysis Results:
- Total Files: {total_files} ({code_files} code files, {test_files} test files)
- Security Issues Found: {len(security_issues)} ({', '.join(security_issues[:3]) if security_issues else 'None detected'})
- Test Coverage Ratio: {test_files}/{code_files} = {(test_files/code_files*100) if code_files > 0 else 0:.1f}%
- Documentation: {'README found ✓' if any('readme' in p for p in paths) else 'No README ✗'}
"""

        prompt = f"""You are 'OMNI', an AI Game Master monitoring a Cyberpunk megacorporation's codebase.

{analysis_context}

Final Health Metrics:
- Security Score: {security_score}/100 {'(CRITICAL!)' if security_score < 60 else '(Good)' if security_score > 80 else '(Warning)'}
- Test Coverage: {test_coverage}/100 {'(CRITICAL!)' if test_coverage < 40 else '(Good)' if test_coverage > 70 else '(Warning)'}
- Readability: {readability_grade}/100 {'(Needs work)' if readability_grade < 60 else '(Excellent)' if readability_grade > 80 else '(Acceptable)'}

Generate a short, flavorful, Cyberpunk-style narrative report (max 3-4 sentences) based on REAL analysis data above.
Comment on the actual findings like a mission briefing. Use neon/hacker aesthetics in your tone.

Example: "Initializing neural scan... Security protocols are holding, but test coverage is critically low. We are exposed to rogue zero-days. Enhance the firewall modules immediately, Netrunner."
"""

        narrative = None
        # Model fallback chain: GPT-4o-mini -> Gemini 2.5 Flash Lite
        model_chain = [
            {'type': 'openai', 'name': 'gpt-4o-mini'},
            {'type': 'gemini', 'name': 'gemini-1.5-flash'},
            {'type': 'gemini', 'name': 'gemini-2.5-flash-lite-preview-02-05'},
            {'type': 'gemini', 'name': 'gemini-3.1-flash-lite-preview'},
            {'type': 'gemini', 'name': 'gemini-2.0-flash'},
        ]
        
        for model_info in model_chain:
            try:
                model_type = model_info['type']
                model_name = model_info['name']
                
                if model_type == 'gemini':
                    full_m_name = model_name if model_name.startswith('models/') else f"models/{model_name}"
                    print(f"Trying Health narrative with: {full_m_name}")
                    model = genai.GenerativeModel(full_m_name)
                    response = model.generate_content(prompt, request_options={"timeout": 30})
                    if response and response.text:
                        narrative = response.text.replace('*', '').strip()
                        # Cache success
                        health_narrative_cache[cache_key] = {
                            'narrative': narrative,
                            'timestamp': current_time
                        }
                        print(f"✓ AI response cached for {repo_param} (model: {model_name})")
                        break
                        
                elif model_type == 'openai' and openai_client:
                    print(f"Trying Health narrative with: OpenAI {model_name}")
                    response = openai_client.chat.completions.create(
                        model=model_name,
                        messages=[
                            {"role": "system", "content": "You are OMNI, a cyberpunk AI monitoring codebases."},
                            {"role": "user", "content": prompt}
                        ],
                        max_tokens=150,
                        temperature=0.8
                    )
                    if response.choices and response.choices[0].message.content:
                        narrative = response.choices[0].message.content.strip()
                        # Cache success
                        health_narrative_cache[cache_key] = {
                            'narrative': narrative,
                            'timestamp': current_time
                        }
                        print(f"✓ AI response cached for {repo_param} (model: {model_name})")
                        break
                        
            except Exception as e:
                error_msg = str(e)
                print(f"Health trial with {model_info['name']} failed: {error_msg}")
                # Quota aşımı kontrolü
                if "429" in error_msg or "quota" in error_msg.lower() or "RESOURCE_EXHAUSTED" in error_msg:
                    print("⚠️ Quota exceeded - trying next model")
                    continue
                continue
        
        # Fallback to smart static narrative (AI çalışmazsa)
        if not narrative:
            # Metrik bazlı akıllı fallback
            if security_score < 60:
                sec_status = "CRITICAL SECURITY BREACH DETECTED"
            elif security_score < 80:
                sec_status = "Security protocols need reinforcement"
            else:
                sec_status = "Security firewalls operational"
            
            if test_coverage < 40:
                test_status = "Test coverage dangerously low. Zero-day vulnerabilities imminent"
            elif test_coverage < 70:
                test_status = "Test coverage suboptimal. Deploy additional quality assurance"
            else:
                test_status = "Test infrastructure solid"
            
            if readability_grade < 60:
                doc_status = "Documentation incomplete. Code maintainability at risk"
            elif readability_grade < 80:
                doc_status = "Documentation acceptable but could be enhanced"
            else:
                doc_status = "Excellent documentation standards maintained"
            
            narrative = f"Neural scan of {repo_name} complete. {sec_status}. {test_status}. {doc_status}. Netrunner, proceed with tactical awareness."

        return jsonify({
            'metrics': {
                'security': security_score,
                'test_coverage': test_coverage,
                'readability': readability_grade
            },
            'narrative': narrative
        })
        
    except Exception as e:
        print(f"❌ CRITICAL ERROR in get_code_health: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f"Failed to fetch health metrics: {str(e)}"}), 500

@app.route('/api/history/<int:history_id>/similar', methods=['GET'])
def get_similar_questions(history_id: int):
    """Benzer topluluk sorularını bul ve döndür."""
    # Mevcut soruyu al
    current_history = History.query.get_or_404(history_id)
    current_question = (current_history.user_question or '').lower().strip()
    
    if not current_question:
        return jsonify({'similar': []})
    
    # Anahtar kelimeleri çıkar (stop words hariç)
    stop_words = {'bir', 'bu', 've', 'de', 'da', 'ile', 'için', 'mi', 'mı', 'mu', 'mü',
                  'ne', 'nasıl', 'neden', 'kim', 'nerede', 'hangi', 'kaç', 'ben', 'sen',
                  'o', 'biz', 'siz', 'onlar', 'var', 'yok', 'olarak', 'gibi', 'daha',
                  'çok', 'en', 'az', 'olan', 'the', 'a', 'an', 'is', 'are', 'in', 'on',
                  'to', 'for', 'of', 'and', 'or', 'how', 'what', 'why', 'where', 'when',
                  'bana', 'sana', 'ona', 'bize', 'size', 'onlara', 'beni', 'seni', 'onu', 
                  'bizi', 'sizi', 'onları', 'benim', 'senin', 'onun', 'bizim', 'sizin', 'onların',
                  'yap', 'et', 'iste', 'soru', 'cevap', 'çözüm', 'bunu', 'şunu', 'böyle', 'şöyle'}
    
    # Kelimeleri ayır ve filtrele
    words = [w for w in current_question.split() if len(w) > 2 and w not in stop_words]
    
    if not words:
        return jsonify({'similar': []})
    
    # Sadece Community postlarından ara (mevcut soru hariç)
    community_posts = History.query.filter(
        History.selected_model == 'Community',
        History.id != history_id
    ).order_by(History.timestamp.desc()).limit(100).all()
    
    similar = []
    for post in community_posts:
        post_question = (post.user_question or '').lower()
        # Kaç anahtar kelime eşleşiyor?
        match_count = sum(1 for word in words if word in post_question)
        if match_count > 0:
            similar.append({
                'post': serialize_history(post),
                'match_score': match_count
            })
    
    # En çok eşleşen 5 tanesini döndür
    similar.sort(key=lambda x: x['match_score'], reverse=True)
    top_similar = [item['post'] for item in similar[:5]]
    
    return jsonify({'similar': top_similar})


@app.route('/api/history/<int:history_id>/answers', methods=['POST'])
@jwt_required()
def create_answer(history_id: int):
    try:
        history = History.query.get_or_404(history_id)
        # Use silent=True to avoid 400 error if content-type is multipart/form-data
        data = request.get_json(silent=True) or {}
        user = get_current_user()
        
        # Handle both JSON and Form Data
        if request.content_type and 'multipart/form-data' in request.content_type:
            body = (request.form.get('body') or '').strip()
            code_snippet = request.form.get('code_snippet', '')
        else:
            # Fallback to JSON
            body = (data.get('body') or '').strip()
            code_snippet = data.get('code_snippet', '')

        if not body:
            return jsonify({'error': 'Answer text cannot be empty'}), 400

        image_path = None
        if request.files and 'image' in request.files:
            image_file = request.files['image']
            if image_file and image_file.filename:
                try:
                    filename = secure_filename(f"ans_{int(time.time())}_{image_file.filename}")
                    image_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    image_file.save(image_path)
                except Exception as e:
                    print(f"File upload error: {e}")
                    return jsonify({'error': f'Failed to upload file: {str(e)}'}), 500

        answer = Answer(
            history_id=history_id,
            author_id=user.id,
            author=user.display_name,
            body=body,
            code_snippet=code_snippet,
            image_path=image_path
        )
        db.session.add(answer)
        
        # Bildirim oluştur (gönderi sahibine)
        if history.conversation and history.conversation.user_id and history.conversation.user_id != user.id:
            try:
                notification = Notification(
                    user_id=history.conversation.user_id,
                    type='comment',
                    message=f'{user.display_name} gönderinize çözüm ekledi',
                    related_user_id=user.id,
                    related_post_id=history_id
                )
                db.session.add(notification)
            except Exception as ne:
                print(f"Notification error (ignored): {ne}")
        
        
        # Çözüm paylaşan kişiye XP ver
        award_xp(user.id, XP_REWARDS['share_solution'], "Sharing a Solution", source='share_solution')
        
        db.session.commit()
        return jsonify({'answer': serialize_answer(answer)}), 201
        
    except Exception as e:
        db.session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'An error occurred: {str(e)}'}), 500





@app.route('/api/answers/<int:answer_id>', methods=['DELETE'])
@jwt_required()
def delete_answer(answer_id: int):
    answer = Answer.query.get_or_404(answer_id)
    user = get_current_user()
    if user.id != answer.author_id and not user.is_admin:
        return jsonify({'error': 'You do not have permission to delete this answer.'}), 403
    db.session.delete(answer)
    db.session.commit()
    return jsonify({'status': 'deleted'})


@app.route('/api/community/posts/<int:history_id>', methods=['GET'])
def get_community_post(history_id):
    # Public endpoint for community posts
    item = History.query.get_or_404(history_id)
    if item.selected_model != 'Community':
        return jsonify({'error': 'This is not a community post'}), 404
    
    return jsonify(serialize_history(item))


@app.route('/api/community/posts', methods=['POST'])
@jwt_required()
def create_community_post():
    user = get_current_user()
    
    # Handle multipart/form-data
    if request.content_type and 'multipart/form-data' in request.content_type:
        title = (request.form.get('title') or '').strip()
        code = request.form.get('code', '')
        solution = (request.form.get('solution') or '').strip()
        
        image_path = None
        if request.files and 'image' in request.files:
            image_file = request.files['image']
            if image_file and image_file.filename:
                filename = secure_filename(f"comm_{int(time.time())}_{image_file.filename}")
                image_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                image_file.save(image_path)
    else:
        # JSON fallback
        data = request.json or {}
        title = (data.get('title') or '').strip()
        code = data.get('code', '')
        solution = (data.get('solution') or '').strip()
        image_path = None

    if not title:
        return jsonify({'error': 'Title/Question required.'}), 400

    # 1. Create Conversation
    conversation = Conversation(user_id=user.id, title=title[:50], source='web')
    db.session.add(conversation)
    db.session.commit()

    # 2. Create History Item (The "Question")
    history = History(
        conversation_id=conversation.id,
        user_question=title,
        code_snippet=code,
        ai_response="This is a community post. You can check the solutions below.",
        selected_model='Community',
        summary=title[:50],
        image_path=image_path
    )
    db.session.add(history)
    db.session.commit()

    # 3. Create Answer (Optional "Solution")
    if solution:
        answer = Answer(
            history_id=history.id,
            author_id=user.id,
            author=user.display_name,
            body=solution,
            code_snippet=code if not image_path else None # If image is main content, code might be secondary, but logic remains same
        )
        db.session.add(answer)
        db.session.commit()

    # Toplulukta paylaşım yapan kişiye XP ver
    award_xp(user.id, XP_REWARDS['community_post'], "Creating a Community Post", source='community_post')

    return jsonify({
        'status': 'success',
        'history_id': history.id,
        'conversation_id': conversation.id
    }), 201


@app.route('/api/community/feed', methods=['GET'])
def get_community_feed():
    # Sadece 'Community' olarak işaretlenmiş, silinmemiş postları getir
    items = History.query.filter_by(selected_model='Community', is_deleted=False)\
        .order_by(History.timestamp.desc())\
        .limit(50)\
        .all()
    
    # Kullanıcı giriş yapmış mı kontrol et
    user_id = None
    followed_ids = set()
    liked_post_ids = set()
    try:
        verify_jwt_in_request(optional=True)
        identity = get_jwt_identity()
        if identity:
            user_id = int(identity)
            # Takip edilenleri al
            followed_ids = {f.following_id for f in UserFollow.query.filter_by(follower_id=user_id).all()}
            # Beğenilen postları al
            liked_post_ids = {l.history_id for l in PostLike.query.filter_by(user_id=user_id).all()}
    except Exception:
        pass

    feed_data = []
    for h in items:
        data = serialize_history(h)
        # Takip durumu ekle
        if data['author_id'] and user_id and data['author_id'] != user_id:
             data['is_following'] = data['author_id'] in followed_ids
        else:
             data['is_following'] = False
        # Beğeni durumu ekle
        data['user_has_liked'] = h.id in liked_post_ids
        feed_data.append(data)

    return jsonify({'feed': feed_data})


@app.route('/uploads/<filename>')
@app.route('/api/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)


# ========== SNIPPET ENDPOINTS ==========

@app.route('/api/snippets', methods=['GET'])
@jwt_required()
def list_snippets():
    """List user's saved code snippets."""
    user = get_current_user()
    snippets = Snippet.query.filter_by(user_id=user.id)\
        .order_by(Snippet.created_at.desc()).all()
    
    return jsonify({'snippets': [{
        'id': s.id,
        'title': s.title,
        'code': s.code,
        'language': s.language,
        'created_at': s.created_at.strftime('%Y-%m-%d %H:%M')
    } for s in snippets]})


@app.route('/api/snippets', methods=['POST'])
@jwt_required()
def create_snippet():
    """Save new code snippet."""
    user = get_current_user()
    data = request.json or {}
    
    title = (data.get('title') or '').strip()
    code = data.get('code') or ''
    language = data.get('language') or 'plaintext'
    
    if not title or not code:
        return jsonify({'error': 'Title and code required.'}), 400
    
    snippet = Snippet(
        user_id=user.id,
        title=title,
        code=code,
        language=language
    )
    db.session.add(snippet)
    db.session.commit()
    
    return jsonify({
        'status': 'success',
        'snippet': {
            'id': snippet.id,
            'title': snippet.title,
            'code': snippet.code,
            'language': snippet.language,
            'created_at': snippet.created_at.strftime('%Y-%m-%d %H:%M')
        }
    }), 201


@app.route('/api/snippets/<int:snippet_id>', methods=['DELETE'])
@jwt_required()
def delete_snippet(snippet_id):
    """Delete code snippet."""
    user = get_current_user()
    snippet = Snippet.query.filter_by(id=snippet_id, user_id=user.id).first_or_404()
    
    db.session.delete(snippet)
    db.session.commit()
    
    return jsonify({'status': 'deleted'})


# ==========================================
# KULLANICI TAKİP SİSTEMİ
# ==========================================

@app.route('/api/users/<int:user_id>/follow', methods=['POST'])
@jwt_required()
def follow_user(user_id):
    """Follow user."""
    current_user = get_current_user()
    
    # Kendini takip edemez
    if current_user.id == user_id:
        return jsonify({'error': 'You cannot follow yourself'}), 400
    
    # Takip edilecek kullanıcı var mı?
    target_user = db.session.get(User, user_id)
    if not target_user:
        return jsonify({'error': 'User not found'}), 404
    
    # Zaten takip ediyor mu?
    existing_follow = UserFollow.query.filter_by(
        follower_id=current_user.id, 
        following_id=user_id
    ).first()
    
    if existing_follow:
        return jsonify({'error': 'You are already following this user'}), 400
    
    # Takip et
    follow = UserFollow(follower_id=current_user.id, following_id=user_id)
    db.session.add(follow)
    
    # Bildirim oluştur (Eğer daha önce benzer bir bildirim yoksa)
    existing_notification = Notification.query.filter_by(
        user_id=user_id,
        type='follow',
        related_user_id=current_user.id
    ).first()

    if not existing_notification:
        notification = Notification(
            user_id=user_id,  # Takip edilen kullanıcıya bildirim
            type='follow',
            message=f'{current_user.display_name} started following you',
            related_user_id=current_user.id
        )
        db.session.add(notification)
    
    db.session.commit()
    
    return jsonify({
        'status': 'followed',
        'message': f'{target_user.display_name} is being followed'
    })


@app.route('/api/users/<int:user_id>/follow', methods=['DELETE'])
@jwt_required()
def unfollow_user(user_id):
    """Unfollow user."""
    current_user = get_current_user()
    
    follow = UserFollow.query.filter_by(
        follower_id=current_user.id, 
        following_id=user_id
    ).first()
    
    if not follow:
        return jsonify({'error': 'You are not following this user'}), 400
    
    db.session.delete(follow)
    db.session.commit()
    
    return jsonify({'status': 'unfollowed'})


@app.route('/api/users/<int:user_id>/followers', methods=['GET'])
def get_followers(user_id):
    """List user's followers."""
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({'error': 'User not found'}), 404
    
    followers = UserFollow.query.filter_by(following_id=user_id).all()
    
    return jsonify({
        'followers': [
            {
                'id': f.follower.id,
                'display_name': f.follower.display_name,
                'profile_image': f"/uploads/{os.path.basename(f.follower.profile_image)}" if f.follower.profile_image else None,
                'followed_at': f.created_at.strftime('%Y-%m-%d %H:%M')
            }
            for f in followers
        ],
        'count': len(followers)
    })


@app.route('/api/users/<int:user_id>/following', methods=['GET'])
def get_following(user_id):
    """List user's following."""
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({'error': 'User not found'}), 404
    
    following = UserFollow.query.filter_by(follower_id=user_id).all()
    
    return jsonify({
        'following': [
            {
                'id': f.following.id,
                'display_name': f.following.display_name,
                'profile_image': f"/uploads/{os.path.basename(f.following.profile_image)}" if f.following.profile_image else None,
                'followed_at': f.created_at.strftime('%Y-%m-%d %H:%M')
            }
            for f in following
        ],
        'count': len(following)
    })


@app.route('/api/users/<int:user_id>/profile', methods=['GET'])
def get_user_profile(user_id):
    """Get user profile info."""
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({'error': 'User not found'}), 404
    
    # Takipçi ve takip sayıları
    followers_count = UserFollow.query.filter_by(following_id=user_id).count()
    following_count = UserFollow.query.filter_by(follower_id=user_id).count()
    
    # Mevcut kullanıcı takip ediyor mu?
    is_following = False
    current_user = get_current_user()
    if current_user:
        is_following = UserFollow.query.filter_by(
            follower_id=current_user.id, 
            following_id=user_id
        ).first() is not None
    
    # Kullanıcının gönderileri (Sadece Community postları)
    posts = History.query.join(Conversation).filter(
        Conversation.user_id == user_id,
        History.selected_model == 'Community'  # Sadece topluluk gönderileri
    ).order_by(History.timestamp.desc()).limit(10).all()
    
    # Serialize user with proper profile_image path
    user_data = serialize_user(user)
    user_data.update({
        'followers_count': followers_count,
        'following_count': following_count,
        'is_following': is_following
    })
    
    return jsonify({
        'user': user_data,
        'posts': [serialize_history(h) for h in posts]
    })


@app.route('/api/feed/following', methods=['GET'])
@jwt_required()
def get_following_feed():
    """Get following users' feed."""
    current_user = get_current_user()
    
    # Takip edilen kullanıcı ID'leri
    following_ids = [f.following_id for f in UserFollow.query.filter_by(follower_id=current_user.id).all()]
    
    if not following_ids:
        return jsonify({'feed': [], 'message': 'You are not following anyone yet'})
    
    # Takip edilenlerin gönderileri
    # Sadece 'Community' olarak işaretlenmiş (paylaşılmış) gönderileri getir
    posts = History.query.join(Conversation).filter(
        Conversation.user_id.in_(following_ids)
    ).filter(
        History.selected_model == 'Community'
    ).order_by(History.timestamp.desc()).limit(50).all()
    
    feed_data = []
    for h in posts:
        if h.conversation and h.conversation.user:
            item_data = serialize_history(h)
            is_liked = False
            if current_user:
                like_check = PostLike.query.filter_by(user_id=current_user.id, history_id=h.id).first()
                if like_check: is_liked = True
            item_data['is_liked'] = is_liked
            item_data['author'] = {
                'id': h.conversation.user.id,
                'display_name': h.conversation.user.display_name,
                'profile_image': f"/uploads/{os.path.basename(h.conversation.user.profile_image)}" if h.conversation.user.profile_image else None
            }
            feed_data.append(item_data)
    return jsonify({'feed': feed_data})

    # Legacy code (unreachable)
    return jsonify({
        'feed': [
            {
                **serialize_history(h),
                'author': {
                    'id': h.conversation.user.id if h.conversation.user else None,
                    'display_name': h.conversation.user.display_name if h.conversation.user else 'Anonymous',
                    'profile_image': f"/uploads/{os.path.basename(h.conversation.user.profile_image)}" if h.conversation.user and h.conversation.user.profile_image else None
                }
            }
            for h in posts if h.conversation and h.conversation.user
        ]
    })


@app.route('/api/notifications/all', methods=['GET'])
@jwt_required()
def get_all_notifications():
    """Kullanıcının tüm bildirimlerini getir (active + non-deleted)."""
    current_user = get_current_user()
    
    # Yeni format bildirimler (Notification tablosu) - Lifecycle Aware
    notifications = Notification.query.filter_by(
        user_id=current_user.id,
        is_deleted=False,
        lifecycle_state='active'
    ).order_by(Notification.created_at.desc())\
     .limit(50)\
     .all()
    
    return jsonify({
        'notifications': [
            {
                'id': n.id,
                'type': n.type,
                'message': n.message,
                'is_read': n.is_read,
                'related_user_id': n.related_user_id,
                'related_post_id': n.related_post_id,
                'created_at': n.created_at.strftime('%Y-%m-%d %H:%M')
            }
            for n in notifications
        ]
    })


@app.route('/api/notifications/<int:notification_id>/read', methods=['POST'])
@jwt_required()
def mark_single_notification_read(notification_id):
    """Bildirimi okundu olarak işaretle."""
    current_user = get_current_user()
    
    notification = Notification.query.filter_by(
        id=notification_id, 
        user_id=current_user.id
    ).first()
    
    if not notification:
        return jsonify({'error': 'Notification not found'}), 404
    
    notification.is_read = True
    db.session.commit()
    
    return jsonify({'status': 'read'})


@app.route('/api/debug/init-db', methods=['POST'])
def init_db():
    try:
        from models import db
        with app.app_context():
            db.create_all()
        return jsonify({'status': 'Database initialized'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/debug/static', methods=['GET'])
def debug_static():
    import glob
    try:
        files = glob.glob(os.path.join(app.static_folder, '**', '*'), recursive=True)
        return jsonify({
            'static_folder': app.static_folder,
            'exists': os.path.exists(app.static_folder),
            'files': [f.replace(app.static_folder, '') for f in files]
        })
    except Exception as e:
        return jsonify({'error': str(e)})



# ============================================
# FAVORITES API
# ============================================

@app.route('/api/favorites', methods=['GET'])
@jwt_required()
def get_favorites():
    """Kullanıcının favori yanıtlarını getir"""
    identity = get_jwt_identity()
    current_user = db.session.get(User, int(identity))
    if not current_user:
        return jsonify({'error': 'User not found'}), 404
    
    # Sadece silinmemiş favorileri getir
    favorites = Favorite.query.filter_by(user_id=current_user.id, is_deleted=False)\
        .order_by(Favorite.created_at.desc()).all()
    
    result = []
    for fav in favorites:
        # Bağlı History kaydı silinmemiş olmalı
        history = History.query.filter_by(id=fav.history_id, is_deleted=False).first()
        if history:
            conversation = db.session.get(Conversation, history.conversation_id)
            result.append({
                'id': fav.id,
                'history_id': history.id,
                'user_question': history.user_question,
                'ai_response': history.ai_response,
                'code_snippet': history.code_snippet,
                'model': history.selected_model,
                'conversation_title': conversation.title if conversation else None,
                'created_at': fav.created_at.isoformat()
            })
    
    return jsonify(result)


@app.route('/api/favorites/<int:history_id>', methods=['POST'])
@jwt_required()
def add_favorite(history_id):
    """Yanıtı favorilere ekle"""
    identity = get_jwt_identity()
    current_user = db.session.get(User, int(identity))
    
    if not current_user:
        return jsonify({'error': f'User not found for identity: {identity}'}), 404
    
    history = db.session.get(History, history_id)
    if not history:
        return jsonify({'error': 'History not found'}), 404
    
    # Zaten favoride mi kontrol et
    existing = Favorite.query.filter_by(user_id=current_user.id, history_id=history_id).first()
    if existing:
        return jsonify({'status': 'added', 'id': existing.id, 'message': 'Already in favorites'}), 200
    
    favorite = Favorite(user_id=current_user.id, history_id=history_id)
    db.session.add(favorite)
    db.session.commit()
    
    return jsonify({'status': 'added', 'id': favorite.id})


@app.route('/api/favorites/<int:history_id>', methods=['DELETE'])
@jwt_required()
def remove_favorite(history_id):
    """Yanıtı favorilerden kaldır"""
    identity = get_jwt_identity()
    current_user = db.session.get(User, int(identity))

    if not current_user:
        return jsonify({'error': 'User not found'}), 404
    
    favorite = Favorite.query.filter_by(user_id=current_user.id, history_id=history_id).first()
    if not favorite:
        return jsonify({'status': 'removed', 'message': 'Not in favorites'}), 200
    
    db.session.delete(favorite)
    db.session.commit()
    
    return jsonify({'status': 'removed'})


@app.route('/api/favorites/check/<int:history_id>', methods=['GET'])
@jwt_required()
def check_favorite(history_id):
    """Yanıtın favoride olup olmadığını kontrol et"""
    current_user = User.query.filter_by(email=get_jwt_identity()).first()
    if not current_user:
        return jsonify({'error': 'User not found'}), 404
    
    favorite = Favorite.query.filter_by(user_id=current_user.id, history_id=history_id).first()
    return jsonify({'is_favorite': favorite is not None})

@app.route('/api/files/<filename>')
def serve_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)



def call_gemini_with_retry(prompt, model_name='gemini-2.5-flash', max_retries=2):
    """Calls Gemini API with exponential backoff for 429 errors."""
    for i in range(max_retries):
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(prompt)
            if response and response.text:
                return response.text
            return None
        except Exception as e:
            err_msg = str(e)
            if "429" in err_msg or "quota" in err_msg.lower():
                if i < max_retries - 1:
                    wait_time = (i + 1) * 3
                    print(f"Gemini 429 received. Waiting {wait_time}s before retry {i+1}...")
                    time.sleep(wait_time)
                    continue
                else:
                    print("Gemini quota exceeded after retries. Triggering fallback.")
                    return None # Explicitly return None to trigger fallback in route
            raise e
    return None

# --- GITHUB FEATURES ROTALARI ---




# --- FEEDBACK API ---

@app.route('/api/feedback', methods=['POST'])
def submit_feedback():
    """AI yanıtına 👍 / 👎 geri bildirimi gönder."""
    from models import Feedback
    data = request.get_json() or {}
    history_id = data.get('history_id')
    rating = data.get('rating')  # +1 veya -1

    print(f"DEBUG: Feedback received - history_id: {history_id}, rating: {rating}")

    if not history_id or rating not in (1, -1):
        return jsonify({'error': 'history_id and rating (+1 or -1) required'}), 400

    # Kullanıcı varsa kaydet, yoksa anonim (user_id=None)
    user_id = None
    try:
        verify_jwt_in_request(optional=True)
        identity = get_jwt_identity()
        if identity:
            user_id = int(identity)
    except Exception as e:
        print(f"DEBUG: JWT Verify failed (optional): {e}")
        pass

    try:
        # Tipi garantiye al (SQLAlchemy query hatalarını önlemek için)
        h_id = int(history_id)
        u_id = user_id

        existing = Feedback.query.filter_by(history_id=h_id, user_id=u_id).first()
        if existing:
            if existing.rating == rating:
                # Aynı oyu tekrar verirse → geri al (toggle)
                # Dislike ise detaylı geri bildirimi de sil
                from models import FeedbackDetail
                FeedbackDetail.query.filter_by(history_id=h_id, user_id=u_id).delete()
                db.session.delete(existing)
                db.session.commit()
                return jsonify({'action': 'removed', 'rating': rating})
            else:
                # Farklı oy → güncelle
                existing.rating = rating
                db.session.commit()
                return jsonify({'action': 'updated', 'rating': rating})
        else:
            fb = Feedback(history_id=h_id, user_id=u_id, rating=rating)
            db.session.add(fb)
            db.session.commit()
            return jsonify({'action': 'added', 'rating': rating})
    except Exception as e:
        print(f"DEBUG: Feedback error: {e}")
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/api/feedback/<int:history_id>', methods=['GET'])
def get_feedback(history_id):
    """Bir mesajın toplam 👍/👎 sayısını ve kullanıcının oyunu döndürür."""
    from models import Feedback
    thumbs_up = Feedback.query.filter_by(history_id=history_id, rating=1).count()
    thumbs_down = Feedback.query.filter_by(history_id=history_id, rating=-1).count()

    user_rating = None
    try:
        verify_jwt_in_request(optional=True)
        identity = get_jwt_identity()
        if identity:
            fb = Feedback.query.filter_by(history_id=history_id, user_id=int(identity)).first()
            if fb:
                user_rating = fb.rating
    except Exception:
        pass

    return jsonify({'thumbs_up': thumbs_up, 'thumbs_down': thumbs_down, 'user_rating': user_rating})


# --- PROJECT / WORKSPACE API ---

@app.route('/api/projects', methods=['GET'])
@jwt_required()
def list_projects():
    """Kullanıcının tüm projelerini listele."""
    from models import Project
    user_id = int(get_jwt_identity())
    projects = Project.query.filter_by(user_id=user_id).order_by(Project.updated_at.desc()).all()
    return jsonify({'projects': [
        {
            'id': p.id,
            'name': p.name,
            'description': p.description,
            'file_count': p.files.count(),
            'created_at': p.created_at.isoformat(),
            'updated_at': p.updated_at.isoformat() if p.updated_at else None,
        }
        for p in projects
    ]})


@app.route('/api/projects', methods=['POST'])
@jwt_required()
def create_project():
    """Yeni proje oluştur."""
    from models import Project
    user_id = int(get_jwt_identity())
    data = request.get_json() or {}
    name = data.get('name', '').strip()
    if not name:
        return jsonify({'error': 'Project name is required'}), 400
    project = Project(user_id=user_id, name=name, description=data.get('description', ''))
    db.session.add(project)
    db.session.commit()
    return jsonify({'id': project.id, 'name': project.name}), 201



@app.route('/api/projects/<int:project_id>/files', methods=['GET'])
@jwt_required()
def list_project_files(project_id):
    """Projenin dosyalarını listele."""
    from models import Project, ProjectFile
    user_id = int(get_jwt_identity())
    project = Project.query.filter_by(id=project_id, user_id=user_id).first_or_404()
    files = project.files.order_by(ProjectFile.name).all()
    return jsonify({'files': [
        {'id': f.id, 'name': f.name, 'language': f.language, 'content': f.content,
         'created_at': f.created_at.isoformat()}
        for f in files
    ]})


@app.route('/api/projects/<int:project_id>/files', methods=['POST'])
@jwt_required()
def add_project_file(project_id):
    """Projeye dosya ekle."""
    from models import Project, ProjectFile
    user_id = int(get_jwt_identity())
    project = Project.query.filter_by(id=project_id, user_id=user_id).first_or_404()
    data = request.get_json() or {}
    name = str(data.get('name', '') or '').replace('\x00', '').strip()
    if not name:
        return jsonify({'error': 'File name is required'}), 400

    raw_content = data.get('content', '')
    if raw_content is None:
        raw_content = ''
    if not isinstance(raw_content, str):
        raw_content = str(raw_content)

    encoding = str(data.get('encoding', 'text') or 'text').strip().lower()
    mime_type = str(data.get('mime_type', '') or '').strip().lower()
    language = str(data.get('language', 'plaintext') or 'plaintext').replace('\x00', '')

    if encoding == 'base64':
        try:
            binary_data = base64.b64decode(raw_content, validate=True)
        except Exception:
            return jsonify({'error': 'Invalid base64 payload'}), 400

        lang_lc = language.lower()
        name_lc = name.lower()
        is_pdf = name_lc.endswith('.pdf') or mime_type == 'application/pdf' or lang_lc == 'pdf'
        is_docx = name_lc.endswith('.docx') or mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or lang_lc == 'docx'

        if is_pdf:
            try:
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(binary_data))
                parts = []
                total_chars = 0
                for page in reader.pages:
                    txt = page.extract_text() or ''
                    if not txt:
                        continue
                    remaining = 20000 - total_chars
                    if remaining <= 0:
                        break
                    segment = txt[:remaining]
                    parts.append(segment)
                    total_chars += len(segment)
                content = '\n'.join(parts).replace('\x00', '')
            except Exception:
                return jsonify({'error': 'PDF text extraction failed'}), 400

            if not content.strip():
                return jsonify({'error': 'No readable text found in PDF'}), 400
        elif is_docx:
            try:
                from docx import Document
                doc = Document(io.BytesIO(binary_data))
                content = '\n'.join(p.text for p in doc.paragraphs if p.text).replace('\x00', '')
                if len(content) > 20000:
                    content = content[:20000]
            except Exception:
                return jsonify({'error': 'DOCX text extraction failed'}), 400

            if not content.strip():
                return jsonify({'error': 'No readable text found in DOCX'}), 400
        else:
            return jsonify({'error': 'Only PDF and DOCX are supported for base64 uploads'}), 400

        language = 'plaintext'
    else:
        # DB driver rejects NUL chars in SQL string params.
        nul_count = raw_content.count('\x00')
        content = raw_content.replace('\x00', '')

        if nul_count > 0 and not content.strip():
            return jsonify({'error': 'Binary file content is not supported. Please upload text content.'}), 400

    pf = ProjectFile(
        project_id=project.id,
        name=name,
        content=content,
        language=language
    )
    db.session.add(pf)
    db.session.commit()
    return jsonify({'id': pf.id, 'name': pf.name}), 201


@app.route('/api/projects/<int:project_id>/files/<int:file_id>', methods=['DELETE'])
@jwt_required()
def delete_project_file(project_id, file_id):
    """Proje dosyasını sil."""
    from models import Project, ProjectFile
    user_id = int(get_jwt_identity())
    project = Project.query.filter_by(id=project_id, user_id=user_id).first_or_404()
    pf = ProjectFile.query.filter_by(id=file_id, project_id=project.id).first_or_404()
    db.session.delete(pf)
    db.session.commit()
    return jsonify({'message': 'File deleted'})

@app.route('/api/projects/<int:project_id>', methods=['DELETE'])
@jwt_required()
def delete_project(project_id):
    """Projeyi ve ona ait tüm konuşmaları sil."""
    from models import Project, Conversation
    user_id = int(get_jwt_identity())
    project = Project.query.filter_by(id=project_id, user_id=user_id).first_or_404()
    
    # Proje silindiğinde ona ait tüm konuşmaları da sil
    Conversation.query.filter_by(project_id=project_id).delete()
    
    db.session.delete(project)
    db.session.commit()
    return jsonify({'message': 'Project deleted successfully'})

@app.route('/api/feedback/detail', methods=['POST'])
@jwt_required(optional=True)
def submit_feedback_detail():
    """Gelişmiş geri bildirim (dislike sonrası kategori/yorum) gönder."""
    from models import FeedbackDetail
    data = request.json or {}
    history_id = data.get('history_id')
    category = data.get('category')
    comment = data.get('comment', '')
    user_id = get_jwt_identity()

    if not history_id or not category:
        return jsonify({'error': 'history_id and category are required'}), 400

    detail = FeedbackDetail(
        history_id=history_id,
        user_id=user_id,
        category=category,
        comment=comment
    )
    db.session.add(detail)
    db.session.commit()
    return jsonify({'message': 'Feedback detail submitted successfully'}), 201


@app.route('/api/projects/<int:project_id>/context', methods=['GET'])
@jwt_required()
def get_project_context(project_id):
    """Projenin tüm dosyalarını birleşik AI bağlamı olarak döndür."""
    from models import Project, ProjectFile
    user_id = int(get_jwt_identity())
    project = Project.query.filter_by(id=project_id, user_id=user_id).first_or_404()
    files = project.files.order_by(ProjectFile.name).all()

    context_parts = [f"# Project: {project.name}"]
    if project.description:
        context_parts.append(f"Description: {project.description}\n")
    for f in files:
        context_parts.append(f"\n## File: {f.name} ({f.language})\n```{f.language}\n{f.content}\n```")

    return jsonify({'context': '\n'.join(context_parts), 'file_count': len(files)})


@app.route('/api/projects/<int:project_id>/semantic_search', methods=['POST'])
@jwt_required()
def semantic_search_project_context(project_id):
    """Return embedding-ranked project chunks for a query (debug/inspection)."""
    from models import Project

    user_id = int(get_jwt_identity())
    project = Project.query.filter_by(id=project_id, user_id=user_id).first_or_404()

    data = request.get_json(silent=True) or {}
    query = (data.get('query') or data.get('question') or '').strip()
    top_k = data.get('top_k', 6)

    if not query:
        return jsonify({'error': 'query is required'}), 400

    try:
        top_k = int(top_k)
    except Exception:
        top_k = 6
    top_k = max(1, min(20, top_k))

    result = get_project_semantic_hits(project, query, top_k=top_k)
    if not result:
        return jsonify({
            'query': query,
            'project_id': project.id,
            'project_name': project.name,
            'hits': [],
            'message': 'No semantic hits found or embedding unavailable. Falling back to static context in /api/ask.'
        })

    return jsonify({
        'query': query,
        'project_id': project.id,
        'project_name': project.name,
        'query_model': result['query_model'],
        'total_chunks': result['total_chunks'],
        'top_k': top_k,
        'hits': result['hits'],
    })


# ==========================================
# TEMA VE KİŞİSELLEŞTİRME MAĞAZASI (THEME STORE)
# ==========================================

@app.route('/api/themes', methods=['GET'])
@jwt_required()
def get_user_theme():
    """Kullanıcının tercih ettiği temayı ve açılmış temaları getir."""
    from models import UserTheme
    user = get_current_user()
    if not user:
        return jsonify({'message': 'User session cleared or user deleted'}), 200
    
    user_id = user.id
    theme_pref = UserTheme.query.filter_by(user_id=user_id).first()
    
    if not theme_pref:
        # Varsayılan oluştur
        theme_pref = UserTheme(
            user_id=user_id,
            active_theme='dark',
            unlocked_themes=json.dumps(['light', 'dark'])
        )
        db.session.add(theme_pref)
        db.session.commit()
    
    unlocked = []
    try:
        unlocked = json.loads(theme_pref.unlocked_themes)
    except:
        unlocked = ['light', 'dark']
        
    return jsonify({
        'active_theme': theme_pref.active_theme,
        'unlocked_themes': unlocked
    })

@app.route('/api/themes', methods=['POST'])
@jwt_required()
def update_user_theme():
    """Kullanıcının aktif temasını güncelle veya yeni tema satın al."""
    print("DEBUG: update_user_theme POST called")
    from models import UserTheme, User
    user = get_current_user()
    if not user:
        return jsonify({'message': 'User session cleared or user deleted'}), 200
    
    user_id = user.id
    data = request.json or {}
    
    action = data.get('action') # 'set_active' or 'unlock'
    theme_name = data.get('theme')
    
    if not theme_name:
        return jsonify({'error': 'Tema adı gerekli'}), 400
        
    theme_pref = UserTheme.query.filter_by(user_id=user_id).first()
    
    if not theme_pref:
        theme_pref = UserTheme(
            user_id=user_id,
            active_theme='dark',
            unlocked_themes=json.dumps(['light', 'dark'])
        )
        db.session.add(theme_pref)
    
    unlocked = []
    try:
        unlocked = json.loads(theme_pref.unlocked_themes)
    except:
        unlocked = ['light', 'dark']
        
    if action == 'set_active':
        if theme_name not in unlocked:
            return jsonify({'error': 'Bu temanın kilidi henüz açılmamış'}), 403
        
        theme_pref.active_theme = theme_name
        db.session.commit()
        return jsonify({'message': 'Aktif tema güncellendi', 'theme': theme_name})
        
    elif action == 'unlock':
        cost = data.get('cost', 50)  # Varsayılan 50 Coin (XP değil)
        
        if theme_name in unlocked:
            return jsonify({'error': 'Bu tema zaten açık'}), 400
            
        if user.coins < cost:
            return jsonify({'error': f'Yetersiz Coin. ({cost} Coin gerekli)'}), 403
            
        # Coin harcatma
        user.coins -= cost
        unlocked.append(theme_name)
        theme_pref.unlocked_themes = json.dumps(unlocked)
        db.session.commit()
        
        return jsonify({
            'message': f'{theme_name} teması başarıyla açıldı',
            'unlocked_themes': unlocked,
            'remaining_coins': user.coins
        })
        
    return jsonify({'error': 'Geçersiz işlem'}), 400




# ==========================================
# GERÇEK ZAMANLI İŞBİRLİĞİ (COLLABORATION)
# ==========================================

@app.route('/api/collaboration/share', methods=['POST'])
@jwt_required()
def share_conversation():
    """Konuşma için paylaşım linki/token'ı oluştur."""
    from models import SharedSession, Conversation
    user_id = int(get_jwt_identity())
    data = request.json or {}
    conversation_id = data.get('conversation_id')
    
    if not conversation_id:
        return jsonify({'error': 'conversation_id gerekli'}), 400
        
    # Konuşma sahipliği kontrolü
    conv = Conversation.query.filter_by(id=conversation_id, user_id=user_id).first()
    if not conv:
        return jsonify({'error': 'Konuşma bulunamadı veya erişim yetkiniz yok'}), 404
        
    # Mevcut aktif session var mı bak
    existing_session = SharedSession.query.filter_by(conversation_id=conversation_id, is_active=True).first()
    if existing_session:
        return jsonify({
            'share_token': existing_session.share_token,
            'message': 'Mevcut paylaşım linki döndürüldü'
        })
        
    # Yeni token oluştur
    token = str(uuid.uuid4())
    new_session = SharedSession(
        conversation_id=conversation_id,
        owner_id=user_id,
        share_token=token,
        is_active=True
    )
    db.session.add(new_session)
    db.session.commit()
    
    return jsonify({
        'share_token': token,
        'message': 'Yeni paylaşım linki oluşturuldu'
    })

@app.route('/api/collaboration/session/<token>', methods=['GET'])
def get_shared_session(token):
    """Token üzerinden paylaşılan konuşma detaylarını getir."""
    from models import SharedSession, History
    
    session = SharedSession.query.filter_by(share_token=token, is_active=True).first()
    if not session:
        return jsonify({'error': 'Geçersiz veya süresi dolmuş paylaşım linki'}), 404
        
    # Konuşma geçmişini getir
    history = History.query.filter_by(conversation_id=session.conversation_id).order_by(History.timestamp.asc()).all()
    
    history_data = []
    for h in history:
        history_data.append({
            'id': h.id,
            'user_question': h.user_question,
            'ai_response': h.ai_response,
            'code_snippet': h.code_snippet,
            'selected_model': h.selected_model,
            'timestamp': h.timestamp.isoformat(),
            'image_url': f"/uploads/{h.image_path}" if h.image_path else None
        })
        
    return jsonify({
        'conversation_id': session.conversation_id,
        'history': history_data,
        'owner_display_name': session.owner.display_name,
        'title': session.conversation.title or 'Paylaşılan Sohbet'
    })

@app.route('/api/collaboration/session/<token>/send', methods=['POST'])
def send_to_session(token):
    """Live Sync: Soruyu al, AI'dan stream et, her chunk'ı Socket.io ile yayınla."""
    from models import SharedSession, History
    import threading

    session = SharedSession.query.filter_by(share_token=token, is_active=True).first()
    if not session:
        return jsonify({'error': 'Geçersiz paylaşım linki'}), 404

    data = request.json or {}
    question = data.get('question', '').strip()
    model_name = data.get('model', 'gemini-2.5-flash-lite')
    sender_name = data.get('sender_name', 'Guest')
    client_nonce = data.get('client_nonce')

    if not question:
        return jsonify({'error': 'Mesaj boş olamaz'}), 400

    # Önce history kaydı oluştur (eğer auto ise arka planda asıl model ile güncellenecek)
    history_entry = History(
        conversation_id=session.conversation_id,
        user_question=question,
        selected_model=model_name,
        ai_response=''
    )
    db.session.add(history_entry)
    db.session.commit()
    history_id = history_entry.id
    conversation_id = session.conversation_id

    # Oda'ya "soru gönderildi" bildirimi yap
    socketio.emit('collab_question', {
        'question': question,
        'sender': sender_name,
        'history_id': history_id,
        'token': token,
        'client_nonce': client_nonce
    }, room=token)

    def _stream_ai_to_room(app_ctx, history_id, conversation_id, question, model_name, token):
        """Background thread: AI'dan stream, her chunk'ı socket.io ile yayınla."""
        from models import History
        full_response = ''
        try:
            with app_ctx:
                # Arka planda model yönlendirme (Lazy Routing)
                if model_name == 'auto':
                    prefs = {}
                    intent = detect_intent(question, '')
                    detected_lang = language_detector.detect(question, '')
                    model_name, _ = model_router.route(detected_lang, intent, prefs)
                    
                    # Veritabanında (History) seçilen modeli güncelle
                    h_entry = db.session.get(History, history_id)
                    if h_entry:
                        h_entry.selected_model = model_name
                        db.session.commit()

                # Gemini streaming
                if 'gemini' in model_name.lower():
                    try:
                        model_obj = genai.GenerativeModel(model_name)
                        stream = model_obj.generate_content(question, stream=True)
                        for chunk in stream:
                            if chunk.text:
                                full_response += chunk.text
                                socketio.emit('collab_stream_chunk', {
                                    'chunk': chunk.text,
                                    'history_id': history_id,
                                    'token': token
                                }, room=token)
                    except Exception as e:
                        print(f'Gemini stream error in collab: {e}')
                        # Fallback to non-streaming
                        try:
                            model_obj = genai.GenerativeModel('gemini-2.5-flash-lite')
                            resp = model_obj.generate_content(question)
                            full_response = resp.text or 'Yanıt üretilemedi'
                            socketio.emit('collab_stream_chunk', {
                                'chunk': full_response,
                                'history_id': history_id,
                                'token': token
                            }, room=token)
                        except Exception as e2:
                            full_response = f'Hata: {str(e2)}'

                # OpenAI streaming
                elif 'gpt' in model_name.lower():
                    try:
                        openai_client = OpenAI(api_key=os.getenv('OPENAI_API_KEY', ''))
                        stream = openai_client.chat.completions.create(
                            model=model_name,
                            messages=[{'role': 'user', 'content': question}],
                            stream=True
                        )
                        for chunk in stream:
                            delta = chunk.choices[0].delta.content if chunk.choices else None
                            if delta:
                                full_response += delta
                                socketio.emit('collab_stream_chunk', {
                                    'chunk': delta,
                                    'history_id': history_id,
                                    'token': token
                                }, room=token)
                    except Exception as e:
                        full_response = f'OpenAI Hata: {str(e)}'

                # Claude streaming
                elif 'claude' in model_name.lower():
                    try:
                        anthropic_client = Anthropic(api_key=os.getenv('ANTHROPIC_API_KEY', ''))
                        with anthropic_client.messages.stream(
                            model=model_name,
                            max_tokens=4096,
                            messages=[{'role': 'user', 'content': question}]
                        ) as stream:
                            for text_chunk in stream.text_stream:
                                full_response += text_chunk
                                socketio.emit('collab_stream_chunk', {
                                    'chunk': text_chunk,
                                    'history_id': history_id,
                                    'token': token
                                }, room=token)
                    except Exception as e:
                        full_response = f'Claude Hata: {str(e)}'

                else:
                    # Default fallback
                    full_response = 'Model desteklenmiyor. Lütfen Gemini, GPT veya Claude seçin.'

                # DB'ye tam yanıtı kaydet
                h = History.query.get(history_id)
                if h:
                    h.ai_response = full_response
                    db.session.commit()

                # Stream tamamlandı sinyali
                socketio.emit('collab_stream_done', {
                    'history_id': history_id,
                    'full_response': full_response,
                    'token': token
                }, room=token)

        except Exception as ex:
            print(f'Collab stream thread error: {ex}')
            socketio.emit('collab_stream_done', {
                'history_id': history_id,
                'full_response': f'Hata oluştu: {str(ex)}',
                'token': token,
                'error': True
            }, room=token)

    # Background task başlat (eventlet veya threading safe)
    app_ctx = app.app_context()
    socketio.start_background_task(
        _stream_ai_to_room,
        app_ctx, 
        history_id, 
        conversation_id, 
        question, 
        model_name, 
        token
    )

    return jsonify({'status': 'streaming', 'history_id': history_id})


@app.route('/api/collaboration/session/<token>/review', methods=['GET'])
def get_session_review(token):
    """Return current review status and comment thread for a shared session."""
    session = SharedSession.query.filter_by(share_token=token, is_active=True).first()
    if not session:
        return jsonify({'error': 'Geçersiz paylaşım linki'}), 404

    review = CollaborationReview.query.filter_by(session_id=session.id).first()
    comments = CollaborationComment.query.filter_by(session_id=session.id).order_by(CollaborationComment.created_at.desc()).limit(100).all()

    return jsonify({
        'status': review.status if review else 'open',
        'updated_at': review.updated_at.isoformat() if review and review.updated_at else None,
        'updated_by': review.updated_by_name if review else None,
        'comments': [
            {
                'id': c.id,
                'author': c.author_name,
                'comment': c.comment,
                'created_at': c.created_at.isoformat() if c.created_at else None,
            }
            for c in comments
        ]
    })


@app.route('/api/collaboration/session/<token>/review/comment', methods=['POST'])
def add_session_review_comment(token):
    """Add a review comment to shared session thread."""
    session = SharedSession.query.filter_by(share_token=token, is_active=True).first()
    if not session:
        return jsonify({'error': 'Geçersiz paylaşım linki'}), 404

    data = request.json or {}
    comment_text = (data.get('comment') or '').strip()
    if not comment_text:
        return jsonify({'error': 'Yorum boş olamaz'}), 400

    user = get_current_user()
    author_name = user.display_name if user else (data.get('guest_name') or 'Guest Reviewer')

    comment = CollaborationComment(
        session_id=session.id,
        author_user_id=user.id if user else None,
        author_name=author_name,
        comment=comment_text,
    )
    db.session.add(comment)
    db.session.commit()

    return jsonify({
        'message': 'Yorum eklendi',
        'comment': {
            'id': comment.id,
            'author': comment.author_name,
            'comment': comment.comment,
            'created_at': comment.created_at.isoformat() if comment.created_at else None,
        }
    })


@app.route('/api/collaboration/session/<token>/review/status', methods=['POST'])
def set_session_review_status(token):
    """Update review status for shared collaboration session."""
    session = SharedSession.query.filter_by(share_token=token, is_active=True).first()
    if not session:
        return jsonify({'error': 'Geçersiz paylaşım linki'}), 404

    data = request.json or {}
    status = (data.get('status') or '').strip().lower()
    allowed = {'open', 'revision_requested', 'approved'}
    if status not in allowed:
        return jsonify({'error': 'Geçersiz durum'}), 400

    user = get_current_user()
    actor_name = user.display_name if user else (data.get('guest_name') or 'Guest Reviewer')

    review = CollaborationReview.query.filter_by(session_id=session.id).first()
    if not review:
        review = CollaborationReview(session_id=session.id, status=status)
        db.session.add(review)
    review.status = status
    review.updated_by_user_id = user.id if user else None
    review.updated_by_name = actor_name
    review.updated_at = _utcnow()

    db.session.commit()

    return jsonify({
        'message': 'Review durumu güncellendi',
        'status': review.status,
        'updated_by': review.updated_by_name,
        'updated_at': review.updated_at.isoformat() if review.updated_at else None,
    })

@app.route('/api/stats/weekly', methods=['GET'])
@jwt_required()
def get_weekly_stats():
    """Haftalık kullanım istatistiklerini getir."""
    from models import History, User, XPEvent
    from datetime import datetime as dt, timedelta
    
    user_id = int(get_jwt_identity())
    now = _utcnow()
    last_7_days = now - timedelta(days=7)
    prev_7_days = now - timedelta(days=14)
    
    # Bu haftaki sorular
    current_questions = History.query.filter(
        History.conversation.has(user_id=user_id),
        History.timestamp >= last_7_days
    ).all()
    
    # Geçen haftaki sorular
    prev_questions = History.query.filter(
        History.conversation.has(user_id=user_id),
        History.timestamp >= prev_7_days,
        History.timestamp < last_7_days
    ).count()
    
    # Model kullanımı
    model_usage = {}
    for q in current_questions:
        m = q.selected_model or 'unknown'
        model_usage[m] = model_usage.get(m, 0) + 1

    # Gerçek XP event tablosundan haftalık hesap
    xp_breakdown = {
        'asking_question': 0,
        'sharing_solution': 0,
        'creating_community_post': 0,
        'received_likes': 0,
        'daily_login': 0,
        'streak_bonus': 0,
        'daily_login_and_streak': 0,
        'other': 0
    }

    weekly_events = XPEvent.query.filter(
        XPEvent.user_id == user_id,
        XPEvent.created_at >= last_7_days
    ).all()

    real_xp_earned = 0
    for ev in weekly_events:
        ev_amount = ev.amount or 0
        real_xp_earned += ev_amount

        if ev.source == 'ask_question':
            xp_breakdown['asking_question'] += ev_amount
        elif ev.source == 'share_solution':
            xp_breakdown['sharing_solution'] += ev_amount
        elif ev.source == 'community_post':
            xp_breakdown['creating_community_post'] += ev_amount
        elif ev.source in ('received_like_post', 'received_like_answer'):
            xp_breakdown['received_likes'] += ev_amount
        elif ev.source == 'daily_login':
            xp_breakdown['daily_login'] += ev_amount
        elif ev.source == 'streak_bonus':
            xp_breakdown['streak_bonus'] += ev_amount
        else:
            xp_breakdown['other'] += ev_amount

    xp_breakdown['daily_login_and_streak'] = xp_breakdown['daily_login'] + xp_breakdown['streak_bonus']
        
    # Günlük dağılım (Son 7 gün, kronolojik) - tarih bazlı anahtarlar
    daily_by_date = {}
    for i in range(6, -1, -1):
        date_obj = (now - timedelta(days=i)).date()
        daily_by_date[date_obj.isoformat()] = {
            'date': date_obj.isoformat(),
            'label': date_obj.strftime('%a'),
            'count': 0
        }

    for q in current_questions:
        date_key = q.timestamp.date().isoformat()
        if date_key in daily_by_date:
            daily_by_date[date_key]['count'] += 1

    # Geriye dönük uyumluluk için eski formatı da döndür
    daily_stats = {item['label']: item['count'] for item in daily_by_date.values()}
            
    user = User.query.get(user_id)
    total_xp_earned, user_level = resolve_effective_progress(user)
    needs_commit = False
    if (user.total_xp_earned or 0) != total_xp_earned:
        user.total_xp_earned = total_xp_earned
        needs_commit = True
    if (user.level or 1) != user_level:
        user.level = user_level
        needs_commit = True
    if needs_commit:
        db.session.commit()
    
    return jsonify({
        'current_week': {
            'total_questions': len(current_questions),
            'model_usage': model_usage,
            'xp_earned': real_xp_earned,
            'xp_breakdown': xp_breakdown,
            'daily_distribution': daily_stats,
            'daily_points': list(daily_by_date.values())
        },
        'previous_week_total': prev_questions,
        'user_stats': {
            'xp': user.xp,
            'level': user_level,
            'streak': user.streak_days
        }
    })

# ==========================================
# BILLING & TOKEN ECONOMY ENDPOINTS
#💳 Stripe entegrasyon rotaları
# ==========================================

_STRIPE_SECRET_KEY = os.getenv('STRIPE_SECRET_KEY', '')
_STRIPE_WEBHOOK_SECRET = os.getenv('STRIPE_WEBHOOK_SECRET', '')
_STRIPE_PUBLISHABLE_KEY = os.getenv('STRIPE_PUBLISHABLE_KEY', '')

if _STRIPE_SECRET_KEY:
    stripe.api_key = _STRIPE_SECRET_KEY
    print("Stripe client configured.")
else:
    print("Warning: STRIPE_SECRET_KEY not set. Billing disabled.")

#💳 Iyzico entegrasyon rotaları
_IYZICO_API_KEY = os.getenv('IYZICO_API_KEY', '')
_IYZICO_SECRET_KEY = os.getenv('IYZICO_SECRET_KEY', '')
_IYZICO_BASE_URL = os.getenv('IYZICO_BASE_URL', 'https://sandbox-api.iyzipay.com')

def _normalize_iyzico_base_url(value: str) -> str:
    base_url = (value or '').strip().rstrip('/')
    if base_url.startswith('https://'):
        return base_url[len('https://'):]
    if base_url.startswith('http://'):
        return base_url[len('http://'):]
    return base_url

iyzi_options = {
    'api_key': _IYZICO_API_KEY,
    'secret_key': _IYZICO_SECRET_KEY,
    'base_url': _normalize_iyzico_base_url(_IYZICO_BASE_URL)
}

def _find_token_package(package_id):
    """Hem DB ID (int) hem de string slug/name bazlı paket araması yapan yardımcı fonksiyon."""
    if not package_id:
        return None
        
    package_id_str = str(package_id).lower().replace('-', ' ')
    pkg = None
    
    try:
        # 1. DB'den Ara (ID ile)
        if isinstance(package_id, int) or (isinstance(package_id, str) and package_id.isdigit()):
            pkg = TokenPackage.query.filter_by(id=int(package_id), is_active=True).first()
        
        # 2. DB'den Ara (İsim/Slug ile)
        if not pkg:
            pkg = TokenPackage.query.filter(
                (db.func.lower(TokenPackage.name) == package_id_str) |
                (db.func.lower(TokenPackage.name) == str(package_id).lower())
            ).filter_by(is_active=True).first()

        # 3. Fallback: DB'de yoksa DEFAULT_TOKEN_PACKAGES'tan seç
        if not pkg:
            for default_pkg in DEFAULT_TOKEN_PACKAGES:
                if default_pkg['id'] == package_id or default_pkg['name'].lower() == package_id_str:
                    # Mock nesne oluştur
                    pkg = type('PseudoPkg', (), {
                        'id': default_pkg['id'],
                        'name': default_pkg['name'],
                        'description': default_pkg.get('description', ''),
                        'tokens': default_pkg['tokens'],
                        'price_usd': default_pkg['price_usd'],
                        'price_try': default_pkg.get('price_try'),
                        'bonus_pct': default_pkg.get('bonus_pct', 0),
                    })()
                    break
    except Exception as e:
        print(f"Package lookup error: {e}")
        
    return pkg

@app.route('/api/billing/iyzico/config', methods=['GET'])
def iyzico_config():
    """Frontend'in Iyzico'nun aktif olup olmadığını anlaması için config endpoint."""
    credentials_configured = bool(_IYZICO_API_KEY and _IYZICO_SECRET_KEY)
    enabled = bool(credentials_configured)

    return jsonify({
        'enabled': enabled,
        'base_url': _IYZICO_BASE_URL if enabled else None,
        'credentials_configured': credentials_configured,
        'reason': None if enabled else 'credentials_missing',
    })

@app.route('/api/billing/iyzico/checkout-session', methods=['POST'])
def iyzico_checkout_session():
    """Iyzico Checkout oturumu oluşturur ve checkout URL'sini döndürür."""
    user, _auth_mode, _key_record, err_payload, status_code = _resolve_authenticated_user(
        preferred_api_key_client='vscode',
        allow_jwt=True,
    )
    if err_payload:
        return jsonify(err_payload), status_code

    user_id = user.id

    if not _IYZICO_API_KEY or not _IYZICO_SECRET_KEY:
        return jsonify({'error': 'Iyzico is not configured on the server.'}), 503

    data = request.get_json()
    if not data:
        return jsonify({'error': 'No JSON data in request'}), 400
    package_id = data.get('package_id')
    # Fallback is opt-in so explicit Iyzico choice does not silently redirect to Stripe.
    allow_stripe_fallback = bool(data.get('allow_stripe_fallback', False))
    package = _find_token_package(package_id)
    if not package:
        return jsonify({'error': 'Invalid package'}), 400

    try:
        # İsim/Soyisim ayrımı
        full_name = user.display_name or "Code Alchemist User"
        name_parts = full_name.split(' ', 1)
        first_name = name_parts[0]
        last_name = name_parts[1] if len(name_parts) > 1 else "User"

        frontend_base_url = _get_iyzico_frontend_base_url(request.headers.get('Origin'))
        callback_url = request.host_url.rstrip('/') + '/api/billing/iyzico/callback'

        # Iyzico Checkout Form Initialize Request
        iyzipay_request = {
            'locale': 'tr',
            'conversationId': str(user_id) + "_" + _utcnow().strftime('%Y%m%d%H%M%S'),
            'price': str(package.price_try or (package.price_usd * 32)), 
            'paidPrice': str(package.price_try or (package.price_usd * 32)),
            'currency': 'TRY',
            'basketId': 'B' + str(user_id),
            'paymentGroup': 'PRODUCT',
            'callbackUrl': callback_url,
            'buyer': {
                'id': str(user_id),
                'name': first_name,
                'surname': last_name,
                'email': user.email,
                'identityNumber': '11111111111', # Placeholder
                'registrationAddress': 'Online Course Access',
                'ip': request.remote_addr,
                'city': 'Istanbul',
                'country': 'Turkey'
            },
            'shippingAddress': {
                'contactName': full_name,
                'city': 'Istanbul',
                'country': 'Turkey',
                'address': 'Digital Product'
            },
            'billingAddress': {
                'contactName': full_name,
                'city': 'Istanbul',
                'country': 'Turkey',
                'address': 'Digital Product',
                'zipCode': '34000'
            },
            'basketItems': [
                {
                    'id': 'TP' + str(package.id),
                    'name': package.name + " Tokens",
                    'category1': 'Tokens',
                    'itemType': 'VIRTUAL',
                    'price': str(package.price_try or (package.price_usd * 32))
                }
            ]
        }

        checkout_form_initialize = iyzipay.CheckoutFormInitialize().create(iyzipay_request, iyzi_options)
        raw_response = checkout_form_initialize.read()
        checkout_result = json.loads(raw_response.decode('utf-8'))

        if checkout_result.get('status') == 'success':
            # Ödemeyi bekleme aşamasında DB'ye kaydet
            purchase = TokenPurchase(
                user_id=user_id,
                package_id=package.id,
                package_name=package.name,
                tokens_granted=package.tokens,
                amount_cents=int((package.price_try or (package.price_usd * 32)) * 100),
                currency='TRY',
                stripe_checkout_session_id='iyz_' + checkout_result.get('token', ''), # Token'ı buraya saklıyoruz
                status='pending',
                metadata_json=json.dumps({'frontend_base_url': frontend_base_url})
            )
            db.session.add(purchase)
            db.session.commit()

            return jsonify({
                'payment_url': checkout_result.get('paymentPageUrl'),
                'token': checkout_result.get('token'),
                'gateway': 'iyzico'
            })
        else:
            if not allow_stripe_fallback:
                error_code = checkout_result.get('errorCode')
                status_code = 503 if error_code == '1001' else 400
                error_message = checkout_result.get('errorMessage', 'Iyzico checkout failed')
                if error_code == '1001':
                    error_message = 'Iyzico API credentials are invalid or inactive.'
                return jsonify({
                    'error': error_message,
                    'gateway': 'iyzico',
                    'code': error_code
                }), status_code

            # Iyzico fail olursa (yalnizca opt-in durumunda) Stripe'a fallback yap
            print(f"[FALLBACK] Iyzico failed: {checkout_result.get('errorMessage')}. Using Stripe instead.")

            # Stripe Checkout Session oluştur
            try:
                stripe_session = stripe.checkout.Session.create(
                    payment_method_types=['card'],
                    line_items=[{
                        'price_data': {
                            'currency': 'usd',
                            'product_data': {
                                'name': package.name + " Tokens",
                                'description': f"{package.tokens} tokens for Code Alchemist"
                            },
                            'unit_amount': int(package.price_usd * 100)
                        },
                        'quantity': 1
                    }],
                    mode='payment',
                    success_url=frontend_base_url + '?billing=success&gateway=stripe',
                    cancel_url=frontend_base_url + '?billing=cancel',
                    metadata={
                        'user_id': user_id,
                        'package_id': package.id,
                        'package_name': package.name,
                        'tokens': package.tokens
                    }
                )

                # Stripe session'ı DB'ye kaydet
                purchase = TokenPurchase(
                    user_id=user_id,
                    package_id=package.id,
                    package_name=package.name,
                    tokens_granted=package.tokens,
                    amount_cents=int(package.price_usd * 100),
                    currency='USD',
                    stripe_checkout_session_id=stripe_session.id,
                    status='pending',
                    metadata_json=json.dumps({'gateway': 'stripe'})
                )
                db.session.add(purchase)
                db.session.commit()

                return jsonify({
                    'payment_url': stripe_session.url,
                    'token': stripe_session.id,
                    'gateway': 'stripe'
                })
            except Exception as stripe_err:
                print(f"[ERROR] Stripe fallback also failed: {stripe_err}")
                return jsonify({'error': 'Payment gateway unavailable'}), 503

    except Exception as e:
        print(f"Iyzico checkout error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/billing/iyzico/callback', methods=['POST'])
def iyzico_callback():
    """Iyzico ödeme tamamlanma olayını dinler ve token bakiyesini günceller."""
    token = request.form.get('token')
    if not token:
        return "Internal Error: Token missing", 400

    # Token ile ödeme sonucunu sorgula
    iyzi_request = {'locale': 'tr', 'token': token}
    result = iyzipay.CheckoutForm().retrieve(iyzi_request, iyzi_options)
    result_data = json.loads(result.read().decode('utf-8'))

    if result_data.get('status') == 'success' and result_data.get('paymentStatus') == 'SUCCESS':
        # Başarılı ödeme
        purchase = TokenPurchase.query.filter_by(stripe_checkout_session_id='iyz_' + token).first()
        if purchase and purchase.status != 'completed':
            purchase.status = 'completed'
            purchase.completed_at = _utcnow()

            frontend_base_url = _get_iyzico_frontend_base_url()
            if purchase.metadata_json:
                try:
                    metadata = json.loads(purchase.metadata_json)
                    if isinstance(metadata, dict) and metadata.get('frontend_base_url'):
                        frontend_base_url = str(metadata['frontend_base_url']).rstrip('/')
                except Exception:
                    pass
            
            # Token yüklemesi
            balance = TokenBalance.query.filter_by(user_id=purchase.user_id).first()
            if not balance:
                balance = TokenBalance(user_id=purchase.user_id, balance=0)
                db.session.add(balance)
            
            balance.balance += purchase.tokens_granted
            
            # Log
            log = TokenTransaction(
                user_id=purchase.user_id,
                amount=purchase.tokens_granted,
                type='purchase',
                description=f"Tokens purchased via Iyzico (Package: {purchase.package_name})",
                reference_id='iyz_' + token
            )
            db.session.add(log)
            db.session.commit()

            return redirect(f'{frontend_base_url}/?billing=success&gateway=iyzico')
    
    return redirect(f'{_get_iyzico_frontend_base_url()}/?billing=error&gateway=iyzico')


@app.route('/api/tokens/renewal-status', methods=['GET'])
@jwt_required()
def get_renewal_status():
    """Kullanıcının kendi otomatik yenileme durumunu görmesi için."""
    user = get_current_user()
    if not user:
        return jsonify({'error': 'Unauthorized'}), 401
    
    wallet = get_or_create_token_balance(user)
    has_purchase = TokenPurchase.query.filter_by(user_id=user.id, status='completed').first() is not None

    return jsonify({
        'monthly_renewal_enabled': wallet.monthly_renewal_enabled,
        'monthly_renewal_day': wallet.monthly_renewal_day,
        'last_renewal_at': wallet.last_renewal_at.isoformat() if wallet.last_renewal_at else None,
        'can_enable': has_purchase
    })


@app.route('/api/tokens/renewal-status', methods=['PUT'])
@jwt_required()
def update_renewal_status():
    """Kullanıcının kendi otomatik yenilemesini açıp kapatması için."""
    user = get_current_user()
    if not user:
        return jsonify({'error': 'Unauthorized'}), 401
    
    data = request.get_json(silent=True) or {}
    enabled = data.get('enabled')
    
    if enabled:
        # Sadece en az bir kez paket almış olanlar auto-renew açabilir
        has_purchase = TokenPurchase.query.filter_by(user_id=user.id, status='completed').first()
        if not has_purchase:
            return jsonify({'error': 'Otomatik yenilemeyi aktif etmek için en az bir paket satın almış olmanız gerekmektedir.'}), 400
        
    wallet = get_or_create_token_balance(user)
    wallet.monthly_renewal_enabled = bool(enabled)
    
    # Eğer yenileme günü set edilmemişse ve aktif ediliyorsa, üyelik tarihini set edelim
    if wallet.monthly_renewal_enabled and wallet.monthly_renewal_day is None:
        wallet.monthly_renewal_day = min(28, user.created_at.day if user.created_at else _utcnow().day)
        
    db.session.commit()
    return jsonify({
        'success': True,
        'monthly_renewal_enabled': wallet.monthly_renewal_enabled,
        'monthly_renewal_day': wallet.monthly_renewal_day
    })


@app.route('/api/billing/config', methods=['GET'])
def billing_config():
    """Frontend'in Stripe'ın aktif olup olmadığını anlaması için config endpoint."""
    enabled = bool(_STRIPE_SECRET_KEY and _STRIPE_PUBLISHABLE_KEY)
    return jsonify({
        'enabled': enabled,
        'public_key': _STRIPE_PUBLISHABLE_KEY if enabled else None,
    })


@app.route('/api/billing/packages', methods=['GET'])
def billing_packages():
    """Aktif token paketlerini listeler. DB boşsa varsayılanlar döner."""
    try:
        pkgs = TokenPackage.query.filter_by(is_active=True).order_by(TokenPackage.price_usd).all()
        if pkgs:
            packages_data = [
                {
                    'id': p.id,
                    'name': p.name,
                    'description': p.description,
                    'tokens': p.tokens,
                    'price_usd': p.price_usd,
                    'bonus_pct': p.bonus_pct,
                    'highlight': (p.bonus_pct is not None and p.bonus_pct > 0),
                }
                for p in pkgs
            ]
        else:
            # DB boşsa statik fallback
            packages_data = [
                {**pkg, 'id': pkg['name'].lower().replace(' ', '-'), 'highlight': pkg.get('bonus_pct', 0) > 0}
                for pkg in DEFAULT_TOKEN_PACKAGES
            ]
        return jsonify({'packages': packages_data})
    except Exception as e:
        print(f"billing_packages error: {e}")
        return jsonify({'error': 'Could not fetch packages.'}), 500


@app.route('/api/tokens/usage', methods=['GET'])
def token_usage():
    """Kullanıcının token bakiyesi ve son işlem geçmişini döndürür."""
    user, _auth_mode, _key_record, err_payload, status_code = _resolve_authenticated_user(
        preferred_api_key_client='vscode',
        allow_jwt=True,
    )
    if err_payload:
        return jsonify(err_payload), status_code

    user_id = user.id
    limit = max(1, min(int(request.args.get('limit', 10)), 50))
    try:
        balance_record = get_or_create_token_balance(user)

        transactions = (
            TokenTransaction.query
            .filter_by(user_id=user_id)
            .order_by(TokenTransaction.created_at.desc())
            .limit(limit)
            .all()
        )
        return jsonify({
            'balance': balance_record.balance,
            'total_spent': balance_record.total_spent,
            'transactions': [
                {
                    'id': tx.id,
                    'amount': tx.amount,
                    'type': tx.type,
                    'description': tx.description,
                    'created_at': tx.created_at.isoformat(),
                }
                for tx in transactions
            ],
        })
    except Exception as e:
        print(f"token_usage error: {e}")
        return jsonify({'error': 'Could not fetch token usage.'}), 500


@app.route('/api/quota/status', methods=['GET'])
@jwt_required()
def get_quota_status_endpoint():
    """Kullanıcının haftalık/günlük kota durumunu döndürür."""
    try:
        user_id = get_jwt_identity()
        if not user_id:
            return jsonify({'error': 'User not found'}), 401
        
        user = User.query.get(int(user_id))
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        wallet = get_or_create_token_balance(user)
        if wallet.weekly_reset_at is None or wallet.daily_reset_at is None:
            init_quota_for_new_user(wallet)
        
        quota_data = get_quota_status(user.id)
        if quota_data:
            return jsonify(quota_data), 200
        else:
            return jsonify({'error': 'Could not get quota data'}), 500
    except Exception as e:
        print(f"get_quota_status error: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Could not fetch quota status: {str(e)}'}), 500


@app.route('/api/billing/checkout-session', methods=['POST'])
def create_checkout_session():
    """Stripe Checkout oturumu oluşturur ve checkout URL'sini döndürür."""
    if not _STRIPE_SECRET_KEY:
        return jsonify({'error': 'Billing is not configured on the server. Contact admin.'}), 503

    user, _auth_mode, _key_record, err_payload, status_code = _resolve_authenticated_user(
        preferred_api_key_client='vscode',
        allow_jwt=True,
    )
    if err_payload:
        return jsonify(err_payload), status_code

    user_id = user.id
    data = request.get_json(silent=True) or {}
    package_id = data.get('package_id')

    # Paketi bul (hem int DB id hem string slug/name için)
    pkg = None
    package_id_str = str(package_id).lower().replace('-', ' ')
    
    try:
        # 1. DB'den Ara (ID ile)
        if isinstance(package_id, int) or package_id.isdigit():
            pkg = TokenPackage.query.filter_by(id=int(package_id), is_active=True).first()
        
        # 2. DB'den Ara (İsim/Slug ile)
        if not pkg:
            pkg = TokenPackage.query.filter(
                (db.func.lower(TokenPackage.name) == package_id_str) |
                (db.func.lower(TokenPackage.name) == str(package_id).lower())
            ).filter_by(is_active=True).first()

        # 3. Fallback: Seeding yapılmamışsa veya DB'de yoksa DEFAULT_TOKEN_PACKAGES'tan seç
        if not pkg:
            for default_pkg in DEFAULT_TOKEN_PACKAGES:
                # Hem ID hem İsim kontrolü
                if default_pkg['id'] == package_id or default_pkg['name'].lower() == package_id_str:
                    # Mock nesne oluştur
                    pkg = type('PseudoPkg', (), {
                        'id': None,
                        'name': default_pkg['name'],
                        'description': default_pkg.get('description', ''),
                        'tokens': default_pkg['tokens'],
                        'price_usd': default_pkg['price_usd'],
                        'bonus_pct': default_pkg.get('bonus_pct', 0),
                    })()
                    break

    except Exception as e:
        print(f"Package matching error: {e}")

    if not pkg:
        return jsonify({'error': f'Package not found: {package_id}'}), 404

    # Bonus token hesabı
    bonus_tokens = int(pkg.tokens * pkg.bonus_pct / 100) if pkg.bonus_pct else 0
    total_tokens = pkg.tokens + bonus_tokens
    price_cents = int(round(pkg.price_usd * 100))

    # Render/production URL tespiti
    host = request.host_url.rstrip('/')
    success_url = f"{host}/?payment=success&tokens={total_tokens}"
    cancel_url = f"{host}/?payment=canceled"

    try:
        session = stripe.checkout.Session.create(
            payment_method_types=['card'],
            mode='payment',
            customer_email=user.email,
            line_items=[{
                'price_data': {
                    'currency': 'usd',
                    'unit_amount': price_cents,
                    'product_data': {
                        'name': f"{pkg.name} — {total_tokens} Tokens",
                        'description': pkg.description or f"{pkg.tokens} tokens" + (f" + %{pkg.bonus_pct} bonus" if pkg.bonus_pct else ""),
                    },
                },
                'quantity': 1,
            }],
            success_url=success_url,
            cancel_url=cancel_url,
            metadata={
                'user_id': str(user_id),
                'package_name': pkg.name,
                'tokens_granted': str(total_tokens),
            },
        )
    except stripe.error.StripeError as e:
        print(f"Stripe checkout error: {e}")
        return jsonify({'error': f'Payment session could not be created: {str(e)}'}), 502
    except Exception as e:
        print(f"Unexpected checkout error: {e}")
        return jsonify({'error': 'An unexpected error occurred. Please try again.'}), 500

    # Veritabanına bekleyen satın alma kaydı ekle
    try:
        purchase = TokenPurchase(
            user_id=user_id,
            package_id=pkg.id if hasattr(pkg, 'id') and pkg.id else None,
            package_name=pkg.name,
            tokens_granted=total_tokens,
            amount_cents=price_cents,
            currency='usd',
            stripe_checkout_session_id=session.id,
            status='pending',
        )
        db.session.add(purchase)
        db.session.commit()
    except Exception as db_err:
        print(f"DB insert for checkout failed: {db_err}")
        db.session.rollback()
        # Stripe oturumu oluştu ama DB'ye yazamadık — yine de yönlendir, webhook düzeltir.

    return jsonify({'checkout_url': session.url})


@app.route('/api/billing/webhook', methods=['POST'])
def stripe_webhook():
    """Stripe ödeme tamamlanma olayını dinler ve token bakiyesini günceller."""
    payload = request.get_data(as_text=False)
    sig_header = request.headers.get('Stripe-Signature', '')

    if not _STRIPE_WEBHOOK_SECRET:
        return jsonify({'error': 'Webhook secret not configured.'}), 500

    try:
        event = stripe.Webhook.construct_event(payload, sig_header, _STRIPE_WEBHOOK_SECRET)
    except stripe.error.SignatureVerificationError:
        print("Webhook signature verification failed.")
        return jsonify({'error': 'Invalid signature'}), 400
    except Exception as e:
        print(f"Webhook parse error: {e}")
        return jsonify({'error': 'Webhook parse error'}), 400

    if event['type'] == 'checkout.session.completed':
        session_obj = event['data']['object']
        stripe_session_id = session_obj.get('id')
        stripe_payment_intent = session_obj.get('payment_intent')

        if not stripe_session_id:
            print('Webhook: missing stripe session id. Skipping event.')
            return jsonify({'received': True}), 200

        # Satın alma kaydını bul
        purchase = TokenPurchase.query.filter_by(
            stripe_checkout_session_id=stripe_session_id
        ).first()

        if not purchase:
            # Metadata'dan kullanıcı bilgisini çekelim (fallback)
            meta = session_obj.get('metadata', {})
            user_id_str = meta.get('user_id')
            tokens_granted_str = meta.get('tokens_granted')
            package_name = meta.get('package_name', 'Unknown')

            if not user_id_str or not tokens_granted_str:
                print(f"Webhook: No purchase record and no metadata for session {stripe_session_id}")
                return jsonify({'received': True}), 200

            # Kayıt yoksa oluştur
            try:
                purchase = TokenPurchase(
                    user_id=int(user_id_str),
                    package_name=package_name,
                    tokens_granted=int(tokens_granted_str),
                    amount_cents=int(session_obj.get('amount_total', 0)),
                    currency=session_obj.get('currency', 'usd'),
                    stripe_checkout_session_id=stripe_session_id,
                    stripe_payment_intent_id=stripe_payment_intent,
                    status='pending',
                )
                db.session.add(purchase)
                db.session.flush()
            except Exception as e:
                print(f"Webhook: Could not create fallback purchase: {e}")
                db.session.rollback()
                return jsonify({'received': True}), 200

        # --- Idempotency: Zaten completed mi? ---
        if purchase.status == 'completed':
            print(f"Webhook: Purchase {purchase.id} already completed. Skipping.")
            return jsonify({'received': True}), 200

        # Idempotency hardening: purchase row'unu atomik sekilde claim et.
        # Sadece pending durumundaki ilk worker processing'e cekebilir.
        claimed = (
            TokenPurchase.query
            .filter_by(id=purchase.id, status='pending')
            .update({'status': 'processing'}, synchronize_session=False)
        )
        if claimed == 0:
            db.session.rollback()
            refreshed = TokenPurchase.query.filter_by(id=purchase.id).first()
            current_status = refreshed.status if refreshed else 'missing'
            print(f"Webhook: Purchase {purchase.id} already claimed (status={current_status}).")
            return jsonify({'received': True}), 200

        db.session.flush()

        user_id = purchase.user_id
        tokens_to_add = purchase.tokens_granted

        try:
            existing_tx = TokenTransaction.query.filter_by(
                user_id=user_id,
                type='purchase',
                reference_id=stripe_session_id,
            ).first()
            if existing_tx:
                purchase.status = 'completed'
                purchase.stripe_payment_intent_id = stripe_payment_intent
                if not purchase.completed_at:
                    purchase.completed_at = _utcnow()
                db.session.commit()
                print(f"Webhook: Duplicate purchase event ignored for session {stripe_session_id}.")
                return jsonify({'received': True}), 200

            # TokenBalance güncelle veya oluştur
            balance_record = TokenBalance.query.filter_by(user_id=user_id).first()
            if not balance_record:
                balance_record = TokenBalance(user_id=user_id, balance=0, total_spent=0)
                db.session.add(balance_record)
                db.session.flush()

            balance_record.balance += tokens_to_add

            # TokenTransaction kaydet
            tx = TokenTransaction(
                user_id=user_id,
                amount=tokens_to_add,
                type='purchase',
                description=f"Token purchase: {purchase.package_name or 'Package'} ({tokens_to_add} tokens)",
                reference_id=stripe_session_id,
            )
            db.session.add(tx)

            # Satın alma kaydını tamamlandı olarak işaretle
            purchase.status = 'completed'
            purchase.stripe_payment_intent_id = stripe_payment_intent
            purchase.completed_at = _utcnow()

            db.session.commit()
            print(f"Webhook: Granted {tokens_to_add} tokens to user {user_id}.")
        except Exception as e:
            db.session.rollback()
            print(f"Webhook: DB update failed for user {user_id}: {e}")
            return jsonify({'error': 'DB update failed'}), 500

    return jsonify({'received': True}), 200


# ==========================================
# API KEYS & VS CODE EXTENSION ENDPOINTS
# ==========================================
import secrets
from models import ApiKey


def _extract_api_key_from_request() -> str:
    """Extract API key from X-API-Key or Authorization header."""
    x_api_key = (request.headers.get('X-API-Key') or '').strip()
    if x_api_key:
        return x_api_key

    auth = (request.headers.get('Authorization') or '').strip()
    if not auth:
        return ''

    if auth.startswith('ca-'):
        return auth

    if auth.lower().startswith('bearer '):
        parts = auth.split(None, 1)
        if len(parts) == 2:
            token = parts[1].strip()
            if token.startswith('ca-'):
                return token

    return ''


def _detect_api_key_client(key_value: str) -> str:
    token = (key_value or '').lower()
    if token.startswith('ca-web-'):
        return 'web'
    if token.startswith('ca-vsc-'):
        return 'vscode'
    if token.startswith('ca-'):
        return 'legacy'
    return 'unknown'


def _hash_api_key_value(raw_key: str) -> str:
    return hashlib.sha256((raw_key or '').encode('utf-8')).hexdigest()


def _build_stored_api_key(raw_key: str) -> str:
    token = (raw_key or '').strip()
    if not token.startswith('ca-'):
        return ''

    client = _detect_api_key_client(token)
    if client == 'vscode':
        prefix = 'ca-vsc'
    elif client == 'web':
        prefix = 'ca-web'
    else:
        prefix = 'ca'

    return f"{prefix}-h-{_hash_api_key_value(token)}"


def _preview_api_key(stored_key: str) -> str:
    token = (stored_key or '').strip().lower()
    if '-h-' in token:
        if token.startswith('ca-vsc-'):
            return f"ca-vsc-***{token[-6:]}"
        if token.startswith('ca-web-'):
            return f"ca-web-***{token[-6:]}"
        if token.startswith('ca-'):
            return f"ca-***{token[-6:]}"
    return f"{stored_key[:7]}***{stored_key[-4:]}" if len(stored_key) > 11 else "***"


def _find_api_key_record(raw_key: str):
    token = (raw_key or '').strip()
    if not token:
        return None

    stored_candidate = _build_stored_api_key(token)
    if stored_candidate:
        record = ApiKey.query.filter_by(key=stored_candidate, is_active=True).first()
        if record:
            return record

    # Backward compatibility for legacy plaintext records.
    record = ApiKey.query.filter_by(key=token, is_active=True).first()
    if record and stored_candidate:
        try:
            record.key = stored_candidate
            db.session.commit()
        except Exception:
            db.session.rollback()
    return record


def _migrate_plaintext_api_keys_to_hash() -> int:
    migrated = 0
    try:
        records = ApiKey.query.all()
        for record in records:
            current = (record.key or '').strip()
            if not current.startswith('ca-'):
                continue
            if '-h-' in current:
                continue

            stored = _build_stored_api_key(current)
            if not stored:
                continue

            record.key = stored
            migrated += 1

        if migrated > 0:
            db.session.commit()
        return migrated
    except Exception:
        db.session.rollback()
        return 0

from functools import wraps
def requires_api_key(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        api_key_header = _extract_api_key_from_request()
        if not api_key_header:
            return jsonify({'error': 'X-API-Key header is missing'}), 401
        
        key_record = _find_api_key_record(api_key_header)
        if not key_record or not key_record.is_active:
            return jsonify({'error': 'Invalid or revoked API Key'}), 401
        
        # Security: ensure last_used_at is updated
        key_record.last_used_at = datetime.datetime.now(datetime.timezone.utc).replace(tzinfo=None)
        db.session.commit()
            
        user = User.query.get(key_record.user_id)
        if not user:
             return jsonify({'error': 'User not found associated with this API key'}), 401
             
        return f(user, *args, **kwargs)
    return decorated
    return decorated


with app.app_context():
    migrated_count = _migrate_plaintext_api_keys_to_hash()
    if migrated_count > 0:
        print(f"Migrated {migrated_count} plaintext API keys to hashed storage.")


def _resolve_authenticated_user(preferred_api_key_client: str | None = None, allow_jwt: bool = True):
    """Resolve user via JWT first (optional), then API key headers.

    Returns: (user, auth_mode, key_record, error_payload, status_code)
    """
    if allow_jwt:
        identity = _get_safe_jwt_identity()
        if identity:
            try:
                user = db.session.get(User, int(identity))
            except Exception:
                user = None
            if user:
                return user, 'jwt', None, None, None

    api_key_value = _extract_api_key_from_request()
    if not api_key_value:
        return None, None, None, {'error': 'Missing credentials. Provide Authorization Bearer token or X-API-Key.'}, 401

    key_record = _find_api_key_record(api_key_value)
    if not key_record:
        return None, None, None, {'error': 'Invalid or revoked API Key'}, 401

    client_type = _detect_api_key_client(key_record.key)
    if preferred_api_key_client and client_type not in (preferred_api_key_client, 'legacy'):
        return None, None, None, {
            'error': f'This endpoint requires a {preferred_api_key_client} API key.',
            'expected_client': preferred_api_key_client,
            'provided_client': client_type,
        }, 403

    user = db.session.get(User, key_record.user_id)
    if not user:
        return None, None, None, {'error': 'User not found for API key'}, 404

    key_record.last_used_at = _utcnow()
    db.session.commit()
    return user, 'api_key', key_record, None, None

@app.route('/api/keys', methods=['GET'])
@jwt_required()
def get_api_keys():
    user_id = get_jwt_identity()
    # Sadece aktif anahtarları getir
    keys = ApiKey.query.filter_by(user_id=user_id, is_active=True).order_by(ApiKey.created_at.desc()).all()
    return jsonify({
        'keys': [
            {
                'id': k.id,
                'name': k.name,
                'client': _detect_api_key_client(k.key),
                'key_preview': _preview_api_key(k.key),
                'created_at': k.created_at.isoformat(),
                'last_used_at': k.last_used_at.isoformat() if k.last_used_at else None
            } for k in keys
        ]
    })

@app.route('/api/keys', methods=['POST'])
@jwt_required()
def create_api_key():
    user_id = get_jwt_identity()
    data = request.get_json() or {}
    name = data.get('name', 'My API Key').strip()
    client = str(data.get('client', 'vscode') or 'vscode').strip().lower()

    if client not in {'vscode', 'web'}:
        return jsonify({'error': 'client must be one of: vscode, web'}), 400

    if not name:
        return jsonify({'error': 'Name is required'}), 400

    # Web ve VS Code için ayrı API key namespace'leri kullan.
    prefix = 'ca-vsc' if client == 'vscode' else 'ca-web'
    token = f"{prefix}-{secrets.token_hex(16)}"
    stored_token = _build_stored_api_key(token)
    
    new_key = ApiKey(user_id=user_id, name=name, key=stored_token)
    db.session.add(new_key)
    db.session.commit()
    
    # Tam token sadece bir kez döndürülür
    return jsonify({
        'message': 'API Key created successfully',
        'key': {
            'id': new_key.id,
            'name': new_key.name,
            'client': client,
            'token': token,
            'created_at': new_key.created_at.isoformat()
        }
    })


@app.route('/v1/auth/login', methods=['POST'])
def vscode_login_and_issue_api_key():
    """Login endpoint for VS Code clients; returns a vscode-scoped API key."""
    data = request.get_json(silent=True) or {}
    email = (data.get('email') or '').strip().lower()
    password = data.get('password') or ''
    key_name = (data.get('key_name') or 'VS Code Extension').strip()

    if not email or not password:
        return jsonify({'error': 'email and password are required'}), 400

    user = User.query.filter_by(email=email).first()
    if not user or not verify_password(password, user.password_hash):
        return jsonify({'error': 'Incorrect email or password.'}), 401

    token = f"ca-vsc-{secrets.token_hex(16)}"
    stored_token = _build_stored_api_key(token)
    new_key = ApiKey(user_id=user.id, name=key_name or 'VS Code Extension', key=stored_token)
    db.session.add(new_key)
    db.session.commit()

    return jsonify({
        'message': 'VS Code API key issued successfully.',
        'token_type': 'vscode_api_key',
        'api_key': token,
        'key': {
            'id': new_key.id,
            'name': new_key.name,
            'client': 'vscode',
            'created_at': new_key.created_at.isoformat(),
        },
        'user': {
            'id': user.id,
            'email': user.email,
            'display_name': user.display_name,
        }
    })


@app.route('/v1/auth/vscode/login', methods=['GET', 'POST'])
@app.route('/api/v1/auth/vscode/login', methods=['GET', 'POST'])
@app.route('/api/auth/vscode/login', methods=['GET', 'POST'])
def vscode_login_page():
    """Browser login page for VS Code flow.

    User logs in on web page, then extension polls `/v1/auth/vscode/poll` with state.
    """
    _cleanup_vscode_login_state()

    if request.method == 'GET':
        state = (request.args.get('state') or '').strip()
        if not _is_valid_vscode_state(state):
            return jsonify({'error': 'Invalid state'}), 400

        html = f"""<!doctype html>
<html lang='tr'>
<head>
  <meta charset='utf-8'>
  <meta name='viewport' content='width=device-width,initial-scale=1'>
  <title>CodeAlchemist VS Code Login</title>
  <style>
    body {{ font-family: Segoe UI, Arial, sans-serif; margin: 0; background: #0b1220; color: #e5e7eb; }}
    .wrap {{ max-width: 520px; margin: 48px auto; padding: 24px; background: #111827; border: 1px solid #1f2937; border-radius: 12px; }}
    h1 {{ font-size: 22px; margin: 0 0 8px; }}
    p {{ color: #9ca3af; margin: 0 0 18px; }}
    label {{ display: block; margin: 10px 0 6px; font-size: 13px; color: #cbd5e1; }}
    input {{ width: 100%; box-sizing: border-box; padding: 11px; border: 1px solid #334155; border-radius: 8px; background: #0f172a; color: #e5e7eb; }}
    button {{ margin-top: 14px; width: 100%; padding: 12px; border: 0; border-radius: 8px; background: #2563eb; color: #fff; font-weight: 600; cursor: pointer; }}
    .hint {{ font-size: 12px; color: #94a3b8; margin-top: 12px; }}
  </style>
</head>
<body>
  <div class='wrap'>
    <h1>VS Code için giriş yap</h1>
    <p>Bu giriş, VS Code extension için özel bir API anahtarı oluşturur ve otomatik olarak kaydedilir.</p>
    <form method='post' action='/v1/auth/vscode/login'>
      <input type='hidden' name='state' value='{state}'>
      <label>Email</label>
      <input type='email' name='email' required autocomplete='username'>
      <label>Şifre</label>
      <input type='password' name='password' required autocomplete='current-password'>
      <button type='submit'>Giriş Yap</button>
    </form>
    <div class='hint'>Girişten sonra bu pencereyi kapatabilirsiniz.</div>
  </div>
</body>
</html>"""
        return Response(html, mimetype='text/html')

    state = (request.form.get('state') or '').strip()
    email = (request.form.get('email') or '').strip().lower()
    password = request.form.get('password') or ''

    if not _is_valid_vscode_state(state):
        return jsonify({'error': 'Invalid state'}), 400
    if not email or not password:
        return jsonify({'error': 'email and password are required'}), 400

    user = User.query.filter_by(email=email).first()
    if not user or not verify_password(password, user.password_hash):
        html = """<!doctype html><html><head><meta charset='utf-8'><title>Login failed</title></head><body style='font-family:Segoe UI,Arial,sans-serif;background:#0b1220;color:#e5e7eb;padding:24px;'><h2>Giriş başarısız</h2><p>Email veya şifre hatalı. VS Code'dan tekrar deneyin.</p></body></html>"""
        return Response(html, mimetype='text/html'), 401

    token = f"ca-vsc-{secrets.token_hex(16)}"
    stored_token = _build_stored_api_key(token)
    new_key = ApiKey(user_id=user.id, name='VS Code Extension', key=stored_token)
    db.session.add(new_key)
    db.session.commit()

    # Database-backed shared state (instead of in-memory _vscode_login_state)
    try:
        # Check if state exists, update if it does, otherwise create
        existing_state = VSCodeLoginState.query.filter_by(state=state).first()
        if existing_state:
            existing_state.api_key = token
            existing_state.user_id = user.id
            existing_state.expires_at = time.time() + VSCODE_LOGIN_STATE_TTL_SECONDS
        else:
            new_state = VSCodeLoginState(
                state=state,
                api_key=token,
                user_id=user.id,
                expires_at=time.time() + VSCODE_LOGIN_STATE_TTL_SECONDS
            )
            db.session.add(new_state)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        print(f"[VSCODE-AUTH] Login state save error: {e}")
        html = """<body style='background:#0b1220;color:#e5e7eb;'><h2>Hata</h2><p>Giriş kaydedilemedi. Lütfen tekrar deneyin.</p></body>"""
        return Response(html, mimetype='text/html'), 500

    html = """<!doctype html>
<html lang='tr'>
<head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'><title>CodeAlchemist Login</title>
<style>body{font-family:Segoe UI,Arial,sans-serif;background:#0b1220;color:#e5e7eb;margin:0}.wrap{max-width:560px;margin:48px auto;padding:24px;background:#111827;border:1px solid #1f2937;border-radius:12px}h1{margin:0 0 10px;font-size:24px}.ok{color:#22c55e;font-weight:700}</style></head>
<body><div class='wrap'><h1 class='ok'>Giriş başarılı</h1><p>Giriş başarılı, bu pencereyi kapatabilirsiniz.</p></div></body>
</html>"""
    return Response(html, mimetype='text/html')


@app.route('/v1/auth/vscode/poll', methods=['GET'])
@app.route('/api/v1/auth/vscode/poll', methods=['GET'])
@app.route('/api/auth/vscode/poll', methods=['GET'])
def vscode_login_poll():
    """Extension polling endpoint for browser-based login completion."""
    _cleanup_vscode_login_state()
    state = (request.args.get('state') or '').strip()
    if not _is_valid_vscode_state(state):
        return jsonify({'status': 'invalid_state'}), 400

    payload = VSCodeLoginState.query.filter_by(state=state).first()
    if not payload:
        return jsonify({'status': 'pending'})

    if time.time() > float(payload.expires_at or 0):
        try:
            db.session.delete(payload)
            db.session.commit()
        except:
            db.session.rollback()
        return jsonify({'status': 'expired'})

    api_key = str(payload.api_key or '').strip()
    if not api_key:
        return jsonify({'status': 'pending'})

    return jsonify({
        'status': 'ready',
        'api_key': api_key,
    })


@app.route('/v1/auth/vscode/consume', methods=['POST'])
@app.route('/api/v1/auth/vscode/consume', methods=['POST'])
@app.route('/api/auth/vscode/consume', methods=['POST'])
def vscode_login_consume():
    """Mark login state as consumed by extension and remove it from database."""
    _cleanup_vscode_login_state()
    data = request.get_json(silent=True) or {}
    state = str(data.get('state') or '').strip()
    if not _is_valid_vscode_state(state):
        return jsonify({'status': 'invalid_state'}), 400

    record = VSCodeLoginState.query.filter_by(state=state).first()
    if record:
        try:
            db.session.delete(record)
            db.session.commit()
            return jsonify({'status': 'consumed'})
        except:
            db.session.rollback()
            return jsonify({'status': 'error'})
    
    return jsonify({'status': 'missing'})

@app.route('/v1/auth/vscode/generate-otp', methods=['POST'])
@requires_api_key
def vscode_generate_otp(user):
    """Generates a secure, 2-minute OTP for VS Code to Browser auth sync."""
    otp = secrets.token_urlsafe(32)
    # 2 minute expiry
    expires_at = datetime.datetime.now(datetime.timezone.utc).replace(tzinfo=None) + timedelta(minutes=2)
    
    new_otp = VSCodeOTP(
        otp_code=otp,
        user_id=user.id,
        expires_at=expires_at
    )
    db.session.add(new_otp)
    db.session.commit()
    
    return jsonify({
        'status': 'ok',
        'otp': otp,
        'expires_at': expires_at.isoformat()
    })

@app.route('/api/auth/consume-otp', methods=['POST'])
def vscode_consume_otp():
    """Consumes an OTP and returns a fresh JWT token for the browser session."""
    data = request.get_json(silent=True) or {}
    otp_code = data.get('otp')
    
    if not otp_code:
        return jsonify({'error': 'OTP is required'}), 400
        
    otp_record = VSCodeOTP.query.filter_by(otp_code=otp_code).first()
    
    if not otp_record:
        return jsonify({'error': 'Invalid or already used OTP'}), 404
        
    if otp_record.expires_at < datetime.datetime.now(datetime.timezone.utc).replace(tzinfo=None):
        db.session.delete(otp_record)
        db.session.commit()
        return jsonify({'error': 'OTP has expired'}), 403
        
    user = otp_record.user
    
    # Generate JWT for browser (Web Auth)
    access_token = create_access_token(identity=str(user.id))
    
    # Clean up IMMEDIATELY (one-time use)
    db.session.delete(otp_record)
    db.session.commit()
    
    return jsonify({
        'status': 'ok',
        'token': access_token,
        'user': {
            'id': user.id,
            'email': user.email,
            'display_name': user.display_name,
            'tokens': get_or_create_token_balance(user).balance
        }
    })

@app.route('/api/keys/<int:key_id>', methods=['DELETE'])
@jwt_required()
def revoke_api_key(key_id):
    user_id = get_jwt_identity()
    key_record = ApiKey.query.filter_by(id=key_id, user_id=user_id).first()
    if not key_record:
        return jsonify({'error': 'Key not found'}), 404
        
    # Soft delete - Veritabanı tutarlılığı için satırı silmiyoruz
    key_record.is_active = False
    db.session.commit()
    return jsonify({'message': 'Key revoked successfully'})


@app.route('/v1/cancel', methods=['POST'])
def cancel_request():
    """Signals that a specific request should be aborted."""
    data = request.get_json(silent=True) or {}
    request_id = data.get('request_id')
    if request_id:
        CANCELLED_REQUESTS[request_id] = True
        print(f"[CANCEL] Request {request_id} has been marked for cancellation.")
        return jsonify({'status': 'ok'})
    return jsonify({'error': 'request_id missing'}), 400


@app.route('/api/cancel', methods=['POST'])
@jwt_required()
def web_cancel_request():
    """Web-client cancellation endpoint."""
    data = request.get_json(silent=True) or {}
    request_id = data.get('request_id')
    if request_id:
        CANCELLED_REQUESTS[request_id] = True
        print(f"[CANCEL-WEB] Request {request_id} marked for cancellation.")
        return jsonify({'status': 'ok'})
    return jsonify({'error': 'request_id missing'}), 400


@app.route('/v1/ask', methods=['POST'])
def external_ask():
    """
    Endpoint for VS Code Extension and external tools.
    Expects header: X-API-Key
    Expects JSON body: { "question": "...", "code": "..." }
    """
    _vsc_balance = 0
    _vsc_cost = 0
    api_key_header = _extract_api_key_from_request()
    if not api_key_header:
        return jsonify({'error': 'X-API-Key header is missing'}), 401
        
    key_record = _find_api_key_record(api_key_header)
    if not key_record:
        return jsonify({'error': 'Invalid or revoked API Key'}), 401

    key_client = _detect_api_key_client(key_record.key)
    if key_client == 'web':
        return jsonify({'error': 'Web API keys cannot be used for /v1 endpoints. Use a vscode API key.'}), 403
    
    # Son kullanım tarihini güncelle
    key_record.last_used_at = _utcnow()
    db.session.commit()


    # Tolerate invalid/missing JSON to avoid generic 400 from Werkzeug
    data = request.get_json(silent=True)
    if data is None:
        try:
            raw = (request.data or b"").decode("utf-8", errors="replace").strip()
            data = json.loads(raw) if raw else {}
        except Exception:
            data = {}
    question = data.get('question', '')
    code = data.get('code', '')
    conversation_id = data.get('conversation_id')
    session_id = str(data.get('session_id') or '').strip()
    request_id = data.get('request_id') or str(uuid.uuid4())
    active_file = str(data.get('active_file') or data.get('file_path') or '').strip()
    requested_model = str(data.get('model') or GEMINI_MODEL.replace('models/', '')).strip()
    if requested_model == 'auto':
        requested_model = GEMINI_MODEL.replace('models/', '')
    agent_mode = _parse_bool(data.get('agent_mode'))
    has_stream_flag = 'stream' in data
    stream_flag = _parse_bool(data.get('stream')) if has_stream_flag else None
    payload_project_id = data.get('project_id')

    # ── 💰 Token Balance Check — Preliminary check moved down for accuracy ──────────────────
    workspace_files = _parse_workspace_files_payload(data.get('workspace_files'))
    history_context = data.get('history_context') if isinstance(data.get('history_context'), list) else []
    
    if not question and not code:
        return jsonify({'error': 'Either question or code must be provided'}), 400
        
    user_id = key_record.user_id
    user = User.query.get(user_id)
    prefs = json.loads(user.preferences) if user and user.preferences else {}

    # 📜 CONVERSATION LOOKUP — Handle numeric ID or Title fallback (Deprecated)
    _conv = None
    if conversation_id:
        # Try finding by integer ID first (Modern approach)
        if str(conversation_id).isdigit():
            _conv = db.session.get(Conversation, int(conversation_id))
            if _conv and _conv.user_id != user_id:
                _conv = None
        
        # Fallback to Title lookup (Legacy approach - slated for removal in v2.0)
        if not _conv:
            _conv = Conversation.query.filter_by(
                user_id=user_id,
                title=str(conversation_id),
                is_deleted=False
            ).first()
            if _conv:
                print(f"[ASK] DEPRECATION WARNING: User {user_id} used title-based lookup for conversation '{conversation_id}'. This will be removed in v2.0.")

    if not _conv and session_id:
        _conv = Conversation.query.filter_by(
            user_id=user_id,
            title=session_id,
            is_deleted=False
        ).first()

    if not _conv:
        conversation_title = session_id or (
            str(conversation_id)
            if conversation_id and not str(conversation_id).isdigit()
            else f"Chat {_utcnow().strftime('%Y-%m-%d %H:%M')}"
        )
        _conv = Conversation(
            user_id=user_id,
            title=conversation_title,
            project_id=payload_project_id if payload_project_id else None,
            source='vscode'
        )
        db.session.add(_conv)
        db.session.commit()
        print(f"[ASK] Created/Identified conversation: {_conv.id} for user {user_id}")

    # 🔄 SESSION STATE — Isolate per-request vs per-session
    # Persistent session state (shared across requests in same chat)
    session_state_key = _conversation_state_key(user_id, _conv.id)
    session_state = _external_conversation_state.get(session_state_key, {})
    
    # Transient request state (isolated to this specific rid)
    request_state = {
        'request_id': request_id,
        'start_time': _utcnow().isoformat(),
        'active_file': active_file or session_state.get('active_file')
    }

    if active_file:
        session_state['active_file'] = active_file
        session_state['last_action'] = 'active_file_supplied'

    carried_active_file = str(session_state.get('active_file') or '').strip()
    if not active_file and carried_active_file:
        active_file = carried_active_file

    if _is_short_confirmation_intent(question):
        target_file = active_file or carried_active_file
        if target_file:
            question = (
                f"{question}\n\n"
                f"[System execution instruction: User provided a short confirmation intent. "
                f"Use the last active file as target: {target_file}. "
                f"Proceed directly with modifications for this file and do not ask follow-up questions.]"
            )

            if not code:
                resolved_path, resolved_content = _resolve_workspace_file_content(workspace_files, target_file)
                if resolved_content:
                    code = f"Active file ({resolved_path}) content:\n{resolved_content}"

            session_state['active_file'] = target_file
            session_state['last_action'] = 'confirmation_intent_applied'
            request_state['active_file'] = target_file

    project_for_agent = None
    if payload_project_id:
        try:
            project_id_int = int(payload_project_id)
        except Exception:
            project_id_int = None
        if project_id_int:
            project_for_agent = Project.query.filter_by(id=project_id_int, user_id=user_id).first()

    provider_key, provider_model = _resolve_agent_provider_model(requested_model)
    if not provider_key:
        provider_key = 'gemini'
        provider_model = GEMINI_MODEL.replace('models/', '')

    if agent_mode and not _is_agent_model_supported(provider_key, provider_model):
        return jsonify({
            'error': f'Agent Mode bu modelde desteklenmiyor: {provider_model}',
            'agent_mode_blocked': True,
            'selected_model': provider_model,
            'supported_agent_models': list(SUPPORTED_GEMINI_AGENT_MODELS) + ['gemma-*', 'gpt-*', 'claude-*'],
        }), 400

    # ── 💰 Token Balance Check (STRICT & ACCURATE) ──────────────────
    _vsc_user = User.query.get(key_record.user_id)
    if not _vsc_user:
        return jsonify({'error': 'İstek için yetkili kullanıcı bulunamadı.'}), 401
    
    _has_tokens, _vsc_balance, _vsc_cost = check_tokens(_vsc_user, provider_model)
    print(f"[TOKEN] Pre-check: User={_vsc_user.id}, Balance={_vsc_balance}, Cost={_vsc_cost}, Model={provider_model}")
    
    if not _has_tokens:
        print(f"[TOKEN] BLOCKED: User {_vsc_user.id} has insufficient balance ({_vsc_balance} < {_vsc_cost})")
        return jsonify({
            'error': f'Yetersiz token bakiyesi. Mevcut: {_vsc_balance}, Gerekli: {_vsc_cost}.',
            'error_code': 'insufficient_balance',
            'balance': _vsc_balance,
            'required': _vsc_cost,
        }), 402

    workspace_root = data.get('workspace_root')
    
    tool_runtime = None
    if agent_mode:
        tool_runtime = AgentToolRuntime(
            project=project_for_agent,
            workspace_root=workspace_root,
            workspace_files=workspace_files,
            search_project_callback=_agent_project_search,
            invalidate_project_cache=invalidate_project_embedding_cache,
        )

    # Progress callback for tool execution (optional)
    on_event = None



    # Define a deduction flag and balance tracker
    deduction_occurred = False
    new_balance = _vsc_balance

    def trigger_token_deduction():
        nonlocal deduction_occurred, new_balance
        if deduction_occurred:
            return
        
        # We need a fresh user record for the transaction
        try:
            with app.app_context():
                db.session.rollback()
                charge_user = db.session.get(User, user_id)
                if charge_user:
                    print(f"[TOKEN-CALLBACK] Triggering deduction for user {charge_user.id}. Model: {provider_model}")
                    success, updated_balance = deduct_tokens(
                        charge_user,
                        provider_model,
                        description=f"VSC Agent (Turn): {question[:30]}...",
                        reference_id=request_id
                    )
                    if success:
                        new_balance = updated_balance
                        deduction_occurred = True
                        print(f"[TOKEN-CALLBACK] Deduction successful. New Balance: {new_balance}")
        except Exception as e:
            print(f"[TOKEN-CALLBACK] Error during deduction: {e}")

    # ── Determine if client wants streaming ────────────────────────────────
    if has_stream_flag:
        wants_stream = bool(stream_flag)
    else:
        wants_stream = (not agent_mode) and ('text/event-stream' in (request.headers.get('Accept') or '').lower())

    # ── Non-agent fast streaming path ──────────────────────────────────────
    # Skip run_agent_turn() entirely; stream directly from the provider.
    # This gives the first token as fast as the model can produce it.
    if wants_stream and not agent_mode:
        def _build_vsc_system_prompt():
            lang_hint = _build_language_hint_simple(question, prefs)
            persona_info = ""
            if prefs:
                persona = prefs.get('persona', 'General User')
                expertise = prefs.get('expertise', 'Intermediate')
                persona_info = f"User profile: {persona} (expertise: {expertise}). "
            return (
                "You are a senior software engineering assistant integrated into VS Code. "
                f"{persona_info}"
                f"{lang_hint} "
                "Be concise and practical. For code questions provide working examples. "
                "Never output internal reasoning labels or metadata."
            )

        def _build_vsc_messages():
            msgs = []
            for turn in (history_context or []):
                u = (turn.get('user') or '').strip()
                a = (turn.get('ai') or '').strip()
                if u: msgs.append({'role': 'user', 'content': u})
                if a: msgs.append({'role': 'assistant', 'content': a})
            user_msg = (question or '').strip() or 'Hello'
            if code and code.strip():
                user_msg += f"\n\nRelated code:\n```\n{code.strip()}\n```"
            msgs.append({'role': 'user', 'content': user_msg})
            return msgs

        def _gemini_stream_vsc(sys_prompt, msgs):
            """Yield raw text chunks from Gemini generate_content_stream."""
            gc = getattr(genai, '_client', None)
            if not gc:
                yield 'Error: Gemini client not configured.'
                return
            from google.genai import types as _gt
            _contents = []
            for m in msgs:
                role = 'model' if m['role'] == 'assistant' else 'user'
                _contents.append(_gt.Content(role=role, parts=[_gt.Part.from_text(text=m['content'])]))
            _cfg = _gt.GenerateContentConfig(
                system_instruction=sys_prompt or None,
                temperature=0.2,
                max_output_tokens=2048,
                http_options=_gt.HttpOptions(timeout=to_gemini_timeout(120)),
            )
            try:
                for item in gc.models.generate_content_stream(
                    model=provider_model, contents=_contents, config=_cfg
                ):
                    t = getattr(item, 'text', None) or ''
                    if t:
                        yield t
            except Exception as exc:
                yield f'\n[Gemini error: {exc}]'

        def _claude_stream_vsc(sys_prompt, msgs):
            """Yield raw text chunks from Claude streaming."""
            if not claude_client:
                yield 'Error: Anthropic client not configured.'
                return
            try:
                with claude_client.messages.stream(
                    model=provider_model,
                    max_tokens=2048,
                    system=sys_prompt,
                    messages=msgs,
                ) as s:
                    for text in s.text_stream:
                        yield text
            except Exception as exc:
                yield f'\n[Claude error: {exc}]'

        def _openai_stream_vsc(sys_prompt, msgs):
            """Yield raw text chunks from OpenAI streaming."""
            if not openai_client:
                yield 'Error: OpenAI client not configured.'
                return
            try:
                all_msgs = [{'role': 'system', 'content': sys_prompt}] + msgs
                resp = openai_client.chat.completions.create(
                    model=provider_model,
                    messages=all_msgs,
                    temperature=0.2,
                    max_completion_tokens=2048,
                    stream=True,
                )
                for chunk in resp:
                    delta = chunk.choices[0].delta if chunk.choices else None
                    t = getattr(delta, 'content', None) or ''
                    if t:
                        yield t
            except Exception as exc:
                yield f'\n[OpenAI error: {exc}]'

        def _get_chunk_generator(sys_prompt, msgs):
            if provider_key == 'gemini':
                return _gemini_stream_vsc(sys_prompt, msgs)
            if provider_key == 'anthropic':
                return _claude_stream_vsc(sys_prompt, msgs)
            if provider_key == 'openai':
                return _openai_stream_vsc(sys_prompt, msgs)
            return iter(['Unsupported provider: ' + provider_key])

        def _build_language_hint_simple(q, p):
            lang = (p or {}).get('preferred_language')
            if lang:
                return f'Always respond in {lang}.'
            import re as _re
            if _re.search(r'[\u00e7\u011f\u0131\u00f6\u015f\u00fc]', q or ''):
                return 'Respond in Turkish.'
            return 'Respond in the same language as the user.'

        def generate_stream_native():
            sys_prompt = _build_vsc_system_prompt()
            msgs = _build_vsc_messages()

            # Deduct tokens at stream start
            trigger_token_deduction()

            meta = {
                'meta': True,
                'agent_mode': False,
                'selected_model': provider_model,
                'agent_provider': provider_key,
                'conversation_id': final_conv_id,
                'balance': new_balance,
            }
            yield f"data: {json.dumps(meta)}\n\n"

            full_text = ''
            try:
                for chunk in _get_chunk_generator(sys_prompt, msgs):
                    if chunk:
                        full_text += chunk
                        yield f"data: {json.dumps({'text': chunk})}\n\n"
            except Exception as exc:
                yield f"data: {json.dumps({'text': f'[Stream error: {exc}]'})}\n\n"

            # Persist history after stream completes
            try:
                with app.app_context():
                    db.session.rollback()
                    _fresh_conv = db.session.get(Conversation, final_conv_id) if final_conv_id else None
                    if _fresh_conv:
                        _hist = History(
                            conversation_id=_fresh_conv.id,
                            user_question=question,
                            ai_response=full_text,
                            selected_model=provider_model,
                            timestamp=_utcnow()
                        )
                        db.session.add(_hist)
                        db.session.commit()
            except Exception as _he:
                print(f'[HISTORY] Stream save error: {_he}')

            yield f"data: {json.dumps({'done': True, 'steps': 0, 'agent_trace': [], 'agent_changed_files': []})}\n\n"
            yield 'data: [DONE]\n\n'

        return Response(stream_with_context(generate_stream_native()), mimetype='text/event-stream')

    # ── Agent path (or non-streaming fallback) ─────────────────────────────
    # For agent mode OR when client does not want streaming, use run_agent_turn().
    try:
        agent_result = run_agent_turn(
            provider=provider_key,
            model=provider_model,
            question=question,
            code=code,
            prefs=prefs,
            history_context=history_context,
            github_context='',
            tool_runtime=tool_runtime,
            openai_client=openai_client,
            anthropic_client=claude_client,
            gemini_client=getattr(genai, '_client', None),
            request_id=request_id,
            on_event=on_event,
            on_first_llm_success=trigger_token_deduction
        )
    except AgentAbortException:
        print(f"[ASK] Request {request_id} ABORTED by user. Ensuring final balance sync.")
        CANCELLED_REQUESTS.pop(request_id, None)
        db.session.rollback()
        _fresh_user = db.session.get(User, user_id)
        current_balance = get_or_create_token_balance(_fresh_user).balance if _fresh_user else _vsc_balance
        return jsonify({
            'error': 'Request cancelled by user',
            'error_code': 'request_cancelled',
            'balance': current_balance
        }), 499
    except Exception as e:
        print(f"[ASK] Error during agent run: {e}")
        db.session.rollback()
        _fresh_user = db.session.get(User, user_id)
        current_balance = get_or_create_token_balance(_fresh_user).balance if _fresh_user else _vsc_balance
        return jsonify({
            'error': f'Agent error: {str(e)}',
            'error_code': 'agent_error',
            'balance': current_balance
        }), 500

    clipped_agent_meta = _clip_agent_metadata(agent_result.trace, agent_result.changed_files)

    if agent_result and isinstance(agent_result.changed_files, list) and agent_result.changed_files:
        first_changed = agent_result.changed_files[0] if isinstance(agent_result.changed_files[0], dict) else {}
        changed_path = str(first_changed.get('path') or '').strip()
        if changed_path:
            session_state['active_file'] = changed_path
            session_state['last_action'] = 'patch_generated'

    session_state['updated_at'] = _utcnow().isoformat()
    _external_conversation_state[session_state_key] = session_state

    # 💰 TOKEN — Final sync
    if not deduction_occurred and _vsc_user:
        db.session.rollback()
        charge_user = db.session.get(User, user_id)
        if charge_user:
            success, new_balance = deduct_tokens(
                charge_user, provider_model,
                description=f"VSC Agent: {question[:30]}...",
                reference_id=None
            )
            CANCELLED_REQUESTS.pop(request_id, None)

    # Capture primitive IDs for session-safe access in generators
    final_conv_id = _conv.id if _conv else None
    final_proj_id = project_for_agent.id if project_for_agent else None
    
    # 📜 HISTORY PERSISTENCE
    try:
        with app.app_context():
            db.session.rollback()
            current_conv = db.session.get(Conversation, final_conv_id) if final_conv_id else None
            
            if current_conv and final_proj_id and not current_conv.project_id:
                current_conv.project_id = final_proj_id
                db.session.add(current_conv)
                
            if not current_conv:
                current_conv = Conversation(
                    user_id=user_id,
                    title=str(conversation_id) if conversation_id else f"Chat {_utcnow().strftime('%Y-%m-%d %H:%M')}",
                    project_id=final_proj_id,
                    source='extension'
                )
                db.session.add(current_conv)
                db.session.flush()
                final_conv_id = current_conv.id
                
            _hist = History(
                conversation_id=current_conv.id,
                user_question=question,
                ai_response=agent_result.text,
                selected_model=provider_model,
                timestamp=_utcnow()
            )
            db.session.add(_hist)
            db.session.commit()
    except Exception as _e:
        db.session.rollback()
        print(f"[HISTORY] Error saving turn for user {user_id}: {_e}")

    # Agent mode always returns SSE (meta + done, no streaming text)
    if agent_mode:
        def generate_agent():
            try:
                meta = {
                    'meta': True,
                    'agent_mode': True,
                    'selected_model': provider_model,
                    'agent_provider': provider_key,
                    'agent_project_id': final_proj_id,
                    'agent_tool_capable': bool(agent_result.tool_capable),
                    'balance': new_balance,
                    'conversation_id': final_conv_id,
                }
                yield f"data: {json.dumps(meta)}\n\n"
                final_text = (agent_result.text or '').strip() or 'No response generated.'
                for chunk in stream_text_chunks(final_text):
                    yield f"data: {json.dumps({'text': chunk})}\n\n"
                done_payload = {
                    'done': True,
                    'steps': len(clipped_agent_meta['trace'] or []),
                    'agent_trace': clipped_agent_meta['trace'],
                    'agent_changed_files': clipped_agent_meta['changed_files'],
                    'agent_trace_total': clipped_agent_meta['trace_total'],
                    'agent_changed_total': clipped_agent_meta['changed_total'],
                    'agent_trace_truncated': clipped_agent_meta['trace_truncated'],
                    'agent_changed_truncated': clipped_agent_meta['changed_truncated'],
                }
                yield f"data: {json.dumps(done_payload)}\n\n"
                yield 'data: [DONE]\n\n'
            except Exception as e:
                yield f"data: {json.dumps({'text': f'[Backend error: {str(e)}]'})}\n\n"
                yield 'data: [DONE]\n\n'
        return Response(stream_with_context(generate_agent()), mimetype='text/event-stream')

    # Non-streaming JSON fallback
    return jsonify({
        'answer': agent_result.text,
        'steps': len(clipped_agent_meta['trace'] or []),
        'agent_mode': bool(agent_mode),
        'selected_model': provider_model,
        'agent_provider': provider_key,
        'agent_project_id': project_for_agent.id if project_for_agent else None,
        'agent_tool_capable': bool(agent_result.tool_capable),
        'active_file': session_state.get('active_file') if isinstance(session_state, dict) else None,
        'agent_trace': clipped_agent_meta['trace'],
        'agent_changed_files': clipped_agent_meta['changed_files'],
        'agent_trace_total': clipped_agent_meta['trace_total'],
        'agent_changed_total': clipped_agent_meta['changed_total'],
        'agent_trace_truncated': clipped_agent_meta['trace_truncated'],
        'agent_changed_truncated': clipped_agent_meta['changed_truncated'],
        'balance': new_balance,
        'conversation_id': _conv.id if _conv else None,
    })

@app.route('/v1/history', methods=['GET'])
def get_vscode_history():
    """
    Get conversation list for the authenticated VS Code user.
    Only returns sessions with titles starting with 'session-' (VS Code sessions).
    """
    api_key_header = _extract_api_key_from_request()
    if not api_key_header:
        return jsonify({'error': 'X-API-Key header is missing'}), 401
        
    key_record = _find_api_key_record(api_key_header)
    if not key_record:
        return jsonify({'status': 'unauthorized'}), 401
    
    user_id = key_record.user_id
    # VS Code sessions now identified by source field
    convs = Conversation.query.filter(
        Conversation.user_id == user_id,
        Conversation.source == 'vscode',
        Conversation.is_deleted == False
    ).order_by(Conversation.created_at.desc()).all()
    
    return jsonify({
        'status': 'ok',
        'sessions': [{
            'id': c.id,
            'title': c.title,
            'updatedAt': c.created_at.isoformat(),
            'pinned': c.is_pinned
        } for c in convs]
    })

@app.route('/v1/history/<string:conv_title>', methods=['GET'])
def get_conversation_details(conv_title):
    """Get full history for a specific VS Code session title."""
    api_key_header = _extract_api_key_from_request()
    if not api_key_header:
        return jsonify({'error': 'X-API-Key header is missing'}), 401
        
    key_record = _find_api_key_record(api_key_header)
    if not key_record:
        return jsonify({'status': 'unauthorized'}), 401
        
    user_id = key_record.user_id
    conv = Conversation.query.filter_by(user_id=user_id, title=conv_title, is_deleted=False).first()
    if not conv:
        return jsonify({'error': 'Conversation not found'}), 404
        
    history_items = History.query.filter_by(conversation_id=conv.id, is_deleted=False).order_by(History.timestamp.asc()).all()
    
    messages = []
    for h in history_items:
        messages.append({
            'role': 'user',
            'text': h.user_question,
            'createdAt': h.timestamp.isoformat()
        })
        if h.ai_response:
            messages.append({
                'role': 'ai',
                'text': h.ai_response,
                'createdAt': h.timestamp.isoformat()
            })
            
    return jsonify({
        'status': 'ok',
        'messages': messages
    })

@app.route('/v1/status', methods=['GET'])
def external_status():
    """
    Status endpoint for VS Code Extension.
    Requires X-API-Key for authentication.
    Returns detailed system status and user context.
    """
    api_key_header = _extract_api_key_from_request()
    if not api_key_header:
        return jsonify({'error': 'X-API-Key header is missing'}), 401
        
    key_record = _find_api_key_record(api_key_header)
    if not key_record:
        return jsonify({'status': 'unauthorized', 'error': 'Invalid or revoked API Key'}), 401

    key_client = _detect_api_key_client(key_record.key)
    if key_client == 'web':
        return jsonify({'status': 'unauthorized', 'error': 'Web API keys cannot be used for /v1 endpoints.'}), 403
    
    user = User.query.get(key_record.user_id)
    
    # Ensure TokenBalance record exists (SaaS resilience)
    if user:
        get_or_create_token_balance(user)
    
    # Dynamically resolve purchase URL based on configuration or current host
    base_url = _get_iyzico_frontend_base_url()
    purchase_url = f"{base_url}/billing"

    # Diagnostic logging for VS Code connection
    reported_balance = get_or_create_token_balance(user).balance if user else 0
    print(f"[VSCODE] FINAL SYNC: {user.email if user else 'N/A'} | Balance: {reported_balance} | URL: {purchase_url}")

    return jsonify({
        'status': 'ok',
        'version': '1.0.0',
        'user': user.email if user else 'unknown',
        'auth_state': 'authenticated',
        'balance': reported_balance,
        'purchase_url': purchase_url,
        'model_config': {
            'gemini': GEMINI_MODEL,
            'openai': OPENAI_MODEL,
            'anthropic': ANTHROPIC_MODEL
        },
        'server_time': _utcnow().isoformat()
    })

# ============================================================
# 🔐 ADMIN — QUOTA & TOKEN MANAGEMENT ENDPOINTS
# ============================================================

def _require_admin():
    """Returns (user, error_response) — error_response is None if admin check passes."""
    user = get_current_user()
    if not user:
        return None, (jsonify({'error': 'Authentication required.'}), 401)
    if not user.is_admin:
        return None, (jsonify({'error': 'Admin privileges required.'}), 403)
    return user, None


@app.route('/api/admin/quota/defaults', methods=['GET'])
@jwt_required()
def admin_get_quota_defaults():
    """Global quota varsayılanlarını döndürür (tüm kullanıcılar için default limit)."""
    _, err = _require_admin()
    if err:
        return err

    return jsonify({
        'signup_grant_tokens': SIGNUP_GRANT_TOKENS,
        'default_daily_limit': 200,
        'default_weekly_limit': 1000,
        'monthly_grant_tokens': MONTHLY_GRANT_TOKENS,
    })


@app.route('/api/admin/quota/defaults', methods=['PUT'])
@jwt_required()
def admin_update_quota_defaults():
    """
    Global quota varsayılanlarını günceller.
    NOT: Bu endpoint mevcut kullanıcı cüzdanlarını TOPLU olarak günceller.
    Body: { daily_limit?: int, weekly_limit?: int }
    """
    _, err = _require_admin()
    if err:
        return err

    data = request.get_json(silent=True) or {}
    daily_limit = data.get('daily_limit')
    weekly_limit = data.get('weekly_limit')

    if daily_limit is None and weekly_limit is None:
        return jsonify({'error': 'En az bir alan belirtilmeli: daily_limit veya weekly_limit'}), 400

    updated = 0
    try:
        wallets = TokenBalance.query.all()
        for wallet in wallets:
            changed = False
            if daily_limit is not None:
                wallet.daily_limit = int(daily_limit)
                changed = True
            if weekly_limit is not None:
                wallet.weekly_limit = int(weekly_limit)
                changed = True
            if changed:
                updated += 1
        db.session.commit()
        return jsonify({
            'success': True,
            'updated_wallets': updated,
            'new_daily_limit': daily_limit,
            'new_weekly_limit': weekly_limit,
        })
    except Exception as e:
        db.session.rollback()
        print(f"admin_update_quota_defaults error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/admin/users', methods=['GET'])
@jwt_required()
def admin_list_users():
    """Tüm kullanıcıları token cüzdanı bilgileriyle listeler."""
    _, err = _require_admin()
    if err:
        return err

    page = max(1, int(request.args.get('page', 1)))
    per_page = max(1, min(int(request.args.get('per_page', 50)), 200))
    search = (request.args.get('search') or '').strip()

    try:
        query = User.query
        if search:
            like = f'%{search}%'
            query = query.filter(
                (User.email.ilike(like)) | (User.display_name.ilike(like))
            )

        paginated = query.order_by(User.id.desc()).paginate(page=page, per_page=per_page, error_out=False)
        users_data = []
        for u in paginated.items:
            wallet = TokenBalance.query.filter_by(user_id=u.id).first()
            users_data.append({
                'id': u.id,
                'email': u.email,
                'display_name': u.display_name,
                'is_admin': u.is_admin,
                'created_at': u.created_at.isoformat() if u.created_at else None,
                'token_balance': {
                    'balance': wallet.balance if wallet else 0,
                    'total_spent': wallet.total_spent if wallet else 0,
                    'daily_limit': wallet.daily_limit if wallet else 200,
                    'daily_used': wallet.daily_used if wallet else 0,
                    'weekly_limit': wallet.weekly_limit if wallet else 1000,
                    'weekly_used': wallet.weekly_used if wallet else 0,
                    'monthly_renewal_enabled': wallet.monthly_renewal_enabled if wallet else False,
                    'monthly_renewal_day': wallet.monthly_renewal_day if wallet else None,
                    'last_renewal_at': wallet.last_renewal_at.isoformat() if wallet and wallet.last_renewal_at else None,
                } if wallet else None,
            })

        return jsonify({
            'users': users_data,
            'total': paginated.total,
            'page': page,
            'per_page': per_page,
            'pages': paginated.pages,
        })
    except Exception as e:
        print(f"admin_list_users error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/admin/users/<int:target_user_id>/quota', methods=['GET'])
@jwt_required()
def admin_get_user_quota(target_user_id):
    """Belirli kullanıcının kota ve token bilgisini döndürür."""
    _, err = _require_admin()
    if err:
        return err

    user = db.session.get(User, target_user_id)
    if not user:
        return jsonify({'error': 'Kullanıcı bulunamadı.'}), 404

    wallet = get_or_create_token_balance(user)
    if wallet.weekly_reset_at is None or wallet.daily_reset_at is None:
        init_quota_for_new_user(wallet)

    return jsonify({
        'user_id': user.id,
        'email': user.email,
        'display_name': user.display_name,
        'balance': wallet.balance,
        'total_spent': wallet.total_spent,
        'daily_limit': wallet.daily_limit,
        'daily_used': wallet.daily_used,
        'daily_reset_at': wallet.daily_reset_at.isoformat() if wallet.daily_reset_at else None,
        'weekly_limit': wallet.weekly_limit,
        'weekly_used': wallet.weekly_used,
        'weekly_reset_at': wallet.weekly_reset_at.isoformat() if wallet.weekly_reset_at else None,
        'monthly_renewal_enabled': wallet.monthly_renewal_enabled,
        'monthly_renewal_day': wallet.monthly_renewal_day,
        'last_renewal_at': wallet.last_renewal_at.isoformat() if wallet.last_renewal_at else None,
        'updated_at': wallet.updated_at.isoformat() if wallet.updated_at else None,
    })


@app.route('/api/admin/users/<int:target_user_id>/quota', methods=['PUT'])
@jwt_required()
def admin_update_user_quota(target_user_id):
    """
    Belirli kullanıcının kota limitlerini günceller.
    Body: {
        daily_limit?: int,
        weekly_limit?: int,
        monthly_renewal_enabled?: bool,
        monthly_renewal_day?: int (1-28) | null
    }
    """
    _, err = _require_admin()
    if err:
        return err

    user = db.session.get(User, target_user_id)
    if not user:
        return jsonify({'error': 'Kullanıcı bulunamadı.'}), 404

    data = request.get_json(silent=True) or {}

    wallet = get_or_create_token_balance(user)

    try:
        changed_fields = []
        if 'daily_limit' in data:
            val = int(data['daily_limit'])
            if val < 0:
                return jsonify({'error': 'daily_limit negatif olamaz.'}), 400
            wallet.daily_limit = val
            changed_fields.append(f'daily_limit={val}')

        if 'weekly_limit' in data:
            val = int(data['weekly_limit'])
            if val < 0:
                return jsonify({'error': 'weekly_limit negatif olamaz.'}), 400
            wallet.weekly_limit = val
            changed_fields.append(f'weekly_limit={val}')

        if 'monthly_renewal_enabled' in data:
            wallet.monthly_renewal_enabled = bool(data['monthly_renewal_enabled'])
            changed_fields.append(f'monthly_renewal_enabled={wallet.monthly_renewal_enabled}')

        if 'monthly_renewal_day' in data:
            rday = data['monthly_renewal_day']
            if rday is not None:
                rday = int(rday)
                if not (1 <= rday <= 28):
                    return jsonify({'error': 'monthly_renewal_day 1-28 arasında olmalı.'}), 400
            wallet.monthly_renewal_day = rday
            changed_fields.append(f'monthly_renewal_day={rday}')

        if not changed_fields:
            return jsonify({'error': 'Güncellenecek alan bulunamadı.'}), 400

        db.session.commit()
        print(f"[ADMIN] Quota updated for user {user.id} ({user.email}): {', '.join(changed_fields)}")

        return jsonify({
            'success': True,
            'user_id': user.id,
            'email': user.email,
            'daily_limit': wallet.daily_limit,
            'weekly_limit': wallet.weekly_limit,
            'monthly_renewal_enabled': wallet.monthly_renewal_enabled,
            'monthly_renewal_day': wallet.monthly_renewal_day,
        })
    except Exception as e:
        db.session.rollback()
        print(f"admin_update_user_quota error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/admin/users/<int:target_user_id>/grant-tokens', methods=['POST'])
@jwt_required()
def admin_grant_tokens(target_user_id):
    """
    Admin olarak belirli kullanıcıya token yükler.
    Body: { amount: int, description?: str }
    """
    admin_user, err = _require_admin()
    if err:
        return err

    user = db.session.get(User, target_user_id)
    if not user:
        return jsonify({'error': 'Kullanıcı bulunamadı.'}), 404

    data = request.get_json(silent=True) or {}
    amount = data.get('amount')
    description = data.get('description', f'Admin tarafından eklendi — {admin_user.display_name}')

    if not amount or int(amount) <= 0:
        return jsonify({'error': 'Geçerli bir miktar belirtin (pozitif tam sayı).'}), 400

    amount = int(amount)
    try:
        wallet = get_or_create_token_balance(user)
        wallet.balance += amount

        tx = TokenTransaction(
            user_id=user.id,
            amount=amount,
            type='bonus',
            description=description,
            reference_id=f'admin:{admin_user.id}',
        )
        db.session.add(tx)
        db.session.commit()

        print(f"[ADMIN] {admin_user.email} granted {amount} tokens to {user.email}. New balance: {wallet.balance}")

        return jsonify({
            'success': True,
            'user_id': user.id,
            'email': user.email,
            'granted': amount,
            'new_balance': wallet.balance,
        })
    except Exception as e:
        db.session.rollback()
        print(f"admin_grant_tokens error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/admin/users/<int:target_user_id>/reset-quota', methods=['POST'])
@jwt_required()
def admin_reset_user_quota(target_user_id):
    """
    Belirli kullanıcının günlük/haftalık kota sayaçlarını sıfırlar.
    Body: { reset_daily?: bool, reset_weekly?: bool }  (default her ikisi de true)
    """
    _, err = _require_admin()
    if err:
        return err

    user = db.session.get(User, target_user_id)
    if not user:
        return jsonify({'error': 'Kullanıcı bulunamadı.'}), 404

    data = request.get_json(silent=True) or {}
    reset_daily = data.get('reset_daily', True)
    reset_weekly = data.get('reset_weekly', True)

    try:
        wallet = get_or_create_token_balance(user)
        if reset_daily:
            wallet.daily_used = 0
            wallet.daily_reset_at = calculate_daily_reset_time()
        if reset_weekly:
            wallet.weekly_used = 0
            wallet.weekly_reset_at = calculate_weekly_reset_time()
        db.session.commit()

        return jsonify({
            'success': True,
            'user_id': user.id,
            'daily_used': wallet.daily_used,
            'weekly_used': wallet.weekly_used,
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/api/admin/users/<int:target_user_id>/force-renewal', methods=['POST'])
@jwt_required()
def admin_force_renewal(target_user_id):
    """
    Belirli kullanıcı için aylık yenilemeyi hemen tetikler.
    Kullanıcının aktif auto_renew purchase'ı varsa renewal yapar.
    """
    admin_user, err = _require_admin()
    if err:
        return err

    user = db.session.get(User, target_user_id)
    if not user:
        return jsonify({'error': 'Kullanıcı bulunamadı.'}), 404

    try:
        renewed = check_and_apply_monthly_renewal(user.id)
        wallet = TokenBalance.query.filter_by(user_id=user.id).first()
        return jsonify({
            'success': True,
            'renewed': renewed,
            'new_balance': wallet.balance if wallet else 0,
            'message': 'Yenileme yapıldı.' if renewed else 'Yenileme zamanı gelmedi veya auto_renew kapalı.',
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/admin/stats', methods=['GET'])
@jwt_required()
def admin_stats():
    """Genel platform istatistiklerini döndürür."""
    _, err = _require_admin()
    if err:
        return err

    try:
        total_users = User.query.count()
        total_tokens_balance = db.session.query(db.func.sum(TokenBalance.balance)).scalar() or 0
        total_tokens_spent = db.session.query(db.func.sum(TokenBalance.total_spent)).scalar() or 0
        total_conversations = Conversation.query.filter_by(is_deleted=False).count()
        total_purchases = TokenPurchase.query.filter_by(status='completed').count()

        return jsonify({
            'total_users': total_users,
            'total_tokens_balance': int(total_tokens_balance),
            'total_tokens_spent': int(total_tokens_spent),
            'total_conversations': total_conversations,
            'total_completed_purchases': total_purchases,
        })
    except Exception as e:
        print(f"admin_stats error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/')
def serve_index():
    return send_from_directory(app.static_folder, 'index.html')

@app.errorhandler(404)
def handle_404(e):
    # Allow API 404s to remain as JSON
    if request.path.startswith('/api/') or request.path.startswith('/v1/'):
        return jsonify({
            "error": "Not Found",
            "details": "The requested API endpoint was not found on this server."
        }), 404
        
    # All other paths (like /tokens) should fallback to index.html for React Router
    return send_from_directory(app.static_folder, 'index.html')

# ==========================================
# SOCKET.IO EVENT HANDLERS — Live Collaboration
# ==========================================

@socketio.on('join_room')
def handle_join_room(data):
    """Kullanıcı, collab token'ı ile odaya katılıyor."""
    from models import SharedSession
    token = data.get('token', '')
    user_name = data.get('user_name', 'Guest')
    if not token:
        return
    with app.app_context():
        session = SharedSession.query.filter_by(share_token=token, is_active=True).first()
        if not session:
            socket_emit('error', {'message': 'Geçersiz collaboration linki'})
            return
    join_room(token)
    # Odadaki herkese bildirim
    socket_emit('user_joined', {
        'user_name': user_name,
        'token': token
    }, room=token)
    print(f'[Socket] {user_name} joined room {token[:8]}...')


@socketio.on('leave_room')
def handle_leave_room(data):
    """Kullanıcı odadan ayrılıyor."""
    token = data.get('token', '')
    user_name = data.get('user_name', 'Guest')
    if token:
        leave_room(token)
        socket_emit('user_left', {
            'user_name': user_name,
            'token': token
        }, room=token)
        print(f'[Socket] {user_name} left room {token[:8]}...')


@socketio.on('connect')
def handle_connect():
    print(f'[Socket] Client connected: {request.sid}')


@socketio.on('disconnect')
def handle_disconnect():
    print(f'[Socket] Client disconnected: {request.sid}')


# ==========================================
# MAIN ENTRY POINT
# ==========================================

with app.app_context():
    db.create_all()

if __name__ == '__main__':

    # Render uses the PORT environment variable
    port = int(os.environ.get("PORT", 5000))
    print(f'Starting SocketIO server (threading mode) on port {port}...')
    # socketio.run() threading async_mode ile WebSocket destekler
    socketio.run(app, host='0.0.0.0', port=port, debug=False, allow_unsafe_werkzeug=True)
